# =============================================================================
# PCA HEDGE FINDER — complete source, v104 (Fixes 1-110)
# =============================================================================
# Built 30 Jul 2026; v104 pass 1 Aug 2026.
#
# ── v104 — READABILITY, RE-ENTRY, AND TWO STATISTICAL QUESTIONS ─────────────
# [Fix 104] EVERY BLOCK IS A TABLE, AND EVERY NUMBER CARRIES ITS MEANING.
#           The report mixed three renderers — styled HTML tables, bare
#           print() key-value lines, and print() rules — so in a notebook half
#           of it was a table and half was raw <pre> at the browser's default
#           monospace. The ticket lines were the worst: a 26-char label and
#           then an unbroken run of numbers, units and conditional clauses.
#           Everything now goes through _section / _kv_block / _print_table
#           with one font stack, and the dense ticket lines (SIZING, RIGHT
#           NOW, n-day HOLDS, HORIZON, refits, MONITOR) gained a WHAT IT MEANS
#           column that states the DECISION, not just the measurement.
# [Fix 105] WHY ROLLING R² GOES NEGATIVE WHILE TE LOOKS FINE. Reported on an
#           equal-weighted AAPL/TSLA/SK-Hynix basket: rolling R² ≈ -0.4 for
#           one stretch, +0.8 for the next. Not a bug — an asymmetry in the
#           standard definition. R² = 1 - Σresid²/Σ(y-ȳ)² removes the mean
#           from the DENOMINATOR but not the numerator, so directional drift
#           is charged in full; TE uses .std(), which removes it. A hedge that
#           halves the vol but drifts one target-sigma scores R² = -0.25 with
#           a perfectly healthy TE. Both are now computed and shown side by
#           side (strict R², variance-only R², and the drift penalty that is
#           exactly the gap), with a count of how many negative windows still
#           cut variance. On the synthetic reproduction: 9 of 13.
#           [Fix 105] also adds the HEDGEABLE CEILING — the share of the
#           target no combination of the candidates can span. For a
#           concentrated book that ceiling, not the model, is what limits R².
# [Fix 106] ONE ACTION AT A TIME. Clicking RUN or DOWNLOAD twice queued a
#           second full run (and a second Bloomberg download). Two causes:
#           nothing was disabled, and _run_engine buffered the engine's stdout
#           to the END of the run, so a multi-minute job showed a blank cell
#           and clicking again was reasonable. Now: single-flight guard,
#           buttons grey out, output streams live, elapsed time reported.
# [Fix 107] The form renders at the end of CELL 1 — no second cell.
# [Fix 108] RECENCY. The composite averages OOS R² over ALL folds, so a config
#           that decayed could still win on the mean; deterioration was only
#           WARNED about afterwards. Now every config carries its recent-fold
#           score and its per-fold decay slope, [1b] reports what a
#           recent-folds-only selection WOULD have picked, and recency_weight
#           lets the selection act on it. Default stays 0 on purpose — see the
#           reasoning at the code site.
# [Fix 109] ALTERNATIVES get the pick's full evidence: exact executable
#           weights, unseen-window OOS, rolling TE/R² through time, and the
#           recency split — so a challenger can actually be judged.
# [Fix 110] A SETTINGS table stating what each knob affected, including the
#           two that are routinely misread: book size is DISPLAY ONLY, and
#           hedge duration changes the REPORT but not the weights.
#
# ── v104.1 — THE DEFAULTS AND THE HORIZON, MEASURED ─────────────────────────
# Both were open questions rather than reported faults, so they were tested
# rather than argued. Synthetic worlds, paired seeds, honest nested OOS.
#
# [Fix 111] The horizon variance ratio may RAISE the reported risk, never
#           lower it. VR was clipped to [0.25, 4.0] and applied both ways, so
#           a VR of 0.25 HALVED the reported h-day risk. Measured, VR has
#           essentially no out-of-sample persistence — corr(VR_in, VR_out)
#           = -0.32..+0.18 — while the true ratio matters enormously
#           (corr(VR_out, realised) = +0.54..+0.86). Shrinking a risk number
#           on an estimate with no forecasting power is not defensible;
#           raising one is merely conservative. Now floored at 1.0.
# [Fix 112] The horizon stays OUT of the selection — now for a stronger
#           reason than [Fix 100] had. A horizon-scaled criterion (daily TE x
#           sqrt(VR_h)) needs no extra data, so it dodges [Fix 100]'s
#           objection entirely. Built and measured against realised
#           non-overlapping h-day TE: it is WORSE than plain daily selection
#           at h = 21/63/126, in both worlds, winning only 4-11 of 25 seeds,
#           and the damage GROWS with the horizon. Selecting on an
#           unforecastable quantity maximises its estimation error. Full
#           numbers at the code site.
# [Fix 114] REVIEW PASS on v104/104.1 — three defects found in my own work:
#           (a) _KVCTX is module-level, and nothing reset it. A run that
#               raised part-way through the report left its buffered rows
#               behind, and since the UI catches exceptions and lets you
#               re-run, the NEXT report opened with rows from the failed one.
#               Now reset at entry, and the two diagnostic try/excepts flush
#               rather than abandon their buffer.
#           (b) the [Fix 108] slope used groupby.apply(include_groups=False),
#               which only exists from pandas 2.2 — on an older kernel a
#               purely DIAGNOSTIC column would have taken the whole selection
#               down with a TypeError. Replaced with an explicit loop.
#           (c) the SHORT-SAMPLE path (sample_window='3m' etc.) still had its
#               own print()-based title/kv, so the one documented workflow
#               most likely to be read by someone in a hurry still produced
#               the packed monospace report [Fix 104] existed to remove — and
#               it is the tier where the reader most needs to be told what a
#               number means, because none of them are validated. Rebuilt on
#               the shared renderer with implications.
#           Also: both new diagnostics now say so when they fail instead of
#           vanishing. Measured and left alone: the per-alternative
#           walk-forwards cost +2.5s on a 254-name universe — cheap enough
#           that shrinking their window count (which would make them
#           incomparable to the pick) is not worth it.
# [Fix 113] The lookback grid is now data-aware. Measured: adding a 504-row
#           lookback is worth +0.015 nested R² (t=+3.2) on a concentrated
#           target — but ONLY with 4+ years of data. At the default 3-year
#           download it is chosen 0 times in 20 because it cannot be trained
#           before each validation window. The report says which case you are
#           in. Related: history beyond the walk-forward's fixed budget is
#           never read at all (2.5y and 5y give bit-identical picks), so a
#           longer LOOKBACK is the only thing more history buys.
#
# HOW TO USE (notebook): paste the whole file into ONE cell and run it. The
# form comes up by itself — there is no second cell any more ([Fix 107]).
# Set HEDGE_UI_AUTOSTART = False above it if you want the scripted path only,
# or call hedge_ui() by hand at any time.
#
# EXPORTING A HEDGE  [Fix 103]
#   The Export button (and export_result() from code) writes the finished
#   hedge as a WORD DOCUMENT — verdict, the trade as a table, risk and
#   monitoring levels, out-of-sample evidence, the exact settings, and the
#   full text report as an appendix — plus weights .csv for execution.
#   Destination: HEDGE_EXPORT_DIR, set in CELL 1 and editable in the form's
#   Data tab. Default:
#         G:\FIN_COMM\DeltaOne\Kenny\PCA
#   Needs `pip install python-docx`; without it a .txt report is written
#   instead so an export never silently fails. If the drive is not mapped
#   (or the notebook runs off Windows) it falls back to ./hedge_reports and
#   says so — note a Windows path is a LEGAL filename on macOS/Linux, so the
#   exporter checks the platform rather than trusting makedirs.
#
# v96-v102 statistical work, both of which overturned an earlier claim of
# mine and are documented at the code sites:
#   [Fix 97]  The "is it working now" check used to fire on the single
#             latest reading (TE > 1.25x its median). Measured on stable
#             synthetic hedges it downgraded 17% of them — one good hedge in
#             six, for nothing. It now scores a BLOCK of recent readings
#             against the hedge's own historical distribution and requires
#             persistence plus two independent signals: 0% false alarms,
#             92% detection of hedges that genuinely broke.
#   [Fix 100] hedge_horizon_days no longer re-cuts the validation windows.
#             Tested (train 600 rows, score the same weights on held-out
#             63-row blocks, 8 paired trials per regime): iid residuals
#             16.39% vs 16.43% (t=+0.44), AR(1) 0.6 residuals 20.57% vs
#             20.58% (t=+0.09) — no detectable benefit, and it costs
#             evidence. A hedge ratio is a covariance property, so daily
#             data targets the same quantity with ~60x more observations.
#             The horizon now scales reported risk, reports what holds of
#             that length actually delivered, and checks the refit need.
#             horizon_sets_windows=True restores the old behaviour.
#
# Verification (synthetic; Bloomberg not required): 65-check feature suite,
# 33-check adversarial suite, window-isolation proof (corrupting all
# pre-window history leaves basket and weights bit-identical), now-cast
# error-rate measurement, the horizon experiment, and a 21-check export
# suite. All passing. Live Bloomberg request flags still need one download
# run against a terminal.
# =============================================================================

# =============================================================================
# CELL 1: SETUP (run once)
# =============================================================================

import numpy as np
import pandas as pd
import blpapi
from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import RidgeCV
from statsmodels.tsa.stattools import adfuller
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
import matplotlib.dates as _mdates
import warnings
import re
import html as _html
import textwrap as _textwrap

warnings.filterwarnings('ignore')

RIDGE_ALPHAS = np.logspace(-4, 2, 25)

# ─────────────────────────────────────────────────────────────────────────────
# INDEX MAP: (market_code, sector_keyword) → Bloomberg Index Ticker
# [Fix 19] The key is the LISTING MARKET of the target (derived from its
# exchange code), except for ADRs where country-of-risk is used as fallback.
# ─────────────────────────────────────────────────────────────────────────────
INDEX_MAP = {
    # US — official S&P 500 sector indices (INDX_MEMBERS friendly)
    ('US', 'Broad'):                 'SPX Index',
    ('US', 'Technology'):            'NDX Index',
    ('US', 'Communication'):         'S5TELS Index',
    ('US', 'Financials'):            'S5FINL Index',
    ('US', 'Health Care'):           'S5HLTH Index',
    ('US', 'Energy'):                'S5ENRS Index',
    ('US', 'Consumer Discretionary'): 'S5COND Index',
    # CN — used for onshore listings AND as the ADR fallback
    # (US-listed China names → HXC, the Nasdaq Golden Dragon index)
    ('CN', 'Broad'):                 'SHSZ300 Index',   # CSI 300 (onshore)
    ('CN', 'Technology'):            'HXC Index',
    ('CN', 'Communication'):         'HXC Index',
    ('CN', 'Consumer Discretionary'): 'HXC Index',
    # HK — HK-listed names (incl. CN-domiciled ones like Tencent/Meituan)
    ('HK', 'Broad'):                 'HSI Index',       # Hang Seng
    ('HK', 'Technology'):            'HSTECH Index',    # Hang Seng Tech
    ('HK', 'Communication'):         'HSTECH Index',
    ('HK', 'Consumer Discretionary'): 'HSTECH Index',
    ('HK', 'Financials'):            'HSI Index',
    # Europe
    ('EU', 'Broad'):                 'SXXP Index',      # Stoxx Europe 600
    ('DE', 'Broad'):                 'DAX Index',
    ('FR', 'Broad'):                 'CAC Index',
    ('GB', 'Broad'):                 'UKX Index',
    # Japan — TOPIX (broader, cap-weighted) instead of Nikkei
    ('JP', 'Broad'):                 'TPX Index',
    ('JP', 'Technology'):            'TPX Index',
    ('JP', 'Financials'):            'TPNBNK Index',
    # Korea, India, Australia, Taiwan
    ('KR', 'Broad'):                 'KOSPI2 Index',
    ('KR', 'Technology'):            'KOSPI2 Index',
    ('IN', 'Broad'):                 'NIFTY Index',
    ('IN', 'Technology'):            'NIFTY Index',
    ('IN', 'Financials'):            'NIFTY Index',
    ('AU', 'Broad'):                 'AS51 Index',
    # [Fix 37] TW50 INDX_MEMBERS is licence-gated on many terminals and
    # returns 0 members — which silently stripped every TAIEX large cap
    # (2330 TT included) from the universe. Values may now be a tuple of
    # fallbacks, tried in order until one returns members.
    ('TW', 'Broad'):                 ('TW50 Index', 'TWSE Index'),
    ('TW', 'Technology'):            ('TW50 Index', 'TWSE Index'),
}

# ─────────────────────────────────────────────────────────────────────────────
# CONTROL ETFS by market (added to the universe as benchmark hedges)
# Each market lists US-listed wrappers AND locally-listed ETFs. Locally-listed
# ETFs trade in the same timezone and currency as local targets, so they avoid
# the async-trading bias and FX mismatch that US-listed wrappers carry.
# ─────────────────────────────────────────────────────────────────────────────
CONTROL_ETFS = {
    'CN': [
        # US-listed wrappers
        'KWEB US Equity', 'FXI US Equity', 'MCHI US Equity',
        'ASHR US Equity', 'CQQQ US Equity', 'GXC US Equity',
        # locally listed
        '510300 CH Equity',   # Huatai-PB CSI 300 ETF (Shanghai)
        '510050 CH Equity',   # ChinaAMC SSE 50 ETF (Shanghai)
        # [Fix 61] mid/small-cap onshore wrappers. Single A-share borrow is
        # scarce and expensive, so these are the practical short-side
        # instruments in China; they also proxy the IC / IM / STAR futures
        # complex, and they are the only way the engine can express a
        # small-cap tilt against an onshore target.
        '510500 CH Equity',   # ChinaAMC CSI 500 — mid cap (IC proxy)
        '512100 CH Equity',   # CSI 1000 — small cap (IM proxy)
        '588000 CH Equity',   # STAR 50
        '159915 CH Equity',   # ChiNext (Shenzhen)
        '2823 HK Equity',     # iShares FTSE China A50 (HK)
        '3188 HK Equity'],    # ChinaAMC CSI 300 (HK)
    'US': [
        # broad / size
        'SPY US Equity', 'QQQ US Equity', 'IWM US Equity', 'DIA US Equity',
        # style factors
        'MTUM US Equity', 'VLUE US Equity', 'QUAL US Equity',
        'USMV US Equity', 'VUG US Equity',
        # core sectors
        'XLK US Equity', 'XLF US Equity', 'XLV US Equity', 'XLE US Equity',
        'XLY US Equity', 'XLC US Equity', 'XLI US Equity'],
    'JP': ['EWJ US Equity', 'DXJ US Equity', 'BBJP US Equity',
           '1306 JT Equity',   # NEXT FUNDS TOPIX ETF (Tokyo)
           '1321 JT Equity'],  # NEXT FUNDS Nikkei 225 ETF (Tokyo)
    'HK': ['EWH US Equity', 'FLHK US Equity',
           '2800 HK Equity',   # Tracker Fund of Hong Kong (HSI)
           '2828 HK Equity',   # Hang Seng China Enterprises ETF
           '3033 HK Equity'],  # CSOP Hang Seng TECH ETF
    'KR': ['EWY US Equity', 'FLKR US Equity',
           '069500 KS Equity'],  # Samsung KODEX 200 (Seoul)
    'TW': ['0050 TT Equity',     # Yuanta Taiwan Top 50 (Taipei)
           '006208 TT Equity'],  # Fubon Taiwan 50 (Taipei)
    'IN': ['INDA US Equity', 'EPI US Equity', 'SMIN US Equity',
           'NIFTYBEES IN Equity'],  # Nippon India Nifty 50 BeES (NSE)
    'DE': ['EWG US Equity',
           'EXS1 GR Equity'],   # iShares Core DAX UCITS (Xetra)
    'GB': ['EWU US Equity',
           'ISF LN Equity',     # iShares Core FTSE 100 (London)
           'VUKE LN Equity'],   # Vanguard FTSE 100 (London)
    'FR': ['EWQ US Equity',
           'CAC FP Equity'],    # Amundi CAC 40 UCITS (Paris)
    'AU': ['EWA US Equity',
           'STW AU Equity',     # SPDR S&P/ASX 200 (ASX)
           'IOZ AU Equity'],    # iShares Core S&P/ASX 200 (ASX)
    # [Fix 43] cross-market style-factor wrappers, added to EVERY universe
    # so momentum / size / value / quality / low-vol have candidates even
    # for non-US targets (US-listed -> treated as async, 2-day returns).
    'FACTOR': ['MTUM US Equity', 'VLUE US Equity', 'QUAL US Equity',
               'USMV US Equity', 'SIZE US Equity', 'IWM US Equity',
               'SPY US Equity', 'IMTM US Equity', 'IVLU US Equity',
               'IQLT US Equity', 'ACWV US Equity'],
    'GLOBAL': ['ACWI US Equity', 'EEM US Equity', 'EFA US Equity', 'URTH US Equity'],
}

# [Fix 14] ETF whitelist (short names) auto-generated — always in sync
ETF_SHORT_SET = {t.split(' ')[0] for lst in CONTROL_ETFS.values() for t in lst}

# ─────────────────────────────────────────────────────────────────────────────
# Exchange code mappings
# ─────────────────────────────────────────────────────────────────────────────
EXCH_MAP = {
    'UQ': 'US', 'UW': 'US', 'UN': 'US', 'US': 'US', 'UA': 'US', 'UP': 'US',
    'HK': 'HK', 'LN': 'LN', 'JT': 'JP', 'GR': 'GR', 'GY': 'GR',
    'FP': 'FP', 'IM': 'IM', 'SM': 'SM', 'AU': 'AU', 'AV': 'AV',
    'IN': 'IN', 'IB': 'IN', 'SJ': 'SJ', 'KS': 'KS', 'KP': 'KS',
    'TT': 'TT', 'SP': 'SP', 'MK': 'MK', 'CH': 'CH',
    # [Fix 32] previously-unmapped codes. 'AT' (ASX single-exchange code,
    # e.g. 'ELS AT Equity') fell through to 'OTHER' while the AU-listed ETFs
    # ('STW AU') mapped to APAC — so Australian stocks and their own local
    # ETFs landed in DIFFERENT timezone groups.
    'AT': 'AU', 'AH': 'AU',            # ASX
    'KQ': 'KS',                        # KOSDAQ
    'C1': 'CH', 'C2': 'CH',            # China Connect lines
    'NA': 'NA', 'SW': 'SW', 'SE': 'SW',  # Amsterdam, SIX
}

# [Fix 19] Normalized exchange code → market key used in INDEX_MAP/CONTROL_ETFS
EXCH_TO_MARKET = {
    'US': 'US', 'HK': 'HK', 'JP': 'JP', 'CH': 'CN', 'KS': 'KR', 'TT': 'TW',
    'IN': 'IN', 'AU': 'AU', 'GR': 'DE', 'LN': 'GB', 'FP': 'FR',
    'IM': 'EU', 'SM': 'EU', 'AV': 'EU', 'SJ': 'EU',
    'NA': 'EU', 'SW': 'EU',            # [Fix 32]
}

# [Fix 32] Approximate cash-session CLOSING HOUR in UTC per normalized
# exchange code (standard time; the ~1h DST drift is absorbed by the
# tolerance). This replaces the old AMER/EMEA/APAC buckets, which were too
# coarse in both directions: Taipei (05:30 UTC) and Mumbai (10:00 UTC) were
# called synchronous because both were 'APAC', while any unmapped exchange
# fell into an 'OTHER' bucket that the async-bias check silently ignored.
EXCH_CLOSE_UTC = {
    'US': 21.0,                        # NYSE / Nasdaq 16:00 ET
    'LN': 16.5,                        # LSE 16:30
    'GR': 16.5, 'FP': 16.5, 'IM': 16.5, 'SM': 16.5, 'AV': 16.5,
    'NA': 16.5, 'SW': 16.5,            # continental Europe ~17:30 CET
    'SJ': 15.0,                        # JSE 17:00 SAST
    'TT': 5.5,                         # TWSE 13:30
    'JP': 6.0,                         # TSE 15:00 JST
    'AU': 6.0,                         # ASX 16:00 AEST
    'KS': 6.5,                         # KRX 15:30 KST
    'CH': 7.0,                         # SSE / SZSE 15:00 CST
    'HK': 8.0,                         # HKEX 16:00
    'SP': 9.0, 'MK': 9.0,              # SGX / Bursa 17:00 local
    'IN': 10.0,                        # NSE / BSE 15:30 IST
}


def _exch_of(full_ticker):
    """'700 HK Equity' → 'HK' (normalized via EXCH_MAP; 'JT' → 'JP' etc.)."""
    parts = full_ticker.split(' ')
    if len(parts) >= 2:
        return EXCH_MAP.get(parts[-2], parts[-2])
    return 'US'


def _close_utc_of(full_ticker):
    """[Fix 32] Approximate closing hour (UTC) of the ticker's exchange,
    or None when the exchange is unknown."""
    return EXCH_CLOSE_UTC.get(_exch_of(full_ticker))


def _close_gap_hours(a, b):
    """[Fix 32] Circular distance in hours between two closing times
    (a 23h gap is really a 1h gap on the clock)."""
    d = abs(a - b) % 24.0
    return min(d, 24.0 - d)


def _fmt_mktcap(mm):
    """[Fix 35] Format a USD-millions floor as '$750mm' / '$5.0bn'. The old
    message divided by 1000 and always printed 'bn', so a 50mm floor showed
    up as '$0bn'."""
    return f"${mm / 1000:.1f}bn" if mm >= 1000 else f"${mm:.0f}mm"


def _apply_exclusions(tickers, exclude):
    """[Fix 36] Manual filter-out of securities the user does not want as
    hedge candidates. Entries match the full ticker ('3156 JP Equity',
    case-insensitive, ' Equity' optional) or the bare short name ('3156').
    Returns (kept, removed)."""
    if not exclude:
        return list(tickers), []
    ex = {str(e).upper().replace(' EQUITY', '').strip() for e in exclude}
    kept, removed = [], []
    for tk in tickers:
        cu = tk.upper().replace(' EQUITY', '').strip()
        if cu in ex or cu.split(' ')[0] in ex:
            removed.append(tk)
        else:
            kept.append(tk)
    return kept, removed


def _resolve_force_include(force_include, columns, target_ticker):
    """[Fix 40] Resolve the user's must-have securities against the
    downloaded price columns. Entries match the full ticker
    ('0050 TT Equity', case-insensitive, ' Equity' optional) or the bare
    short name ('0050'). The target itself is ignored.
    Returns (resolved_full_tickers, unresolved_entries)."""
    if not force_include:
        return [], []
    resolved, unresolved = [], []
    for f in force_include:
        fu = str(f).upper().replace(' EQUITY', '').strip()
        hit = None
        for c in columns:
            cu = c.upper().replace(' EQUITY', '').strip()
            if cu == fu or cu.split(' ')[0] == fu:
                hit = c
                break
        if hit is None:
            unresolved.append(f)
        elif hit != target_ticker:
            resolved.append(hit)
    out, seen = [], set()
    for r in resolved:                 # dedupe, keep user order
        if r not in seen:
            out.append(r)
            seen.add(r)
    return out, unresolved


# [Fix 51] Jupyter detection — used by _print_table to pick a renderer.
try:
    from IPython.display import display as _ipy_display, HTML as _HTML
    from IPython import get_ipython as _get_ipython
    _IN_JUPYTER = _get_ipython() is not None
except Exception:                      # [Fix 87] _HTML kept defined either way
    _IN_JUPYTER = False
    _HTML = None


# ─────────────────────────────────────────────────────────────────────────────
# [Fix 115] REPORT TYPOGRAPHY — ONE renderer, one type scale, one grammar.
# ─────────────────────────────────────────────────────────────────────────────
# [Fix 104] set out to unify the report and only got half way: it introduced
# _kv_block (hand-built HTML) but left _print_table on the pandas Styler. Two
# renderers meant two looks on one page — Styler sized tables to their content
# while kv tables ran to 1100px, Styler set EVERY cell in monospace including
# the prose ones, and the two disagreed about zebra striping. Against the raw
# stdout <pre> that the engine's progress lines still landed in, the page
# carried three fonts and three widths. That is the reported "the font and
# tables are ugly, the structure is messy".
#
# What this pass changes:
#   · ONE table renderer. _print_table and _kv_block both build their markup
#     through _html_table, so every table on the page shares a header band,
#     rules, padding, row rhythm and measure. The pandas Styler path is gone
#     (and with it the jinja2 dependency that [Fix 75] had to guard).
#   · COLUMNS ARE TYPED, not uniformly monospaced. A column of numbers gets
#     tabular figures, right alignment and no wrapping; readings and prose get
#     the UI font. Monospace on a sentence was the ugliest thing on the page.
#   · ONE FACT PER LINE. Readings are written as 'a · b · c'; they are now
#     broken at the '·' so each fact sits on its own line instead of reflowing
#     into a grey block.
#   · PROSE HAS A HIERARCHY. The first sentence of a WHAT IT MEANS cell is the
#     claim and is set in ink; the rest is set muted and a step smaller, so the
#     page can be read at two depths instead of one.
#   · NO INTERNAL TAGS. '[Fix 108]' and friends are changelog markers for
#     whoever maintains this file. They meant nothing to a reader and made the
#     output look like debug spew; they are scrubbed at render time, so no call
#     site can leak one again.
#   · A PLAIN-TEXT TRANSCRIPT. In a notebook the report is displayed as HTML
#     and never touches stdout, so the Word export used to receive only the few
#     progress lines that were still bare print()s. Every renderer now also
#     records a text copy, and Export gets the whole report.
_RF_UI = ("-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,"
          "'Helvetica Neue',Arial,sans-serif")
_RF_NUM = ("ui-monospace,SFMono-Regular,'SF Mono',Menlo,Consolas,"
           "'Liberation Mono',monospace")
_RF_INK = '#1c2733'          # body text / the claim
_RF_MUTE = '#5b6b7c'         # supporting prose
_RF_FAINT = '#8a97a6'        # captions, notes, sub-labels
_RF_HEAD = '#1f3f66'         # header band
_RF_RULE = '#e4e9ef'         # hairlines
_RF_ZEBRA = '#fafbfc'
_RF_MAXW = '1120px'          # one measure for every block on the page
_TXT_W = 96                  # terminal fallback: total table width budget

# Changelog markers ('[Fix 104]', '[Fix 89/95]') are for maintainers of this
# file, never for the reader of a hedge report. Scrubbed at render time.
_FIXTAG = re.compile(r'\[Fix\s*\d+(?:\s*[/,]\s*\d+)*\]\s*')


def _clean(s):
    """Text as it should reach the reader: no internal tags, no double spaces."""
    return re.sub(r'[ \t]{2,}', ' ', _FIXTAG.sub('', str(s))).strip()


# ── plain-text transcript ────────────────────────────────────────────────────
# Populated by every renderer below whichever path it takes, so there is always
# a faithful text copy of what was shown — for the Word export, for a log file,
# and for anyone reading a saved run.
_REPORT_TX = []


def _tx(*lines):
    _REPORT_TX.extend(lines)


def transcript_reset():
    del _REPORT_TX[:]


def transcript_text():
    """The report as plain text, exactly as it was rendered."""
    return '\n'.join(_REPORT_TX)


# ── column typing ────────────────────────────────────────────────────────────
# Typing the COLUMN rather than the cell is deliberate: a column reads as a
# column only when every cell in it is set the same way, and one stray 'n/a'
# should not knock a numeric column back into prose.
_NUMISH = re.compile(r'^[+\-±$<>~≈#]*\s*[\d.,]+')


def _col_kind(cells):
    """'num' | 'prose' | 'text' — decides font, alignment and wrapping."""
    vals = [c.strip() for c in cells if c.strip() and c.strip().lower() != 'n/a']
    if not vals:
        return 'text'
    if sum(bool(_NUMISH.match(v)) for v in vals) >= 0.7 * len(vals):
        return 'num'
    if max(len(v) for v in vals) > 46:
        return 'prose'
    return 'text'


def _facts(s):
    """Split a reading written as 'a · b · c' into its facts.

    The convention throughout this file is to join related readings with ' · '.
    Rendered as one string they reflow into an unreadable run; one per line
    they read like a ticket."""
    parts = [p.strip() for p in re.split(r'\s+·\s+', str(s)) if p.strip()]
    return parts or ['']


def _lede(s):
    """(claim, rest) — the first sentence of an explanation, and the rest.

    Every WHAT IT MEANS cell is written claim-first, so setting the first
    sentence in ink and the remainder muted gives the page a skim layer
    without deleting the detail underneath it."""
    s = _clean(s)
    if not s:
        return '', ''
    m = re.search(r'(?<=[.!?])\s+(?=[A-Z(“"])', s)
    return (s[:m.start()], s[m.end():]) if m else (s, '')


# ── the single HTML table renderer ───────────────────────────────────────────
def _html_table(headers, rows, kinds, title='', note=''):
    n = len(headers)
    # a table carrying an explanation column stretches to the page measure; a
    # data table (a ticket, a weights list, a two-column readout) sizes to its
    # content instead of sitting in a field of white space
    wide = 'prose' in kinds
    th = (f'background:{_RF_HEAD};color:#fff;text-align:left;'
          f'padding:7px 12px;font-size:10.5px;font-weight:600;'
          f'letter-spacing:.06em;text-transform:uppercase;'
          f'font-family:{_RF_UI};white-space:nowrap')
    out = [f'<table style="border-collapse:collapse;border-spacing:0;'
           f'margin:4px 0 14px 0;font-family:{_RF_UI};'
           f'{"width:100%;" if wide else ""}max-width:{_RF_MAXW}">']
    if title:
        out.append(f'<caption style="caption-side:top;text-align:left;'
                   f'font-family:{_RF_UI};font-size:11px;font-weight:700;'
                   f'letter-spacing:.06em;text-transform:uppercase;'
                   f'color:{_RF_HEAD};padding:2px 0 6px 1px;'
                   f'white-space:nowrap">'
                   f'{_html.escape(_clean(title))}</caption>')
    out.append('<tr>' + ''.join(
        f'<th style="{th}{";text-align:right" if kinds[i] == "num" else ""}">'
        f'{_html.escape(_clean(h))}</th>' for i, h in enumerate(headers))
        + '</tr>')
    for ri, row in enumerate(rows):
        sub = str(row[0]).startswith('  ')     # '  label' → a sub-row
        bg = f'background:{_RF_ZEBRA};' if (ri % 2) else ''
        out.append('<tr>')
        for ci in range(n):
            cell, kind = str(row[ci]), kinds[ci]
            td = (f'padding:6px 12px;vertical-align:top;'
                  f'border-bottom:1px solid {_RF_RULE};{bg}')
            if kind == 'num':
                out.append(f'<td style="{td}text-align:right;'
                           f'font-family:{_RF_NUM};font-size:11.5px;'
                           f'font-variant-numeric:tabular-nums;'
                           f'white-space:nowrap;color:{_RF_INK}">'
                           f'{_html.escape(_clean(cell))}</td>')
            elif kind == 'label':
                # [Fix 116] a label may be written as two lines: what the row
                # tells you, then the technical name for it. The reader who
                # does not know the method reads the first line; the reader
                # who does gets the term they were looking for.
                _l1, _, _l2 = str(cell).partition('\n')
                body = _html.escape(_clean(_l1))
                if _l2.strip():
                    body += (f'<div style="font-weight:400;font-size:10.5px;'
                             f'color:{_RF_FAINT};margin-top:2px;'
                             f'letter-spacing:.01em">'
                             f'{_html.escape(_clean(_l2))}</div>')
                out.append(
                    f'<td style="{td}white-space:nowrap;font-size:'
                    f'{11 if sub else 12}px;font-weight:{400 if sub else 600};'
                    f'color:{_RF_FAINT if sub else _RF_INK};'
                    f'padding-left:{26 if sub else 12}px">{body}</td>')
            elif kind == 'prose':
                claim, rest = _lede(cell)
                body = (f'<span style="color:{_RF_INK}">'
                        f'{_html.escape(claim)}</span>')
                if rest:
                    body += (f' <span style="color:{_RF_MUTE};'
                             f'font-size:11.5px">{_html.escape(rest)}</span>')
                out.append(f'<td style="{td}font-size:12px;line-height:1.5;'
                           f'width:46%">{body}</td>')
            else:                              # 'read' / 'text'
                # only a READING is broken into one fact per line; doing it to
                # every text column turned 'lb 252 · pc auto · k 3' in the
                # ALTERNATIVES summary into three lines and made a scan table
                # three times as tall
                parts = (_facts(_clean(cell)) if kind == 'read'
                         else [_clean(cell)])
                body = '<br>'.join(_html.escape(f) for f in parts)
                # a short label ('◄ chosen', 'normal') is a tag, not prose —
                # letting it wrap mid-phrase reads as a layout accident
                _nw = ('white-space:nowrap;'
                       if max((len(x) for x in parts), default=0) <= 22
                       else '')
                out.append(f'<td style="{td}{_nw}font-size:12px;'
                           f'line-height:1.5;font-variant-numeric:tabular-nums;'
                           f'color:{_RF_INK}">{body}</td>')
        out.append('</tr>')
    out.append('</table>')
    if note:
        out.append(f'<div style="font-family:{_RF_UI};font-size:11px;'
                   f'color:{_RF_FAINT};margin:-10px 0 14px 2px;'
                   f'max-width:{_RF_MAXW}">{_html.escape(_clean(note))}</div>')
    return ''.join(out)


def _text_table(headers, rows, kinds, title='', note='', indent='  '):
    """Terminal fallback — the boxed layout, with every cell wrapped so a table
    cannot run off the side of the terminal."""
    n = len(headers)
    cells = [[_facts(_clean(r[i])) if kinds[i] == 'read'
              else ([_clean(x) for x in str(r[i]).split('\n')]  # [Fix 116]
                    if kinds[i] == 'label' else [_clean(r[i])])
              for i in range(n)] for r in rows]
    w = [max([len(_clean(headers[i]))]
             + [max((len(x) for x in row[i]), default=0) for row in cells])
         for i in range(n)]
    budget = _TXT_W - len(indent) - 3 * n - 1
    # Prose gives up width first, but only down to a readable measure; after
    # that every wrappable column shrinks widest-first. Shrinking prose alone
    # produced tables with a 72-char reading column beside a 14-char ribbon of
    # explanation — technically inside the budget, unreadable in practice.
    while sum(w) > budget:
        cand = [i for i in range(n) if kinds[i] != 'num'
                and w[i] > max(len(_clean(headers[i])), 12)]
        if not cand:
            break
        wide_prose = [i for i in cand if kinds[i] == 'prose' and w[i] > 30]
        w[max(wide_prose or cand, key=lambda j: w[j])] -= 1

    def wrap(cell, width):
        out = []
        for piece in cell:
            out += _textwrap.wrap(piece, width) or ['']
        return out

    def rule(l, m, r):
        return indent + l + m.join('─' * (x + 2) for x in w) + r

    def band(vals):
        cols = [wrap(v, w[i]) for i, v in enumerate(vals)]
        out = []
        for k in range(max(len(c) for c in cols)):
            parts = []
            for i, c in enumerate(cols):
                s = c[k] if k < len(c) else ''
                pad = ' ' * (w[i] - len(s))
                parts.append(f' {pad}{s} ' if kinds[i] == 'num'
                             else f' {s}{pad} ')
            out.append(indent + '│' + '│'.join(parts) + '│')
        return out

    lines = ['', f"{indent}{_clean(title).upper()}"] if title else []
    lines.append(rule('┌', '┬', '┐'))
    lines += band([[_clean(h)] for h in headers])
    lines.append(rule('├', '┼', '┤'))
    for row in cells:
        lines += band(row)
    lines.append(rule('└', '┴', '┘'))
    if note:
        lines += [indent + x for x in
                  _textwrap.wrap(_clean(note), _TXT_W - len(indent))]
    return lines


def _emit_table(headers, rows, kinds, title='', note='', indent='  '):
    """The one entry point for both renderers, so a table can look only one
    way. Both paths are fed the same headers/rows/kinds."""
    _kv_flush()                    # ordering — see the note on _KVCTX below
    lines = _text_table(headers, rows, kinds, title=title, note=note,
                        indent=indent)
    _tx(*lines)
    try:
        if _IN_JUPYTER:
            _ipy_display(_HTML(_html_table(headers, rows, kinds, title=title,
                                           note=note)))
            return
    except Exception:
        pass                       # never lose a block to a display problem
    for ln in lines:
        print(ln)


# [Fix 104] ORDERING. kv() buffers its rows so a run of them renders as ONE
# table. Anything else that writes to the report must therefore flush that
# buffer first, or the table jumps ahead of the prose above it. Rather than
# police ~44 call sites by hand, the buffer lives at module level and the
# writers (_section, _emit_table, _P) flush it on entry — so the ordering is
# correct by construction and stays correct when lines are added.
_KV_HEADERS = ('', 'READING', 'WHAT IT MEANS')
_KVCTX = {'title': None, 'note': '', 'rows': [], 'headers': _KV_HEADERS}


def _kv_flush():
    if _KVCTX['rows']:
        _rows, _title = list(_KVCTX['rows']), _KVCTX['title']
        _note, _hdr = _KVCTX['note'], _KVCTX['headers']
        # [Fix 115] clear the context BEFORE rendering — that makes the
        # _kv_flush() at the top of _emit_table a no-op instead of a recursion,
        # and it clears `headers`, which the old version left behind. Leaving
        # them behind is how [1] VALIDATION came to be headed 'WATCH / THIS RUN
        # / WHY AT THIS HORIZON' — inherited from the block above it — with an
        # empty third column under the third heading.
        _kv_reset()
        _kv_block(_title, _rows, note=_note, headers=_hdr)


def _kv_open(title=None, headers=_KV_HEADERS, note=''):
    _kv_flush()
    _KVCTX.update(title=title, headers=headers, note=note, rows=[])


def _kv_reset():
    """[Fix 114] Drop any half-built table WITHOUT rendering it.

    _KVCTX is module-level, so a run that raises part-way through the report
    leaves its buffered rows sitting there — and the UI catches exceptions and
    lets you re-run, so the next run's first table would silently open with
    rows from the failed one. Discarding is right, not flushing: those rows
    describe a run that did not finish."""
    _KVCTX.update(title=None, note='', rows=[], headers=_KV_HEADERS)


def _P(*a, **k):
    """print() for the report — flushes any pending kv table first.
    In Jupyter a bare line renders as styled prose rather than raw <pre>."""
    _kv_flush()
    if not a:
        _tx('')
        if not _IN_JUPYTER:
            print()
        return
    _raw = ' '.join(str(x) for x in a).strip('\n')
    _lead = len(_raw) - len(_raw.lstrip())
    _txt = _clean(_raw)
    if not _txt:
        return
    # '►' takeaways and '⚠' warnings are the lines a reader is meant to stop
    # on; everything else at this level is connective tissue.
    _key = _txt[:1] in '►⚠✓✗△'
    # the text copy is wrapped: it is what the Word appendix and any saved log
    # get, and a 160-column paragraph is unreadable in both
    _pad = ' ' * (_lead + 2)
    _wrapped = _textwrap.wrap(_txt, _TXT_W - len(_pad) - 2,
                              subsequent_indent='  ') or ['']
    _tx(*[_pad + ln for ln in _wrapped])
    try:
        if _IN_JUPYTER and k.pop('html', True):
            _ipy_display(_HTML(
                f'<div style="font-family:{_RF_UI};font-size:12px;'
                f'color:{_RF_INK if _key else _RF_MUTE};'
                f'font-weight:{600 if _key else 400};line-height:1.5;'
                f'margin:3px 0 3px {8 + _lead * 3}px;max-width:{_RF_MAXW}">'
                f'{_html.escape(_txt)}</div>'))
            return
    except Exception:
        pass
    for _ln in _wrapped:
        print(_pad + _ln, **k)


def _section(title, subtitle=''):
    """A report section header — styled in Jupyter, ruled in text."""
    _kv_flush()
    title, subtitle = _clean(title), _clean(subtitle)
    _tx('', '═' * _TXT_W, f" {title}")
    if subtitle:
        _tx(f" {subtitle}")
    _tx('═' * _TXT_W)
    try:
        if _IN_JUPYTER:
            _ipy_display(_HTML(
                f'<div style="margin:26px 0 8px 0;padding:0 0 5px 0;'
                f'border-bottom:2px solid {_RF_HEAD};font-family:{_RF_UI};'
                f'max-width:{_RF_MAXW}">'
                f'<span style="font-size:14.5px;font-weight:700;'
                f'letter-spacing:.03em;color:{_RF_HEAD}">'
                f'{_html.escape(title)}</span>'
                + (f'<span style="font-size:12px;color:{_RF_FAINT};'
                   f'margin-left:10px">{_html.escape(subtitle)}</span>'
                   if subtitle else '')
                + '</div>'))
            return
    except Exception:
        pass
    print(f"\n{'═' * _TXT_W}\n {title}" + (f"\n {subtitle}" if subtitle else '')
          + f"\n{'═' * _TXT_W}")


def _kv_block(title, rows, note='', headers=_KV_HEADERS):
    """The workhorse: rows of (label, reading, implication).

    Every ticket line used to be `kv('SIZING', <one long sentence>)` — a
    26-char label and then an unbroken run of numbers and prose that the eye
    cannot parse and that wraps unpredictably. Splitting the implication into
    its own column is the single biggest readability win, and it forces each
    number to come with a statement of what to DO about it.

    rows may be 2-tuples (label, reading) or 3-tuples with the implication.
    A row whose label starts with '  ' renders as a sub-row (indented, muted).
    """
    rows = [list(r) + [''] * (3 - len(r)) for r in rows if r is not None]
    if not rows:
        return
    headers, kinds = list(headers), ['label', 'read', 'prose']
    if not any(str(r[2]).strip() for r in rows):     # no implications → 2 cols
        headers, rows, kinds = headers[:2], [r[:2] for r in rows], kinds[:2]
    _emit_table(headers, rows, kinds, title=title, note=note)


def _print_table(headers, rows, indent='  ', wrap_last=None, title='',
                 note=''):
    """Render rows as a table. `wrap_last` is kept for call-site compatibility
    and ignored — wrapping is decided per column now, by _col_kind."""
    headers, rows = list(headers), [[str(x) for x in r] for r in rows]
    kinds = [_col_kind([r[i] for r in rows]) for i in range(len(headers))]
    _emit_table(headers, rows, kinds, title=title, note=note, indent=indent)


def _market_of(target_ticker, country_iso):
    """[Fix 19] Market key for index/ETF lookups.

    Rule: use the LISTING market derived from the exchange code. The only
    exception is a US listing whose country of risk is not the US (an ADR,
    e.g. BABA US with COUNTRY_ISO=CN) — there the country mapping is the
    right 'local' universe (HXC = US-listed China names).

    Examples: 700 HK (Tencent, COUNTRY_ISO=CN) → 'HK' → HSTECH/HSI.
              BABA US (COUNTRY_ISO=CN)        → 'CN' → HXC.
              005930 KS (Samsung)             → 'KR' → KOSPI2.
    """
    exch = _exch_of(target_ticker)
    market = EXCH_TO_MARKET.get(exch)
    if exch == 'US' and country_iso != 'US':
        return country_iso
    return market or country_iso


def _fmt_bbg_ticker(raw):
    """[Fix 1] Normalize 'XXX YY' or 'XXX YY Equity' → 'XXX <exch> Equity'."""
    parts = raw.replace(' Equity', '').split(' ')
    if len(parts) >= 2 and parts[0]:
        return f"{parts[0]} {EXCH_MAP.get(parts[-1], parts[-1])} Equity"
    return None


# ─────────────────────────────────────────────────────────────────────────────
# Bloomberg helper functions
# ─────────────────────────────────────────────────────────────────────────────
def bbg_connect(host='localhost', port=8194):
    opts = blpapi.SessionOptions()
    opts.setServerHost(host)
    opts.setServerPort(port)
    s = blpapi.Session(opts)
    # [Fix 13] explicit raise instead of assert (assert is stripped under -O)
    if not s.start():
        raise ConnectionError(
            f"Bloomberg session failed to start ({host}:{port}) — "
            "check that the terminal is logged in")
    if not s.openService("//blp/refdata"):
        raise ConnectionError("Failed to open //blp/refdata service")
    print(f"✓ Bloomberg connected ({host}:{port})")
    return s


# [Fix 3] Lazy session: reconnectable, not frozen into a default argument
SESSION = None

def get_session():
    global SESSION
    if SESSION is None:
        SESSION = bbg_connect()
    return SESSION


def _drain_events(session, on_message, timeout_ms=60000):
    """[Fix 2] Shared event loop: raises on TIMEOUT instead of spinning forever."""
    while True:
        ev = session.nextEvent(timeout_ms)
        if ev.eventType() == blpapi.Event.TIMEOUT:
            raise TimeoutError(
                f"Bloomberg request timed out after {timeout_ms} ms — "
                "check terminal connectivity or reduce the batch size")
        for msg in ev:
            on_message(msg)
        if ev.eventType() == blpapi.Event.RESPONSE:
            break


def bbg_hist_field(session, tickers, field, start, end):
    """[Fix 44] Generic historical BDH for a single numeric field (e.g.
    'CUR_MKT_CAP'). Same batching / draining as bbg_hist; returns a DataFrame
    keyed by full ticker. Used to build point-in-time market-cap paths so the
    size filter can be applied as-of each backtest window instead of with
    today's cap (look-ahead for names that have changed size)."""
    svc = session.getService("//blp/refdata")
    data, bad = {}, []
    batch = 100
    for i in range(0, len(tickers), batch):
        req = svc.createRequest("HistoricalDataRequest")
        for t in tickers[i:i + batch]:
            req.getElement("securities").appendValue(t)
        req.getElement("fields").appendValue(field)
        req.set("periodicitySelection", "DAILY")
        req.set("startDate", start)
        req.set("endDate", end)
        req.set("nonTradingDayFillOption", "NON_TRADING_WEEKDAYS")
        req.set("nonTradingDayFillMethod", "PREVIOUS_VALUE")
        session.sendRequest(req)

        def on_msg(msg):
            if not msg.hasElement("securityData"):
                return
            sd = msg.getElement("securityData")
            sec = sd.getElementAsString("security")
            if sd.hasElement("securityError"):
                bad.append(sec)
                return
            fda = sd.getElement("fieldData")
            recs = []
            for k in range(fda.numValues()):
                fd = fda.getValueAsElement(k)
                try:
                    recs.append({'date': fd.getElementAsDatetime("date"),
                                 'val': fd.getElementAsFloat(field)})
                except Exception:
                    continue
            if recs:
                df = pd.DataFrame(recs)
                df['date'] = pd.to_datetime(df['date'])
                data[sec] = df.set_index('date')['val']

        _drain_events(session, on_msg)
    return pd.DataFrame(data)


def bbg_hist(session, tickers, start, end):
    """Historical PX_LAST. Returns a DataFrame keyed by full ticker.
    Invalid tickers are reported, not silently dropped."""
    svc = session.getService("//blp/refdata")
    data, bad = {}, []
    batch = 100  # keep historical requests reasonably sized
    for i in range(0, len(tickers), batch):
        req = svc.createRequest("HistoricalDataRequest")
        for t in tickers[i:i + batch]:
            req.getElement("securities").appendValue(t)
        req.getElement("fields").appendValue("PX_LAST")
        req.set("periodicitySelection", "DAILY")
        req.set("startDate", start)
        req.set("endDate", end)
        req.set("nonTradingDayFillOption", "NON_TRADING_WEEKDAYS")
        req.set("nonTradingDayFillMethod", "PREVIOUS_VALUE")
        # [Fix 58] Make the price adjustment EXPLICIT. Without these three
        # lines HistoricalDataRequest inherits whatever DPDF the local
        # terminal happens to be set to — on most desks that is split-
        # adjusted but NOT dividend-adjusted, so every ex-div date shows up
        # as a spurious negative return. ETFs and HK/China high-payout names
        # distribute 2-6% a year, which quietly biases betas and hedge
        # ratios. Setting them here also means two colleagues running this
        # notebook on different terminals finally get identical numbers.
        req.set("adjustmentSplit", True)       # splits / rights issues
        req.set("adjustmentNormal", True)      # regular cash dividends
        req.set("adjustmentAbnormal", True)    # special distributions
        session.sendRequest(req)

        def on_msg(msg):
            if not msg.hasElement("securityData"):
                return
            sd = msg.getElement("securityData")
            sec = sd.getElementAsString("security")
            if sd.hasElement("securityError"):          # [Fix 13/15]
                bad.append(sec)
                return
            fda = sd.getElement("fieldData")
            recs = []
            for k in range(fda.numValues()):
                fd = fda.getValueAsElement(k)
                try:
                    recs.append({'date': fd.getElementAsDatetime("date"),
                                 'px': fd.getElementAsFloat("PX_LAST")})
                except Exception:
                    continue
            if recs:
                df = pd.DataFrame(recs)
                df['date'] = pd.to_datetime(df['date'])
                data[sec] = df.set_index('date')['px']

        _drain_events(session, on_msg)
    if bad:
        print(f"  ⚠ {len(bad)} invalid tickers dropped (e.g. {bad[:5]})")
    return pd.DataFrame(data)


def bbg_bulk(session, ticker, field, overrides=None):
    """BDS bulk field — returns a list of dicts (one per row).
    [Fix 49] optional overrides, e.g. {'END_DATE_OVERRIDE': '20240630'} to
    read INDX_MWEIGHT_HIST as of a past date."""
    svc = session.getService("//blp/refdata")
    req = svc.createRequest("ReferenceDataRequest")
    req.getElement("securities").appendValue(ticker)
    req.getElement("fields").appendValue(field)
    if overrides:
        ovr = req.getElement("overrides")
        for k, v in overrides.items():
            o = ovr.appendElement()
            o.setElement("fieldId", k)
            o.setElement("value", str(v))
    session.sendRequest(req)
    rows = []

    def on_msg(msg):
        if not msg.hasElement("securityData"):
            return
        arr = msg.getElement("securityData")
        for i in range(arr.numValues()):
            sd = arr.getValueAsElement(i)
            if sd.hasElement("securityError"):
                continue
            fd = sd.getElement("fieldData")
            if fd.hasElement(field):
                bulk = fd.getElement(field)
                for j in range(bulk.numValues()):
                    row = bulk.getValueAsElement(j)
                    d = {}
                    for k in range(row.numElements()):
                        e = row.getElement(k)
                        d[str(e.name())] = str(e.getValue())
                    rows.append(d)

    _drain_events(session, on_msg)
    return rows


def bbg_ref_multi(session, tickers, fields, overrides=None):
    """BDP reference data for many tickers → {full_ticker: {field: value}}.
    [Fix 25] `overrides` is an optional {fieldId: value} dict."""
    svc = session.getService("//blp/refdata")
    result = {}
    batch = 300
    for i in range(0, len(tickers), batch):
        req = svc.createRequest("ReferenceDataRequest")
        for t in tickers[i:i + batch]:
            req.getElement("securities").appendValue(t)
        for f in fields:
            req.getElement("fields").appendValue(f)
        if overrides:
            ovr = req.getElement("overrides")
            for k, v in overrides.items():
                o = ovr.appendElement()
                o.setElement("fieldId", k)
                o.setElement("value", str(v))
        session.sendRequest(req)

        def on_msg(msg):
            if not msg.hasElement("securityData"):
                return
            arr = msg.getElement("securityData")
            for j in range(arr.numValues()):
                sd = arr.getValueAsElement(j)
                sec = sd.getElementAsString("security")
                if sd.hasElement("securityError"):
                    continue
                fd = sd.getElement("fieldData")
                result[sec] = {}
                for f in fields:
                    if fd.hasElement(f):
                        try:
                            result[sec][f] = fd.getElementAsString(f)
                        except Exception:
                            try:
                                result[sec][f] = fd.getElementAsFloat(f)
                            except Exception:
                                pass

        _drain_events(session, on_msg)
    return result


def bbg_ref_single(session, ticker, fields):
    return bbg_ref_multi(session, [ticker], fields).get(ticker, {})


# ─────────────────────────────────────────────────────────────────────────────
# Universe discovery
# ─────────────────────────────────────────────────────────────────────────────
def _index_members(session, spec):
    """[Fix 37] `spec` is one index ticker or a tuple of fallbacks. Returns
    (ticker_used, members) from the first index that yields any members —
    an index that loads but returns 0 members (licence-gated INDX_MEMBERS)
    now falls through instead of silently emptying a universe leg."""
    tickers = [spec] if isinstance(spec, str) else list(spec or [])
    for n, idx_ticker in enumerate(tickers):
        members = []
        try:
            for r in bbg_bulk(session, idx_ticker, 'INDX_MEMBERS'):
                tk = _fmt_bbg_ticker(r.get('Member Ticker and Exchange Code', ''))
                if tk:
                    members.append(tk)
        except Exception as e:
            print(f"    ⚠ failed to read {idx_ticker} members: {e}")
            continue
        if members:
            return idx_ticker, members
        more = " — trying fallback" if n + 1 < len(tickers) else ""
        print(f"    ⚠ {idx_ticker} returned 0 members (licence-gated?){more}")
    return None, []


def _pit_index_membership(session, index_tickers, start, end, freq_days=91):
    """[Fix 49] Point-in-time index-membership snapshots via INDX_MWEIGHT_HIST
    with an END_DATE_OVERRIDE, roughly quarterly across [start, end]. Returns
    {Timestamp: frozenset(full tickers)} (union across the given indices).
    GRACEFUL FALLBACK: some indices licence-gate this field (e.g. certain
    TW/Asia benchmarks) — any snapshot that fails is skipped, and if NOTHING
    comes back an empty dict is returned with a warning, so the pipeline
    silently reverts to today's-members behavior (survivorship caveat then
    still applies)."""
    idxs = [i for i in dict.fromkeys(index_tickers or []) if i]
    if not idxs:
        return {}
    d0 = pd.Timestamp(datetime.strptime(start, '%Y%m%d'))
    d1 = pd.Timestamp(datetime.strptime(end, '%Y%m%d'))
    dates = list(pd.date_range(d0, d1, freq=f'{freq_days}D'))
    if not dates or dates[-1] != d1:
        dates.append(d1)
    snaps = {}
    for dt in dates:
        members = set()
        for idx in idxs:
            try:
                rows = bbg_bulk(session, idx, 'INDX_MWEIGHT_HIST',
                                overrides={'END_DATE_OVERRIDE':
                                           dt.strftime('%Y%m%d')})
            except Exception:
                continue                      # this snapshot/index: give up
            for r in rows:
                raw = r.get('Index Member') or next(iter(r.values()), '')
                tk = _fmt_bbg_ticker(raw)
                if tk:
                    members.add(tk)
        if members:
            snaps[dt] = frozenset(members)
    if not snaps:
        print(f"    ⚠ INDX_MWEIGHT_HIST unavailable for {idxs} "
              f"(licence-gated?) — point-in-time membership skipped; "
              f"falling back to today's members (survivorship caveat applies)")
    return snaps


def auto_discover_universe(session, target_ticker, max_first_degree_peers=10,
                           min_mktcap_usd_mm=5000, meta=None):
    print(f"  Auto-discovering universe for {target_ticker}...")

    ref_fields = [
        'COUNTRY_ISO',
        'BICS_LEVEL_2_INDUSTRY_GROUP_NAME',
        'BICS_LEVEL_3_INDUSTRY_NAME',
        'BICS_LEVEL_4_SUB_INDUSTRY_NAME',
        'GICS_SECTOR_NAME',
        'GICS_INDUSTRY_NAME',
        'GICS_SUB_INDUSTRY_NAME',   # [Fix 1] correct field name
        'CUR_MKT_CAP',
        'AVERAGE_VOLUME_30D',
    ]
    tgt_ref = bbg_ref_single(session, target_ticker, ref_fields)
    country = (str(tgt_ref.get('COUNTRY_ISO') or 'US'))[:2].upper()

    # [Fix 19] market = listing venue (Tencent 700 HK → 'HK', not 'CN');
    # country-of-risk only kicks in for US-listed ADRs (BABA US → 'CN')
    market = _market_of(target_ticker, country)

    tgt_bics_lvl4 = tgt_ref.get('BICS_LEVEL_4_SUB_INDUSTRY_NAME')
    tgt_bics_lvl3 = tgt_ref.get('BICS_LEVEL_3_INDUSTRY_NAME')
    tgt_bics_lvl2 = tgt_ref.get('BICS_LEVEL_2_INDUSTRY_GROUP_NAME')
    tgt_gics_sub = tgt_ref.get('GICS_SUB_INDUSTRY_NAME')
    tgt_gics_ind = tgt_ref.get('GICS_INDUSTRY_NAME')
    tgt_gics_sec = tgt_ref.get('GICS_SECTOR_NAME') or 'Broad'

    sector_key = 'Broad'
    for key_word in ['Technology', 'Communication', 'Financials',
                     'Health Care', 'Energy', 'Consumer Discretionary']:
        if (tgt_bics_lvl2 and key_word.lower() in str(tgt_bics_lvl2).lower()) or \
           (key_word.lower() in str(tgt_gics_sec).lower()):
            sector_key = key_word
            break

    # 1. Local index mapping — keyed by listing market [Fix 19/37]
    idx_spec = (INDEX_MAP.get((market, sector_key))
                or INDEX_MAP.get((market, 'Broad'))
                or INDEX_MAP.get((country, 'Broad')))
    local_index_ticker, local_members = _index_members(session, idx_spec)
    print(f"    ├─ market={market}, sector={sector_key} → local index "
          f"{local_index_ticker or 'NONE'}: {len(local_members)} names")

    # 2. Global cross-country super pool [Fix 37: fallback-aware]
    broad_specs = sorted({(v if isinstance(v, str) else tuple(v))
                          for k, v in INDEX_MAP.items() if k[1] == 'Broad'},
                         key=str)
    global_super_pool = []
    for spec in broad_specs:
        _, members = _index_members(session, spec)
        global_super_pool.extend(members)
    global_super_pool = sorted(set(global_super_pool))

    pool_ref_data = bbg_ref_multi(session, global_super_pool, ref_fields)

    # 3. Dual-taxonomy matching (BICS OR GICS)
    industry_matches = []
    print(f"    ├─ running BICS + GICS dual-taxonomy cross-country search...")
    for tk, d in pool_ref_data.items():
        if tk == target_ticker:
            continue
        bics_match = (tgt_bics_lvl4 and d.get('BICS_LEVEL_4_SUB_INDUSTRY_NAME') == tgt_bics_lvl4) or \
                     (tgt_bics_lvl3 and d.get('BICS_LEVEL_3_INDUSTRY_NAME') == tgt_bics_lvl3)
        gics_match = (tgt_gics_sub and d.get('GICS_SUB_INDUSTRY_NAME') == tgt_gics_sub) or \
                     (tgt_gics_ind and d.get('GICS_INDUSTRY_NAME') == tgt_gics_ind)
        if bics_match or gics_match:
            industry_matches.append(tk)
    print(f"    ├─ industry matches captured: {len(industry_matches)} names")

    # 4. Bloomberg official peers + second-degree peers
    #    [Fix 1: proper formatting + capped API calls]
    peer_matches = []
    try:
        first_degree = []
        for r in bbg_bulk(session, target_ticker, 'BLOOMBERG_PEERS'):
            tk = _fmt_bbg_ticker(r.get('Peer Ticker', ''))
            if tk:
                first_degree.append(tk)
        peer_matches.extend(first_degree)
        for peer in first_degree[:max_first_degree_peers]:
            try:
                for r in bbg_bulk(session, peer, 'BLOOMBERG_PEERS'):
                    tk = _fmt_bbg_ticker(r.get('Peer Ticker', ''))
                    if tk:
                        peer_matches.append(tk)
            except Exception:
                continue
        print(f"    ├─ Bloomberg peers (incl. 2nd degree): {len(set(peer_matches))} names")
    except Exception as e:
        print(f"    ⚠ failed to read peers: {e}")

    # 5. Control ETFs — from BOTH listing market and country of risk [Fix 19]
    all_etfs = sorted(set(CONTROL_ETFS.get(market, []) +
                          CONTROL_ETFS.get(country, []) +
                          CONTROL_ETFS.get('FACTOR', []) +   # [Fix 43]
                          CONTROL_ETFS['GLOBAL']))

    if meta is not None:                          # [Fix 49] expose discovery
        meta['index_ticker'] = local_index_ticker  # facts for PIT membership
        meta['non_index_shorts'] = (
            {t.split(' ')[0] for t in peer_matches} |
            {t.split(' ')[0] for t in industry_matches} |
            {t.split(' ')[0] for t in all_etfs})

    # Merge + [Fix 7] dedupe by short name but KEEP full tickers
    # (local members take priority on collision)
    merged, seen_short = [], set()
    for tk in local_members + peer_matches + industry_matches + all_etfs:
        if tk == target_ticker:
            continue
        short = tk.split(' ')[0]
        if short == target_ticker.split(' ')[0]:
            continue
        if short not in seen_short:
            merged.append(tk)
            seen_short.add(short)

    # [Fix 25] Liquidity filter: drop non-ETF names below the USD market-cap
    # floor. Without this, microcaps could enter a mega-cap's basket on
    # spurious loading similarity. CRNCY_ADJ_MKT_CAP with a USD override is
    # reported in USD millions; parse failures are kept (fail-open) so a
    # field-permission issue cannot wipe out the universe.
    if min_mktcap_usd_mm and merged:
        try:
            liq = bbg_ref_multi(session, merged, ['CRNCY_ADJ_MKT_CAP'],
                                overrides={'EQY_FUND_CRNCY': 'USD'})
            kept, dropped, unknown = [], 0, 0
            for tk in merged:
                if tk.split(' ')[0] in ETF_SHORT_SET:
                    kept.append(tk)
                    continue
                try:
                    mc = float(liq.get(tk, {}).get('CRNCY_ADJ_MKT_CAP'))
                except (TypeError, ValueError):
                    unknown += 1
                    kept.append(tk)      # unknown → keep (fail-open)
                    continue
                if mc >= min_mktcap_usd_mm:
                    kept.append(tk)
                else:
                    dropped += 1
            # [Fix 35] correct units in the message ('$50mm', not '$0bn') and
            # report how many names passed only because their cap was unknown
            if dropped or unknown:
                print(f"    ├─ liquidity filter: dropped {dropped} names below "
                      f"{_fmt_mktcap(min_mktcap_usd_mm)} market cap"
                      + (f" ({unknown} unknown caps kept)" if unknown else ""))
            merged = kept
        except Exception as e:
            print(f"    ⚠ liquidity filter skipped ({e})")

    print(f"    ✓ final merged universe: {len(merged)} candidates")
    return merged


# ─────────────────────────────────────────────────────────────────────────────
# [Fix 88] DOWNLOAD CACHE — Bloomberg pulls are slow, rate-limited and lost on
# every kernel restart. download_hedge_data() now writes its result to disk and
# it can be reloaded with zero API calls, days later, from another session.
# Files are plain pickles of the returned dict (DataFrames + dicts) under
# HEDGE_CACHE_DIR. Note: a pickle is tied to your pandas version — if you
# upgrade pandas and a load fails, just re-download.
# ─────────────────────────────────────────────────────────────────────────────
import os as _os
import pickle as _pickle
import glob as _glob
import re as _re
import json as _json

HEDGE_CACHE_DIR = _os.path.join(_os.getcwd(), 'hedge_data_cache')

# [Fix 103] Where exported hedge reports go. Set once here, or override per
# call / in the form's Export folder box. If the drive is not mapped on the
# machine running the notebook, the exporter falls back to a local folder and
# says so rather than failing.
HEDGE_EXPORT_DIR = r'G:\FIN_COMM\DeltaOne\Kenny\PCA'


def _slug(text):
    return _re.sub(r'[^A-Za-z0-9]+', '-', str(text)).strip('-')[:60] or 'data'


def save_hedge_data(data, label=None, cache_dir=None):
    """[Fix 88] Persist a download_hedge_data() result. Returns the path.

    [Fix 90] also writes a small JSON sidecar holding the fingerprint (legs,
    download filters, coverage dates, ticker list) so the catalogue and the
    GUI can judge whether a cached set still fits the book you are hedging
    WITHOUT unpickling every file."""
    d = cache_dir or HEDGE_CACHE_DIR
    _os.makedirs(d, exist_ok=True)
    lab = _slug(label or data.get('target_label') or 'data')
    stamp = datetime.now().strftime('%Y%m%d-%H%M')
    path = _os.path.join(d, f"{lab}__{stamp}.pkl")
    with open(path, 'wb') as fh:
        _pickle.dump(data, fh, protocol=4)
    try:                                                   # [Fix 90]
        px = data.get('prices')
        fp = dict(data.get('fingerprint') or {})
        fp.update({
            'label': data.get('target_label'),
            'legs': sorted((data.get('portfolio') or {}).get('legs') or []),
            'first_date': str(px.index[0].date()) if px is not None else None,
            'last_date': str(px.index[-1].date()) if px is not None else None,
            'n_rows': int(px.shape[0]) if px is not None else 0,
            'n_tickers': int(px.shape[1]) if px is not None else 0,
            'columns': sorted(map(str, px.columns)) if px is not None else []})
        with open(path[:-4] + '.json', 'w') as fh:
            _json.dump(fp, fh)
    except Exception:
        pass                       # sidecar is an optimization, never fatal
    print(f"  ✓ data cached → {path}")
    return path


def list_hedge_data(cache_dir=None):
    """[Fix 88/90] Cached downloads, newest first, with the coverage each one
    carries (read from the JSON sidecar — no unpickling)."""
    d = cache_dir or HEDGE_CACHE_DIR
    rows = []
    for p in sorted(_glob.glob(_os.path.join(d, '*.pkl')),
                    key=_os.path.getmtime, reverse=True):
        base = _os.path.basename(p)[:-4]
        lab, _, stamp = base.partition('__')
        rec = {'label': lab, 'saved': stamp,
               'MB': round(_os.path.getsize(p) / 1e6, 1), 'path': p,
               'last_date': None, 'n_tickers': None}
        try:                                               # [Fix 90]
            with open(p[:-4] + '.json') as fh:
                fp = _json.load(fh)
            rec['last_date'] = fp.get('last_date')
            rec['n_tickers'] = fp.get('n_tickers')
            rec['_fp'] = fp
        except Exception:
            pass
        rows.append(rec)
    return pd.DataFrame(rows)


def _busdays_since(date_str):
    """[Fix 90] Business days between a YYYY-MM-DD string and today."""
    if not date_str:
        return None
    try:
        return int(np.busday_count(pd.Timestamp(date_str).date(),
                                   datetime.now().date()))
    except Exception:
        return None


def check_cache_fit(target, data=None, fingerprint=None, stale_after=5):
    """[Fix 90] Is a cached download still the right one for THIS book?

    A download is reusable when every leg you want to hedge has prices in it
    — the engine excludes legs from the candidate pool and treats ex-legs as
    candidates ([Fix 66]), so leg churn inside the same book is fine. What is
    NOT fine is a leg that was never downloaded, or prices that stop before
    the period you care about. Returns a dict with 'status' in
    {'exact', 'compatible', 'stale', 'mismatch'} and a one-line 'message'.
    """
    legs = _coerce_targets(target)[0]
    if data is not None:
        cols = set(map(str, data['prices'].columns))
        fp_legs = sorted((data.get('portfolio') or {}).get('legs') or [])
        last = str(data['prices'].index[-1].date())
    else:
        fp = fingerprint or {}
        cols = set(fp.get('columns') or [])
        fp_legs = sorted(fp.get('legs') or [])
        last = fp.get('last_date')
    missing = [l for l in legs if l not in cols]
    if missing:
        return {'status': 'mismatch', 'missing': missing, 'last_date': last,
                'message': (f"needs a fresh download — "
                            f"{', '.join(m.split(' ')[0] for m in missing)} "
                            f"{'is' if len(missing) == 1 else 'are'} not in "
                            f"this dataset")}
    age = _busdays_since(last)
    same = sorted(legs) == fp_legs
    if age is not None and age > stale_after:
        return {'status': 'stale', 'missing': [], 'last_date': last,
                'age': age,
                'message': (f"prices stop at {last} — {age} business days "
                            f"old; re-download to include recent sessions")}
    base = ("matches this book exactly" if same
            else "covers every leg of this book (the book changed, the "
                 "universe still fits)")
    return {'status': ('exact' if same else 'compatible'), 'missing': [],
            'last_date': last, 'age': age,
            'message': f"{base} · prices through {last}"
                       + (f" ({age} business day{'s' if age != 1 else ''} "
                          f"old)" if age else " (current)")}


def load_hedge_data(which=0, cache_dir=None):
    """[Fix 88] Reload a cached download — zero API calls.

    `which` may be a path, a label fragment ('AAPL'), or an integer index
    into list_hedge_data() (0 = most recent)."""
    if isinstance(which, str) and which.endswith('.pkl') and _os.path.exists(which):
        path = which
    else:
        cat = list_hedge_data(cache_dir)
        if not len(cat):
            raise FileNotFoundError(
                f"[Fix 88] no cached downloads in "
                f"{cache_dir or HEDGE_CACHE_DIR}")
        if isinstance(which, int):
            if which >= len(cat):
                raise IndexError(f"[Fix 88] only {len(cat)} cached downloads")
            path = cat.iloc[which]['path']
        else:
            hit = cat[cat['label'].str.contains(str(which), case=False,
                                                regex=False)]
            if not len(hit):
                raise FileNotFoundError(
                    f"[Fix 88] nothing cached matching {which!r}. Available: "
                    f"{list(cat['label'])}")
            path = hit.iloc[0]['path']
    with open(path, 'rb') as fh:
        data = _pickle.load(fh)
    px = data.get('prices')
    print(f"  ✓ loaded {_os.path.basename(path)} — "
          f"{px.shape[0]} rows × {px.shape[1]} tickers, no API calls")
    return data


# =============================================================================
# CELL 2: ENGINE
# =============================================================================
# ── [Fix 52] METRICS GLOSSARY ────────────────────────────────────────────────
# The report prints four different R-squared numbers and they are routinely
# confused with each other. They are not competing estimates of one quantity;
# each answers a different question, and the differences between them are the
# most informative part of the report. Printed at the top of section [1].
# [Fix 115] was a hand-aligned block of text passed to print(). Hand alignment
# survives a terminal and nothing else — in a notebook the whole thing
# collapsed into one paragraph. It is data, so it is now rows, and it renders
# through the same table as everything else.
R2_GLOSSARY_ROWS = [
    ('1 · Fitted and scored on the\n    same days  (in-sample R²)',
     'flattered — a sanity check only',
     'The model grades its own homework, so this is always the highest number '
     'on the page. It says the fit converged. It says nothing about whether '
     'the hedge works.'),
    ('2 · Fitted on the past, scored\n    on what came next',
     'honest weights, flattered recipe',
     'Fit on the last N days, score on the days that follow, repeat, average. '
     'The WEIGHTS never see the days they are graded on — but the recipe was '
     'chosen by looking at all of them.'),
    ('3 · This recipe, re-fitted\n    before each test stretch',
     'honest weights AND honest days',
     'The chosen recipe re-fitted strictly on data from before each held-back '
     'stretch. A small bias remains, because the recipe itself was picked '
     'using all the data.'),
    ('4 · The whole method re-run\n    from scratch each time',
     'the number to trade on',
     'Everything — choosing the recipe and fitting the weights — redone using '
     'only data from before each held-back stretch. This is closest to what '
     'running this tool live would actually have delivered.'),
]
R2_GLOSSARY_NOTE = ('Expect them to fall in that order: 1 highest, 4 lowest. '
                    'A big gap between 1 and 4 means the recipe was tuned to '
                    'its own data. Trade on 3 and 4; quote 1 only as a '
                    'caveat.')

# [Fix 116] The report is read by people who do not build these models. Where a
# term is unavoidable — it is on a chart axis, in a column header, or it is what
# a desk would actually say out loud — it is defined here once, in the appendix,
# where it costs nothing to a reader who does not need it.
WORDS_USED = [
    ('Hedge / basket', 'the set of names you trade against your position',
     'Shorting the basket is meant to cancel the part of your position that '
     'moves with the rest of the market.'),
    ('Tracks N% of the moves', 'how much of the position the hedge follows',
     'The statistical name is R². 100% would be a perfect mirror; 0% means '
     'the hedge is unrelated to what you hold.'),
    ('Drift / tracking error', 'the part the hedge does NOT cancel',
     'Quoted as %/yr. This is what you actually carry, and what a risk limit '
     'should be set against.'),
    ('Unseen data / out-of-sample', 'days the model had never looked at',
     'The only honest test. Any score measured on the days a model was fitted '
     'to is a description of the past, not a forecast.'),
    ('Its own data / in-sample', 'the days the model WAS fitted on',
     'Always flattering. Useful as a ceiling and a sanity check, never as an '
     'expectation.'),
    ('Test stretch / window', 'a block of consecutive days held back',
     'The model is fitted on everything before it and graded on it, then the '
     'block moves forward.'),
    ('Lookback', 'how many days of history each refit is fitted on',
     'Short adapts faster and is noisier; long is steadier and staler.'),
    ('Gross / net', 'total position size vs the direction left over',
     'Gross $74 per $100 means you trade $74 of hedge for every $100 of book. '
     'Borrow and financing scale with gross.'),
    ('Market drivers', 'the shared movements pulled out of the universe',
     'Formally principal components. The first one is usually "the whole '
     'market went up"; later ones are sector or style moves.'),
    ('Refit', 'recomputing the weights on newer data',
     'The RE-HEDGE button. Weights go stale as relationships move.'),
    ('1 standard deviation', 'a typical, not a worst, outcome',
     'Two-thirds of the time you land inside it. About 1 time in 20 you are '
     'outside twice it.'),
]
# ─────────────────────────────────────────────────────────────────────────────
# [Fix 42] Robust PCA inputs
#   Ordinary PCA on a sample covariance is optimal only for Gaussian data. Daily
#   equity returns are fat-tailed and skewed, so a handful of jump days can
#   dominate the leading components and distort the loading-space "similarity"
#   used to pick candidates. Two low-cost defenses, applied ONLY to the PCA
#   candidate-SELECTION step (never to the RidgeCV hedge ratios, which stay fit
#   on the true returns):
#     • winsorize_pct : clip each column at its pct / 1-pct quantiles before
#       standardizing, so outliers cannot hijack the eigenvectors.
#     • robust_cov='ledoit' : replace the sample covariance with a Ledoit-Wolf
#       shrinkage estimate (well-conditioned when p >> n, the norm here).
# ─────────────────────────────────────────────────────────────────────────────
def _winsorize_returns(df, pct):
    """Clip every column to its [pct, 1-pct] quantiles. pct=0 -> no-op
    (exact pre-Fix-42 behavior)."""
    if not pct or pct <= 0:
        return df
    lo = df.quantile(pct)
    hi = df.quantile(1.0 - pct)
    return df.clip(lower=lo, upper=hi, axis=1)


class _PCAShim:
    """Uniform PCA result (ordinary or Ledoit-Wolf) exposing the same three
    attributes the rest of the code reads off a fitted sklearn PCA object."""
    __slots__ = ('components_', 'explained_variance_', 'explained_variance_ratio_')

    def __init__(self, components_, explained_variance_, explained_variance_ratio_):
        self.components_ = components_
        self.explained_variance_ = explained_variance_
        self.explained_variance_ratio_ = explained_variance_ratio_


def _pca_fit(tr_sc, n_pc, robust_cov='none'):
    """Fit n_pc components on standardized returns. robust_cov='ledoit'
    eigendecomposes a Ledoit-Wolf shrinkage covariance; otherwise ordinary PCA.
    Returns a _PCAShim so downstream loading code is unchanged."""
    X = np.asarray(tr_sc, dtype=float)
    if robust_cov == 'ledoit':
        from sklearn.covariance import LedoitWolf
        cov = LedoitWolf().fit(X).covariance_
        evals, evecs = np.linalg.eigh(cov)
        idx = np.argsort(evals)[::-1][:n_pc]
        evals = np.clip(evals[idx], 1e-12, None)
        comps = evecs[:, idx].T                       # (n_pc x p)
        total = float(np.trace(cov))
        evr = evals / total if total > 0 else np.zeros_like(evals)
        return _PCAShim(comps, evals, evr)
    pca = PCA(n_components=n_pc).fit(X)
    return _PCAShim(pca.components_, pca.explained_variance_,
                    pca.explained_variance_ratio_)


# ─────────────────────────────────────────────────────────────────────────────
# [Fix 43] Style-factor proxies for the factor-leakage diagnostic. Each factor
# maps to a preferred US proxy and an international fallback; the diagnostic uses
# whichever is actually present in the retained universe.
# ─────────────────────────────────────────────────────────────────────────────
FACTOR_PROXIES = {
    'Market':   ('SPY', 'ACWI'),
    'Size(SC)': ('IWM', 'SIZE'),
    'Momentum': ('MTUM', 'IMTM'),
    'Value':    ('VLUE', 'IVLU'),
    'Quality':  ('QUAL', 'IQLT'),
    'LowVol':   ('USMV', 'ACWV'),
}


def _build_universe_factors(returns, caps_map, stock_cols,
                            mom_skip=21, min_names=6):
    """[Fix 47] Build returns-based factor-mimicking portfolios FROM the
    downloaded universe itself, so factor diagnostics work in EVERY market —
    including markets with no factor ETFs at all (e.g. an ASX name like GMD AU,
    where the US factor ETFs are dropped by the timezone filter).
      MKT(univ) : equal-weight mean of the universe stocks (local market proxy)
      SMB(univ) : small-cap minus big-cap tercile return  (size)   [needs caps]
      WML(univ) : winners minus losers tercile by trailing return  (momentum)
      VOL(univ) : high-vol minus low-vol tercile by realized vol  (volatility)
                  [Fix 48]
      REV(univ) : losers minus winners on the prior 21d return  (short-term
                  reversal, rolling) — the dominant A-share effect  [Fix 55]
    Returns a DataFrame (whichever of MKT/SMB/WML could be built) aligned to
    returns.index; empty if the universe is too small."""
    cols = [c for c in stock_cols if c in returns.columns]
    out = {}
    if len(cols) >= 2:
        out['MKT(univ)'] = returns[cols].mean(axis=1)
    capped = [(c, caps_map.get(c)) for c in cols]
    capped = [(c, m) for c, m in capped
              if m is not None and np.isfinite(m) and m > 0]
    if len(capped) >= min_names:
        capped.sort(key=lambda x: x[1])
        k = max(1, len(capped) // 3)
        small = [c for c, _ in capped[:k]]
        big = [c for c, _ in capped[-k:]]
        out['SMB(univ)'] = (returns[small].mean(axis=1)
                            - returns[big].mean(axis=1))
    if len(cols) >= min_names and len(returns) > mom_skip + 20:
        trail = returns[cols].iloc[:-mom_skip].sum().dropna().sort_values()
        if len(trail) >= min_names:
            k = max(1, len(trail) // 3)
            los = trail.index[:k].tolist()
            win = trail.index[-k:].tolist()
            out['WML(univ)'] = (returns[win].mean(axis=1)
                                - returns[los].mean(axis=1))
    if len(cols) >= min_names:                      # [Fix 48] volatility factor
        vol = returns[cols].std().dropna().sort_values()
        if len(vol) >= min_names:
            k = max(1, len(vol) // 3)
            lowv = vol.index[:k].tolist()
            highv = vol.index[-k:].tolist()
            out['VOL(univ)'] = (returns[highv].mean(axis=1)
                                - returns[lowv].mean(axis=1))
    # [Fix 55] REV(univ) — short-term (1-month) reversal.
    #   In A-shares 12-1 momentum (the WML factor above) is famously weak or
    #   outright inverted, while one-month reversal is the dominant
    #   cross-sectional effect. Without this factor the leakage diagnostic
    #   has no way to report the single largest style exposure an onshore
    #   China basket can carry. Unlike the static WML/VOL sorts this factor
    #   is formed on a ROLLING basis — each day, losers-minus-winners on the
    #   PRIOR 21d return, shifted by one day so nothing is known in advance.
    #   The first ~22 rows are NaN; _factor_leakage's finite-mask handles it.
    if len(cols) >= min_names and len(returns) > 42:
        past = returns[cols].rolling(21).sum().shift(1)
        lo_q = past.quantile(1.0 / 3.0, axis=1)
        hi_q = past.quantile(2.0 / 3.0, axis=1)
        rev = (returns[cols].where(past.le(lo_q, axis=0)).mean(axis=1)
               - returns[cols].where(past.ge(hi_q, axis=0)).mean(axis=1))
        out['REV(univ)'] = rev
    return pd.DataFrame(out)


def _factor_leakage(returns, target, weights, caps_map=None, stock_cols=None):
    """[Fix 43/47] Returns-based factor-exposure check. Regress the target AND
    the hedged residual (target - basket.weights) on a set of factors, and
    compare the betas: a residual beta near 0 means the hedge neutralized that
    factor; near the target beta means it did not.

    [Fix 47] Market / Size / Momentum come from factor-mimicking portfolios
    BUILT FROM THE UNIVERSE (so they exist in every market, ETF or not);
    Value / Quality / LowVol come from style ETFs when those survive in the
    universe. [Fix 50] Factors are sequentially orthogonalized before the
    regression so collinearity (e.g. SMB vs VOL) cannot inflate the betas.
    Returns a DataFrame, or None if fewer than two factors are available."""
    present = {}
    for fac in ('Value', 'Quality', 'LowVol'):          # ETF-only factors
        for p in FACTOR_PROXIES[fac]:
            hit = next((c for c in returns.columns
                        if c.split(' ')[0] == p and c != target), None)
            if hit is not None:
                present[fac] = returns[hit]
                break
    if caps_map is not None and stock_cols is not None:  # [Fix 47] local factors
        uf = _build_universe_factors(returns, caps_map, stock_cols)
        for name in uf.columns:
            present[name] = uf[name]
    if len(present) < 2:                                 # fallback: ETF market/size/mom
        for fac in ('Market', 'Size(SC)', 'Momentum'):
            for p in FACTOR_PROXIES[fac]:
                hit = next((c for c in returns.columns
                            if c.split(' ')[0] == p and c != target), None)
                if hit is not None:
                    present[fac] = returns[hit]
                    break
    if len(present) < 2:
        return None
    # [Fix 50] sequential orthogonalization. The factors overlap heavily —
    # small caps ARE more volatile (SMB ~ VOL), style ETFs share market beta —
    # and collinear regressors inflate/destabilize individual betas (this is
    # why adding global factor ETFs can make the printed betas JUMP UP).
    # Regress each factor on the ones before it (fixed order) and keep the
    # residual: each beta then reads as the INCREMENTAL exposure not already
    # explained by the factors listed above it. Diagnostics only — hedge
    # selection and weights are untouched.
    # [Fix 55] REV(univ) sits AFTER WML so its beta reads as reversal exposure
    # not already explained by market / size / momentum.
    _ORDER = ['MKT(univ)', 'SMB(univ)', 'WML(univ)', 'REV(univ)', 'VOL(univ)',
              'Value', 'Quality', 'LowVol', 'Market', 'Size(SC)', 'Momentum']
    names = ([k for k in _ORDER if k in present]
             + [k for k in present if k not in _ORDER])
    y_t = returns[target].values
    basket = [c for c in weights.index if c in returns.columns]
    resid = y_t - returns[basket].values.dot(weights.reindex(basket).values)
    Fraw = np.column_stack([np.asarray(present[n], dtype=float) for n in names])
    m = np.isfinite(Fraw).all(axis=1) & np.isfinite(y_t) & np.isfinite(resid)
    Fo = np.empty((int(m.sum()), len(names)))
    for j in range(len(names)):
        f = Fraw[m, j]
        if j:
            Xp = np.column_stack([np.ones(len(f)), Fo[:, :j]])
            f = f - Xp @ np.linalg.lstsq(Xp, f, rcond=None)[0]
        Fo[:, j] = f
    Xc = np.column_stack([np.ones(len(Fo)), Fo])
    bt = np.linalg.lstsq(Xc, y_t[m], rcond=None)[0][1:]
    br = np.linalg.lstsq(Xc, resid[m], rcond=None)[0][1:]
    return pd.DataFrame({'factor': names,
                         'target_beta': np.round(bt, 3),
                         'residual_beta': np.round(br, 3)})


PORTFOLIO_COL = '__PORTFOLIO__'   # [Fix 46] synthetic column for a multi-name book


def _coerce_targets(target):
    """[Fix 46/64] Normalize the target argument into a hedging job.

    THE FORMAT CONTRACT — three accepted shapes:

      1. str                       single-name hedge
           'AAPL US Equity'

      2. list / tuple of str       EQUAL-WEIGHT LONG book (all legs long,
           ['AAPL US Equity',      same size). Short tickers are fine —
            'NVDA US Equity',      'AAPL' / 'NVDA US' / '700 HK' are
            'MSFT US Equity']      resolved by _fmt_bbg_ticker.

      3. dict {ticker: weight}     WEIGHTED book. Weights are RELATIVE
           {'AAPL US Equity': +0.5,   NOTIONALS in any consistent unit —
            'NVDA US Equity': +0.3,   dollars, percent of book, or plain
            'MSFT US Equity': -0.2}   ratios; sign = direction (+ long,
                                      - short). They do NOT need to sum
                                      to 1: everything is normalized by
                                      GROSS (sum of |w|), which keeps a
                                      market-neutral book well defined
                                      (net-sum normalization would divide
                                      by ~0).

    What the engine then does with it: the book becomes ONE synthetic
    return series, sum(w_i * r_i) per day, and every downstream number —
    R², vol reduction, and the hedge recipe itself — refers to that NET
    series. "per $100" in the recipe means per $100 of GROSS book value.
    So for {'AAPL': +50, 'NVDA': +30, 'MSFT': -20} the book is normalized
    to (+0.5, +0.3, -0.2) and a recipe line of "short $80 QQQ" means:
    for every $100 of gross notional you hold across the three legs,
    short $80 of QQQ against the NET exposure.

    Returns (legs, weights_dict, is_portfolio, label)."""
    if isinstance(target, str):
        # [Fix 67] normalize AND validate. A bare short name ('AAPL') has no
        # exchange code, cannot be resolved by Bloomberg, and used to sail
        # through here only to die inside bbg_hist with a cryptic "no price
        # data" error. Minimum viable form: 'AAPL US'.
        tk = _fmt_bbg_ticker(target)
        if tk is None:
            raise ValueError(
                f"[Fix 67] '{target}' has no exchange code — Bloomberg "
                f"needs one: write 'AAPL US', 'AAPL US Equity' or '700 HK'")
        return [tk], {}, False, tk.split(' ')[0]
    # [Fix 64] fail loudly on malformed books instead of mid-pipeline
    if isinstance(target, dict):
        if not target:
            raise ValueError("[Fix 64] portfolio dict is empty")
        _bad = [str(k) for k in target if _fmt_bbg_ticker(str(k)) is None]
        if _bad:                                              # [Fix 67]
            raise ValueError(
                f"[Fix 67] ticker(s) {_bad} have no exchange code — "
                f"Bloomberg needs one: write 'AAPL US', 'AAPL US Equity' "
                f"or '700 HK', never a bare short name")
        try:
            items = [(_fmt_bbg_ticker(str(k)), float(v))
                     for k, v in target.items()]
        except (TypeError, ValueError):
            raise ValueError(
                "[Fix 64] portfolio dict values must be numbers "
                "(relative notionals, sign = direction) — got "
                f"{ {k: type(v).__name__ for k, v in target.items()} }")
        if all(v == 0.0 for _, v in items):
            raise ValueError("[Fix 64] all portfolio weights are zero")
    elif isinstance(target, (list, tuple)):
        if not target:
            raise ValueError("[Fix 64] portfolio list is empty")
        if not all(isinstance(k, str) for k in target):
            raise ValueError(
                "[Fix 64] a portfolio LIST must contain tickers only "
                "(equal-weight long book). For custom weights or shorts "
                "use a dict: {'AAPL US Equity': +0.5, 'MSFT US Equity': "
                "-0.5}")
        _bad = [str(k) for k in target if _fmt_bbg_ticker(str(k)) is None]
        if _bad:                                              # [Fix 67]
            raise ValueError(
                f"[Fix 67] ticker(s) {_bad} have no exchange code — "
                f"Bloomberg needs one: write 'AAPL US', 'AAPL US Equity' "
                f"or '700 HK', never a bare short name")
        items = [(_fmt_bbg_ticker(str(k)), 1.0) for k in target]
    else:
        raise ValueError(
            f"[Fix 64] target must be a str, list of str, or dict of "
            f"ticker->weight — got {type(target).__name__}")
    legs = [k for k, _ in items]
    if len(set(legs)) != len(legs):
        raise ValueError(f"[Fix 64] duplicate legs after ticker "
                         f"normalization: {legs}")
    if len(legs) == 1:
        # a one-leg "book" is just a single-name hedge — treat it as one
        return [legs[0]], {}, False, legs[0].split(' ')[0]
    gross = sum(abs(v) for _, v in items) or 1.0
    weights = {k: v / gross for k, v in items}
    label = 'PF[' + '+'.join(k.split(' ')[0] for k in legs) + ']'
    return legs, weights, True, label


def _validation_tier(rows, lookback_options, cv_test_window, cv_n_folds,
                     outer_window, outer_folds, purge, basket_size_options,
                     rolling_refit_every):
    """[Fix 77/85] Resolve the validation protocol for a window of `rows`
    return rows. Extracted verbatim from find_best_hedge in [Fix 85] so that
    preview_window() and the GUI describe the protocol that will ACTUALLY
    run — a second copy of this logic would drift from the engine.

      FULL      enough rows — the full walk-forward + nested protocol.
      REDUCED   walk-forward CV + outer windows with shorter lookbacks /
                test windows and fewer folds. Wider error bars.
      HOLDOUT   too short for walk-forward: fit once on the first ~80%,
                score ONCE on the held-out tail. Thin but honest.
      IN-SAMPLE under ~60 rows: no OOS evidence is possible. Descriptive
                only; the basket is capped and the ticket says so loudly.

    Returns a dict of resolved settings plus 'notes' (lines to print)."""
    if rows < 8:
        raise ValueError(f"[Fix 77] only {rows} return rows in the window — "
                         f"need at least 8 even for a descriptive fit")
    need_full = (max(lookback_options) + cv_test_window * cv_n_folds
                 + outer_window * outer_folds)
    t = {'mode': 'FULL', 'rows': rows, 'need_full': need_full, 'hold_rows': 0,
         'lookback_options': tuple(lookback_options),
         'cv_test_window': cv_test_window, 'cv_n_folds': cv_n_folds,
         'outer_window': outer_window, 'outer_folds': outer_folds,
         'min_train_rows': 50,          # [Fix 80] inner-CV minimum train block
         'rolling_refit_every': rolling_refit_every,
         'basket_size_options': tuple(basket_size_options), 'notes': []}
    if rows < need_full:
        _tw = int(np.clip(rows // 8, 10, cv_test_window))
        _nf = int(min(cv_n_folds, 3))
        _of = int(min(outer_folds, 2))
        _lb_max = rows - (_tw + purge) * (_nf + _of)
        if _lb_max >= 30:
            t['mode'] = 'REDUCED'
            t['cv_test_window'] = t['outer_window'] = _tw
            t['cv_n_folds'], t['outer_folds'] = _nf, _of
            t['lookback_options'] = tuple(sorted({max(30, _lb_max // 2),
                                                  _lb_max}))
            # [Fix 80] the inner CV skips folds whose training block is
            # shorter than min_train_rows; with REDUCED lookbacks the fixed
            # 50-row floor rejected every fold. Track the tier's own floor.
            t['min_train_rows'] = min(50, min(t['lookback_options']))
            t['rolling_refit_every'] = min(int(rolling_refit_every), _tw)
            t['notes'] += [
                f"  [Fix 77] validation mode REDUCED — {rows} rows < "
                f"{need_full} needed for the full protocol",
                f"    → lookbacks {t['lookback_options']}, CV {_nf}×{_tw}d,"
                f" outer {_of}×{_tw}d; fewer/shorter windows ⇒ wider error"
                f" bars — mind the CI"]
        elif rows >= 60:
            t['mode'] = 'HOLDOUT'
            t['hold_rows'] = max(12, rows // 5)
            t['notes'] += [
                f"  [Fix 77] validation mode HOLDOUT — {rows} rows is too "
                f"short for walk-forward CV",
                f"    → fit on the first {rows - t['hold_rows'] - purge} "
                f"rows, score ONCE on the last {t['hold_rows']} unseen rows"]
        else:
            t['mode'] = 'IN-SAMPLE'
            t['notes'].append(
                f"  [Fix 77] validation mode IN-SAMPLE — {rows} rows: no "
                f"out-of-sample evidence is possible at this length; treat "
                f"every number as descriptive")
    # cap the basket so the fit window keeps ≥ ~8 obs per regressor
    _fit_len = (min(rows, max(t['lookback_options']))
                if t['mode'] in ('FULL', 'REDUCED') else rows - t['hold_rows'])
    _bs_cap = max(2, _fit_len // 8)
    if max(basket_size_options) > _bs_cap:
        t['basket_size_options'] = tuple(sorted({min(b, _bs_cap)
                                                 for b in basket_size_options}))
        t['notes'].append(f"  [Fix 77] basket sizes capped at {_bs_cap} "
                          f"({_fit_len}-row fit window)")
    return t


TIER_MEANING = {
    'FULL': 'full protocol: walk-forward CV, nested windows, rolling curve, '
            'bootstrap CI — a tradeable verdict',
    'REDUCED': 'walk-forward still runs but on shorter/fewer windows — '
               'a usable verdict with wider error bars',
    'HOLDOUT': 'one held-out tail only — thin evidence; size down and '
               'confirm on a longer window',
    'IN-SAMPLE': 'NO out-of-sample evidence is possible — descriptive only, '
                 'do not size off it',
}


def preview_window(data, sample_window=None, lookback_options=(126, 252),
                   cv_test_window=42, cv_n_folds=5, outer_window=None,
                   outer_folds=3, basket_size_options=(3, 5, 7),
                   quiet=False):
    """[Fix 85] What will this sample_window actually buy me? Reports the
    session count and the validation tier BEFORE you spend minutes on a run,
    using the same _validation_tier the engine uses. Cheap and API-free."""
    idx = data['prices'].index
    if sample_window is not None:
        s, e = _parse_sample_window(sample_window, idx)
        sub = idx[(idx >= s) & (idx <= e)]
    else:
        sub = idx
    ow = outer_window or cv_test_window
    out = {'sessions': len(sub),
           'start': (sub[0] if len(sub) else None),
           'end': (sub[-1] if len(sub) else None)}
    if len(sub) < 10:
        out['mode'] = None
        if not quiet:
            print(f"  ⚠ only {len(sub)} sessions in that window — the engine "
                  f"needs ≥ 10 (and ≥ 8 return rows).")
        return out
    modes = {}
    for pu, lab in ((0, 'daily'), (1, '2-day (cross-timezone)')):
        try:
            t = _validation_tier(len(sub) - 1 - pu, lookback_options,
                                 cv_test_window, cv_n_folds, ow, outer_folds,
                                 pu, basket_size_options, 21)
            modes[lab] = t['mode']
        except ValueError:
            modes[lab] = None
    out['mode'] = modes.get('daily')
    out['mode_two_day'] = modes.get('2-day (cross-timezone)')
    if not quiet:
        print(f"  {out['sessions']} sessions "
              f"({out['start']:%d %b %Y} → {out['end']:%d %b %Y})")
        if out['mode'] == out['mode_two_day']:
            print(f"  → validation tier: {out['mode']}")
        else:
            print(f"  → validation tier: {out['mode']} (daily) / "
                  f"{out['mode_two_day']} if the universe is cross-timezone")
        for m in dict.fromkeys(v for v in modes.values() if v):
            print(f"     {m:<9} {TIER_MEANING[m]}")
    return out


def _parse_sample_window(sample_window, idx):
    """[Fix 76] Resolve `sample_window` into (start, end) Timestamps.

    Accepted forms:
      '3m' / '90d' / '26w' / '1y'      last N months/days/weeks/years up to
                                       the END of the downloaded data
      (start, end) tuple               absolute dates; each entry is anything
                                       pd.Timestamp parses ('2026-07-13',
                                       '13-Jul-2026', a Timestamp) or None
                                       for "no bound on this side"
    """
    import re as _re
    if isinstance(sample_window, str):
        m = _re.match(r'^\s*(\d+)\s*([dwmy])\s*$', sample_window.lower())
        if not m:
            raise ValueError(
                f"[Fix 76] sample_window string must look like '3m', '90d', "
                f"'26w' or '1y' — got {sample_window!r}. For absolute dates "
                f"pass a (start, end) tuple.")
        n, u = int(m.group(1)), m.group(2)
        end = idx[-1]
        off = {'d': pd.Timedelta(days=n), 'w': pd.Timedelta(weeks=n),
               'm': pd.DateOffset(months=n), 'y': pd.DateOffset(years=n)}[u]
        return pd.Timestamp(end - off), pd.Timestamp(end)
    if isinstance(sample_window, (tuple, list)) and len(sample_window) == 2:
        s, e = sample_window
        try:
            start = pd.Timestamp(s) if s is not None else pd.Timestamp(idx[0])
            end = pd.Timestamp(e) if e is not None else pd.Timestamp(idx[-1])
        except (ValueError, TypeError):
            raise ValueError(f"[Fix 76] cannot parse sample_window dates "
                             f"{sample_window!r} — use e.g. "
                             f"('2026-07-13', '2026-07-28')")
        if start >= end:
            raise ValueError(f"[Fix 76] sample_window start {start:%Y-%m-%d} "
                             f"is not before end {end:%Y-%m-%d}")
        return start, end
    raise ValueError(f"[Fix 76] sample_window must be a string like '3m' or "
                     f"a (start, end) tuple — got "
                     f"{type(sample_window).__name__}")


def _verdict_banner(grade, reasons, subtitle=''):
    """[Fix 87] Colour-coded verdict banner in Jupyter; plain text elsewhere.
    Any rendering failure falls through to the text version, so the report
    can never be lost to a display problem.

    [Fix 115] Was set in monospace, which made the one block a reader looks at
    first the least legible thing on the page. It now uses the report's own
    type scale, and it records itself into the transcript like every other
    block."""
    _kv_flush()
    pal = {'✓': ('#1b6b3a', '#edf7f0'), '△': ('#8a5a00', '#fff8ea'),
           '✗': ('#8c1c1c', '#fdeeee'), '?': ('#3d4650', '#eff2f5')}
    grade, subtitle = _clean(grade), _clean(subtitle)
    reasons = [_clean(r) for r in reasons[:4]]
    fg, bg = pal.get(grade[:1], pal['?'])
    _tx('', f"  VERDICT   {grade}")
    if subtitle:
        _tx(f"            {subtitle}")
    _tx(*[f"            · {r}" for r in reasons])
    try:
        if _IN_JUPYTER:
            li = ''.join(
                f'<li style="margin:2px 0">{_html.escape(r)}</li>'
                for r in reasons)
            _ipy_display(_HTML(
                f'<div style="border-left:5px solid {fg};background:{bg};'
                f'padding:12px 16px;margin:10px 0 14px 0;border-radius:3px;'
                f'font-family:{_RF_UI};max-width:{_RF_MAXW}">'
                f'<div style="font-size:17px;font-weight:700;color:{fg};'
                f'letter-spacing:.01em">{_html.escape(grade)}</div>'
                + (f'<div style="font-size:12.5px;color:{_RF_INK};'
                   f'margin-top:3px;line-height:1.5">'
                   f'{_html.escape(subtitle)}</div>' if subtitle else '')
                + f'<ul style="margin:8px 0 0 0;padding-left:18px;'
                  f'font-size:12px;color:{_RF_MUTE};line-height:1.5">'
                  f'{li}</ul></div>'))
            return
    except Exception:
        pass
    print(f"\n  VERDICT   {grade}")
    if subtitle:
        print(f"            {subtitle}")
    for r in reasons:
        print(f"            · {r}")


def _fmt_money_mm(x_mm):
    """[Fix 101] Format a USD-millions figure the way a desk reads it:
    $2.51mm, $316k, $8.40. Avoids '$2,506k'."""
    a = abs(float(x_mm))
    if a >= 1.0:
        return f"${a:,.2f}mm"
    if a >= 0.001:
        return f"${a*1e3:,.0f}k"
    return f"${a*1e6:,.0f}"


def _ordinal(n):
    """[Fix 101] 1st / 2nd / 3rd / 33rd / 11th."""
    n = int(round(n))
    if 10 <= n % 100 <= 20:
        suf = 'th'
    else:
        suf = {1: 'st', 2: 'nd', 3: 'rd'}.get(n % 10, 'th')
    return f"{n}{suf}"


def _horizon_sigma(resid, h, te_ann, two_day):
    """[Fix 92] Expected 1-sigma hedge slippage over an h-row holding period.

    The naive answer is te_ann * sqrt(h/252) — correct only if residuals are
    serially uncorrelated. If they trend (an under-hedged drifting factor),
    h-day risk is LARGER than the daily numbers imply. The variance ratio
    VR = Var(sum of h residuals) / (h * Var(single residual)), estimated on
    overlapping h-row sums, corrects this; VR>1 means residuals compound.
    For 2-day overlapping rows the construction already double-counts, so
    only the iid scaling is used there (VR would be biased). Returns
    (sigma_over_h, vr) with vr=NaN when not estimable."""
    base = te_ann * np.sqrt(h / 252.0)
    r = np.asarray(resid, dtype=float)
    r = r[np.isfinite(r)]
    if two_day or h < 2 or len(r) < max(3 * h, 40):
        return base, np.nan
    s = pd.Series(r)
    rh = s.rolling(h).sum().dropna().values
    v1 = float(np.var(r))
    if v1 <= 1e-18 or len(rh) < 20:
        return base, np.nan
    vr = float(np.var(rh) / (h * v1))
    vr = float(np.clip(vr, 0.25, 4.0))          # tail-guard the estimate
    # ── [Fix 111] THE VR MAY RAISE THE RISK NUMBER, NEVER LOWER IT ─────────
    # Measured (synthetic, 1000 basket/seed pairs, train 500 rows, score the
    # next 500): the in-sample variance ratio has essentially NO
    # out-of-sample persistence —
    #     corr(VR_in, VR_out)     = -0.32 .. +0.18   across h = 21/63/126
    # while the TRUE ratio matters a great deal —
    #     corr(VR_out, realised)  = +0.54 .. +0.86
    # so the quantity is real and large (p10-p90 spans 0.67-1.24 at h=21 and
    # 0.32-1.32 at h=126) but is NOT forecastable from the sample in hand.
    #
    # For SELECTION that is fatal, and it is why the horizon deliberately
    # does not enter the config search — see the [Fix 100]/[Fix 112] note in
    # find_best_hedge. For REPORTING there is no winner's curse, so a noisy
    # estimate is tolerable in the conservative direction only:
    #   vr > 1  keep it. Residuals that compounded are a warning, and a
    #           warning that turns out to be unnecessary costs nothing.
    #   vr < 1  DISCARD it. This is the dangerous branch: a vr of 0.25 —
    #           inside the old clip — HALVED the reported h-day risk on the
    #           strength of an estimate with zero forecasting power. There is
    #           no version of prudent risk reporting in which that is right.
    # The raw estimate is still returned so the report can show it.
    return base * np.sqrt(max(vr, 1.0)), vr


def _block_bootstrap_ci(y, resid, n_boot=2000, block=None, seed=0, ann=1.0):
    """[Fix 70] Circular block bootstrap CI for the pooled OOS R-squared and
    the annualized tracking error. Hedge residuals are serially dependent
    (volatility clustering; 2-day overlapping returns are MA(1) by
    construction), so an iid bootstrap understates uncertainty. Contiguous
    blocks of (target, residual) PAIRS are resampled with wrap-around, which
    preserves both the dependence and the y-vs-residual pairing. Block length
    defaults to ~n^(1/3) (Politis-Romano rate) floored at 5; pass block=8+
    for 2-day overlapping returns. Percentile intervals; the seed is fixed so
    two runs of the same report print the same CI. NOTE: the CI conditions on
    the fitted weights — refit noise sits ON TOP of this interval, so read it
    as a floor on the uncertainty, not the whole of it. Returns dict or None
    (series too short / degenerate)."""
    y = np.asarray(y, dtype=float)
    r = np.asarray(resid, dtype=float)
    n = len(y)
    if n < 30 or len(r) != n:
        return None
    b = int(block) if block else max(5, int(round(n ** (1.0 / 3.0))))
    b = min(b, n)
    k = int(np.ceil(n / b))
    rng = np.random.default_rng(seed)
    starts = rng.integers(0, n, size=(int(n_boot), k))
    idx = (starts[:, :, None] + np.arange(b)[None, None, :]) % n
    idx = idx.reshape(int(n_boot), -1)[:, :n]
    ys, rs = y[idx], r[idx]
    sst = ((ys - ys.mean(axis=1, keepdims=True)) ** 2).sum(axis=1)
    sse = (rs ** 2).sum(axis=1)
    with np.errstate(divide='ignore', invalid='ignore'):
        r2s = np.where(sst > 1e-12, 1.0 - sse / sst, np.nan)
    tes = rs.std(axis=1) * ann
    r2s = r2s[np.isfinite(r2s)]
    if not len(r2s):
        return None
    return {'r2_lo': float(np.percentile(r2s, 2.5)),
            'r2_hi': float(np.percentile(r2s, 97.5)),
            'te_lo': float(np.percentile(tes, 2.5)),
            'te_hi': float(np.percentile(tes, 97.5)),
            'block': b, 'n': n, 'n_boot': int(n_boot)}


def _window_dispersion_stats(detail, pooled, n_eff):
    """[Fix 72] Pure computation behind section [1b] and the trade ticket:
    noise band, per-window flags, vol-regime correlation, weight-stability
    cosine, deterioration trend, worst-window R-squared and tracking error.
    Split out of the print block so the ticket can consume the SAME numbers
    the diagnostics print — no chance of the two drifting apart."""
    out = {'band': None, 'flags': [], 'reg_corr': np.nan, 'min_sim': None,
           'trend': False, 'worst_r2': np.nan, 'worst_te': np.nan}
    if detail is None or not len(detail):
        return out
    out['worst_r2'] = float(detail['r2'].min())
    te_w = detail['tgt_vol'] * (1.0 - detail['vol_red'])
    out['worst_te'] = float(te_w.max())
    if not np.isnan(pooled) and pooled > 0:
        z = np.arctanh(min(np.sqrt(pooled), 0.999))
        zsd = 1.96 / np.sqrt(max(n_eff - 3, 1))
        out['band'] = (float(np.tanh(z - zsd) ** 2),
                       float(np.tanh(z + zsd) ** 2))
    band = out['band']
    for _, r in detail.iterrows():
        if band is None:
            out['flags'].append('')
        elif band[0] <= r['r2'] <= band[1]:
            out['flags'].append('noise-range')
        else:
            out['flags'].append('BELOW band' if r['r2'] < band[0]
                                else 'above band')
    if detail['tgt_vol'].nunique() > 1:
        out['reg_corr'] = float(np.corrcoef(detail['tgt_vol'],
                                            detail['r2'])[0, 1])
    if len(detail) >= 2:
        names = sorted(set().union(*[set(w.index)
                                     for w in detail['weights']]))
        W = np.array([w.reindex(names).fillna(0.0).values
                      for w in detail['weights']])
        sims = []
        for a in range(len(W)):
            for b in range(a + 1, len(W)):
                na, nb = np.linalg.norm(W[a]), np.linalg.norm(W[b])
                if na > 1e-12 and nb > 1e-12:
                    sims.append(float(W[a].dot(W[b]) / (na * nb)))
        if sims:
            out['min_sim'] = min(sims)
    if len(detail) >= 3:
        tr = float(np.corrcoef(detail['window'], detail['r2'])[0, 1])
        out['trend'] = bool(tr > 0.7 and detail.loc[
            detail['window'] == 1, 'r2'].iloc[0] == detail['r2'].min())
    return out


def _ticket_verdict(quality, worst_r2, overfit_gap, min_sim, trend,
                    ref_pca, ref_simple, roll_breach):
    """[Fix 72] One consolidated go / reduce / no-go verdict for the ticket.

    Gates (heuristics, deliberately conservative — override them with
    judgment, not by re-running until the box says yes):
      ✗  pooled OOS R2 < 0.20 (vol cut under ~11%; costs likely eat it), OR
         in-sample-vs-OOS gap >= 0.35 (clear overfit), OR any window with
         NEGATIVE OOS R2 (the hedge ADDED risk there), OR no OOS evidence.
      △  pooled R2 < 0.45, worst window < 0.15, gap >= 0.15, basket turnover
         across windows (min weight cosine <= 0.5), R2 deteriorating into
         the present, live rolling TE already past the kill-switch, or a
         single instrument matching the basket OOS.
      ✓  none of the above.
    Returns (grade, [reasons])."""
    if quality is None or np.isnan(quality):
        return ('? INSUFFICIENT OOS EVIDENCE',
                ['no valid unseen windows — read every number here as '
                 'in-sample'])
    reasons, fail, warn = [], False, False
    if quality < 0.20:
        fail = True
        reasons.append(f'pooled OOS R² {quality:.2f} < 0.20 — too little to '
                       f'pay for')
    if not np.isnan(overfit_gap) and overfit_gap >= 0.35:
        fail = True
        reasons.append(f'in-sample vs OOS gap {overfit_gap:.2f} ≥ 0.35 — '
                       f'clearly overfit')
    if not np.isnan(worst_r2) and worst_r2 < 0.0:
        fail = True
        reasons.append(f'worst window OOS R² {worst_r2:.2f} < 0 — the hedge '
                       f'ADDED risk there')
    if roll_breach:
        warn = True
        reasons.append('latest rolling TE is already above the kill level — '
                       'do not add risk on this hedge today')
    if not fail:
        if quality < 0.45:
            warn = True
            reasons.append(f'pooled OOS R² {quality:.2f} is moderate')
        if not np.isnan(worst_r2) and worst_r2 < 0.15:
            warn = True
            reasons.append(f'worst window OOS R² {worst_r2:.2f} — size off '
                           f'that window, not the average')
        if not np.isnan(overfit_gap) and overfit_gap >= 0.15:
            warn = True
            reasons.append(f'in-sample vs OOS gap {overfit_gap:.2f} — some '
                           f'overfitting')
        if min_sim is not None and min_sim <= 0.5:
            warn = True
            reasons.append(f'windows fitted materially different baskets '
                           f'(min weight cosine {min_sim:+.2f})')
        if trend:
            warn = True
            reasons.append('R² deteriorates toward the LATEST window — '
                           'trust the newest reading, not the mean')
    if (not np.isnan(ref_pca) and not np.isnan(ref_simple)
            and ref_simple >= ref_pca):
        warn = True
        reasons.append('a single instrument matches the basket OOS — prefer '
                       'the simple hedge (see [2])')
    grade = ('✗ DO NOT TRADE AS-IS' if fail
             else ('△ TRADE AT REDUCED SIZE' if warn else '✓ TRADE'))
    if not reasons:
        reasons.append('all gates passed — OOS strong, worst window fine, '
                       'no overfit, stable basket')
    return grade, reasons


def _rank_scores(loadings, returns, target, universe):
    """[Fix 10] z-score eucl / cos / corr before combining into rank_score
    (lower = better hedge candidate).

    [Fix 68] rank on |corr| and |cos|, with the loading distance folded over
    the sign (min of the distance to +v and to -v). Ridge assigns negative
    weights happily and the recipe prints LONG legs, yet the old score sent
    a perfectly ANTI-correlated candidate to the bottom of the ranking
    (corr entered signed with weight -2, so corr=-0.9 read as the worst
    possible evidence) while the single-instrument benchmarks pick by
    |corr| — the basket was denied instruments the benchmarks were allowed.
    The returned 'corr' column stays SIGNED so direction is visible."""
    tgt_load = loadings.loc[target].values
    tgt_norm = np.linalg.norm(tgt_load)
    raw = {}
    for tk in universe:
        if tk not in loadings.index:
            continue
        v = loadings.loc[tk].values
        # [Fix 68] fold the sign: an anti-correlated name sits near -v
        eucl = min(np.linalg.norm(tgt_load - v), np.linalg.norm(tgt_load + v))
        cos = abs(float(np.dot(tgt_load, v) /
                        (tgt_norm * np.linalg.norm(v) + 1e-10)))
        corr = returns[target].corr(returns[tk])
        raw[tk] = {'eucl': eucl, 'cos': cos, 'corr': 0.0 if pd.isna(corr) else corr}
    if not raw:
        return pd.DataFrame(columns=['eucl', 'cos', 'corr', 'rank_score'])
    df = pd.DataFrame(raw).T
    zin = df.assign(corr=df['corr'].abs())        # [Fix 68] rank on |corr|
    z = (zin - zin.mean()) / zin.std(ddof=0).replace(0, 1.0)
    # [Fix 24] realized correlation is the most direct evidence of hedging
    # power — give it double weight so the two (noisier) loading-space
    # metrics cannot outvote it
    df['rank_score'] = z['eucl'] - z['cos'] - 2.0 * z['corr']
    return df.sort_values('rank_score')


def _mask_stale(prices, max_flat_days=7):
    """[Fix 56] Mask stale (forward-filled) prices — the suspension guard.

    A-share names halt for weeks at a time; the forward fill keeps the last
    print, producing long runs of identical prices and therefore zero
    returns. Those zeros dilute every correlation and beta the selection
    relies on: a suspended stock looks like a LOW-beta stock and can sneak
    into the basket (or push a genuinely better name out of it). Any price
    unchanged for more than `max_flat_days` consecutive sessions is masked,
    so the complete-window filters downstream exclude the affected stretch
    instead of learning from stale data.

    The default (7) is deliberately longer than any holiday cluster — Golden
    Week and CNY fills are ~5 weekdays — so ordinary market holidays are
    left untouched. Set mask_stale_days=None in find_best_hedge to disable.
    """
    out = prices.copy()
    n_masked = 0
    for c in out.columns:
        s = out[c]
        flat = s.eq(s.shift())
        grp = (~flat).cumsum()            # one group per constant-price run
        run = flat.groupby(grp).cumsum()
        bad_grp = set(grp[run > max_flat_days])
        if bad_grp:
            out[c] = s.mask(grp.isin(bad_grp) & flat)   # mask the whole run
            n_masked += 1
    if n_masked:
        print(f"  [Fix 56] stale-price guard: masked suspension-like flat "
              f"runs (>{max_flat_days} sessions) in {n_masked} names")
    return out


def _fit_basket(returns, target, basket):
    """[Fix 11/59] RidgeCV handles the highly collinear baskets this method
    produces. Returns (model, in-sample R², residuals).

    [Fix 59] The regressors are vol-standardized before the fit. RidgeCV
    applies ONE penalty to all coefficients, so on raw returns a 45-vol small
    cap and an 11-vol index ETF face very different EFFECTIVE shrinkage and
    the penalty quietly tilts baskets toward high-vol names. Scaling each
    column by its own standard deviation equalizes the penalty; dividing the
    fitted coefficients by the same scale afterwards gives an EXACTLY
    equivalent model in raw-return space, so .predict / .score and every
    downstream call site (weights, gross, per-$100 sizing) are unchanged in
    meaning — only the shrinkage is fairer."""
    X = returns[basket].values
    y = returns[target].values
    sd = X.std(axis=0)
    sd[sd < 1e-12] = 1.0                  # constant column → no rescaling
    reg = RidgeCV(alphas=RIDGE_ALPHAS).fit(X / sd, y)
    reg.coef_ = reg.coef_ / sd            # map back to raw-return space
    return reg, reg.score(X, y), y - reg.predict(X)


def _n_pc_auto(ret_scaled, cap, robust_cov='none'):
    """[Fix 22] Number of significant PCs via the Marchenko-Pastur upper edge.

    [Fix 73] the count is now ALWAYS taken on the raw sample spectrum. The
    MP edge describes the bulk of SAMPLE eigenvalues under a pure-noise
    null; Ledoit-Wolf shrinkage pulls large sample eigenvalues toward the
    grand mean, so the old [Fix 42] branch — shrunk eigenvalues against the
    unshrunk edge — systematically UNDER-counted factors, and
    robust_cov='ledoit' silently ran with fewer PCs than the same data under
    'none'. Shrinkage still stabilizes the eigenVECTORS in _pca_fit (which
    is what the similarity ranking consumes); the eigenvalue COUNT is a
    hypothesis test that belongs on the raw spectrum. `robust_cov` is kept
    in the signature so call sites are unchanged."""
    n, p = ret_scaled.shape
    ev = PCA().fit(ret_scaled).explained_variance_
    mp_edge = (1.0 + np.sqrt(p / n)) ** 2
    k = int(np.sum(ev > mp_edge))
    return int(min(max(2, k), cap, 15))


def download_hedge_data(target, candidates=None, session=None,
                        years=3, min_mktcap_usd_mm=5000, exclude=None,
                        force_include=None, fetch_mktcap_hist=True,
                        fetch_pit_members=True, cache=True):
    """[Fix 31] CELL A for Jupyter: universe discovery + the single Bloomberg
    download. Run once; feed the returned dict to find_best_hedge(..., data=...).
    [Fix 35] Stores current USD market caps so the cap floor can be retuned with
    no API calls (min_mktcap_usd_mm in USD MILLIONS; 5000 = $5bn).
    [Fix 36] `exclude` drops names before download. [Fix 40] `force_include`
    guarantees names are downloaded so they can be forced into every basket.
    [Fix 44] Also fetches HISTORICAL market-cap PATHS (approx-USD) so the size
    filter can be applied point-in-time (as-of each window) instead of with
    today's cap. Set fetch_mktcap_hist=False to skip that extra download.
    [Fix 46] `target` may be a single ticker, an equal-weight list, or a
    weighted dict {'AAPL US Equity': 0.6, '700 HK': 0.4} (a whole book). For a
    book the universe is the UNION of each leg's universe, every leg is removed
    from the candidate pool, and all legs are downloaded."""
    session = session or get_session()
    legs, pf_w, is_portfolio, label = _coerce_targets(target)      # [Fix 46]
    leg_shorts = {l.split(' ')[0] for l in legs}
    disc_meta = []                                                 # [Fix 49]

    if candidates is None:
        if is_portfolio:
            print(f"  Portfolio target {label}: discovering the union of "
                  f"{len(legs)} legs' universes...")
            uni, disc_meta = [], []
            for leg in legs:
                m = {}
                uni += auto_discover_universe(
                    session, leg, min_mktcap_usd_mm=min_mktcap_usd_mm, meta=m)
                disc_meta.append(m)
            seen, candidates = set(), []
            for c in uni:                       # dedupe; never a leg's own name
                s = c.split(' ')[0]
                if s in leg_shorts or s in seen:
                    continue
                candidates.append(c)
                seen.add(s)
        else:
            m = {}
            candidates = auto_discover_universe(
                session, legs[0], min_mktcap_usd_mm=min_mktcap_usd_mm, meta=m)
            disc_meta = [m]

    candidates, removed = _apply_exclusions(candidates, exclude)   # [Fix 36]
    if removed:
        print(f"  Manual exclusions: dropped {len(removed)} names before "
              f"download ({', '.join(t.split(' ')[0] for t in removed[:12])}"
              f"{', ...' if len(removed) > 12 else ''})")
    for f in (force_include or []):                               # [Fix 40]
        tk = _fmt_bbg_ticker(str(f))
        if tk is None:
            print(f"  ⚠ force_include entry '{f}' has no exchange code — "
                  f"cannot download it; pass e.g. '0050 TT'")
            continue
        if tk in removed:
            print(f"  ⚠ {tk} is in both exclude= and force_include= — "
                  f"force_include wins, keeping it")
        if tk not in candidates and tk not in legs:
            candidates.append(tk)

    seen, dedup = set(leg_shorts), []            # dedupe, exclude every leg
    for c in candidates:
        s = c.split(' ')[0]
        if s not in seen:
            dedup.append(c)
            seen.add(s)

    end_dt = datetime.now().strftime('%Y%m%d')
    start_dt = (datetime.now() - timedelta(days=int(years * 365))).strftime('%Y%m%d')
    # [Fix 49] point-in-time index membership — graceful: {} when the field
    # is unavailable for this index (then everything behaves exactly as before)
    pit_members, non_index_shorts = {}, set()
    if fetch_pit_members and disc_meta:
        for m in disc_meta:
            non_index_shorts |= m.get('non_index_shorts', set())
        pit_members = _pit_index_membership(
            session, [m.get('index_ticker') for m in disc_meta],
            start_dt, end_dt)
        if pit_members:
            ever = set().union(*pit_members.values())
            known = {c.split(' ')[0] for c in dedup} | leg_shorts
            extra = sorted(t for t in ever if t.split(' ')[0] not in known)
            extra, _ = _apply_exclusions(extra, exclude)
            if extra:
                print(f"  [Fix 49] adding {len(extra)} PAST index members "
                      f"missing from today's universe (delisted / dropped "
                      f"names — survivorship fix)")
                dedup += extra
            print(f"  ✓ point-in-time membership: {len(pit_members)} "
                  f"snapshots, {len(ever)} names ever members")
    elif fetch_pit_members:
        print("  (candidates were user-supplied — no index discovered, "
              "point-in-time membership skipped)")

    all_tickers = legs + dedup                                    # [Fix 46]
    print(f"Downloading {len(all_tickers)} tickers ({start_dt} → {end_dt})...")
    prices = bbg_hist(session, all_tickers, start_dt, end_dt)
    present_legs = [l for l in legs if l in prices.columns]
    if is_portfolio:
        if len(present_legs) < 2:
            raise ValueError(
                f"Only {len(present_legs)} of {len(legs)} portfolio legs "
                f"returned prices — check the tickers ({legs})")
        if len(present_legs) < len(legs):
            print(f"  ⚠ portfolio legs with no price data: "
                  f"{[l for l in legs if l not in present_legs]}")
    elif legs[0] not in prices.columns:
        raise ValueError(f"No price data for {legs[0]} — check the ticker spelling")
    print(f"  ✓ {prices.shape[0]} price rows × {prices.shape[1]} tickers")

    # [Fix 35] current USD market caps (fail-open)
    mktcap = {}
    try:
        liq = bbg_ref_multi(session, list(prices.columns),
                            ['CRNCY_ADJ_MKT_CAP'],
                            overrides={'EQY_FUND_CRNCY': 'USD'})
        for tk, d_ in liq.items():
            try:
                mktcap[tk] = float(d_.get('CRNCY_ADJ_MKT_CAP'))
            except (TypeError, ValueError):
                continue
        print(f"  ✓ USD market caps stored for {len(mktcap)} tickers")
    except Exception as e:
        print(f"  ⚠ market-cap download skipped ({e}) — the cap filter will "
              f"be unavailable in find_best_hedge")

    # [Fix 44] historical market-cap PATHS (local ccy from Bloomberg), rescaled
    # to approx-USD by matching each name's latest point to its current USD cap.
    mktcap_hist = pd.DataFrame()
    if fetch_mktcap_hist:
        try:
            hist_local = bbg_hist_field(session, list(prices.columns),
                                        'CUR_MKT_CAP', start_dt, end_dt)
            scaled = {}
            for c in hist_local.columns:
                s = hist_local[c].dropna()
                if s.empty:
                    continue
                cur_usd = mktcap.get(c)
                last_local = float(s.iloc[-1])
                if cur_usd and last_local:
                    scaled[c] = hist_local[c] * (cur_usd / last_local)   # ≈ USD
                else:
                    scaled[c] = hist_local[c]        # fail-open: local ccy path
            mktcap_hist = pd.DataFrame(scaled)
            print(f"  ✓ point-in-time market-cap paths for "
                  f"{mktcap_hist.shape[1]} tickers")
        except Exception as e:
            print(f"  ⚠ historical market-cap download skipped ({e}) — the size "
                  f"filter will fall back to today's caps (look-ahead)")

    out = {'target': target, 'target_label': label,
           'portfolio': {'legs': legs, 'weights': pf_w,
                         'is_portfolio': is_portfolio, 'label': label},
           'prices': prices, 'mktcap_usd_mm': mktcap,
           'mktcap_hist': mktcap_hist,
           'pit_members': pit_members,                        # [Fix 49]
           'non_index_shorts': sorted(non_index_shorts),
           'downloaded_at': datetime.now(),                   # [Fix 88]
           'fingerprint': {'legs': sorted(legs), 'years': years,  # [Fix 90]
                           'min_mktcap_usd_mm': min_mktcap_usd_mm,
                           'exclude': sorted(map(str, exclude or [])),
                           'force_include': sorted(map(str,
                                                       force_include or []))}}
    if cache:                                                 # [Fix 88]
        try:
            save_hedge_data(out)
        except Exception as e:
            print(f"  ⚠ could not cache this download ({e}) — it is still "
                  f"returned in memory")
    return out


def find_best_hedge(
    target_ticker,
    candidates=None,
    data=None,                    # [Fix 31] dict from download_hedge_data()
    session=None,
    lookback_options=(126, 252),
    pc_options='auto',            # 'auto' or a list of ints, e.g. [3, 5, 8]
    basket_size_options=(3, 5, 7),
    cv_test_window=42,
    cv_n_folds=5,
    outer_folds=3,                # [Fix 20] nested outer validation windows
    outer_window=None,            # defaults to cv_test_window
    async_adjust='auto',          # [Fix 5] 'auto' / True / False
    restrict_to_target_tz='auto',  # [Fix 26/33] 'auto' / True / False
    tz_tolerance_hours=3.0,       # [Fix 32/33] closes within this many hours
                                  # of the target's count as synchronous
    min_mktcap_usd_mm=5000,       # [Fix 25/35] market-cap floor, USD MILLIONS
                                  # (5000 = $5bn); now also works with data=
    exclude=None,                 # [Fix 36] securities to filter out, e.g.
                                  # ['002230', '3156 JP Equity']
    force_include=None,           # [Fix 40] securities that MUST be held in
                                  # every basket, in ALL outputs — e.g.
                                  # ['0050 TT'] or ['2317 TT Equity']; they
                                  # bypass the cap floor / tz restriction /
                                  # exclude, and occupy basket slots in every
                                  # CV fold, nested window, and final refit
    recency_weight=0.0,           # [Fix 108] 0 = every CV fold counts equally
                                  # (the pre-v104 behaviour). >0 tilts the
                                  # selection toward RECENT folds: a fold f
                                  # folds back in time gets weight
                                  # exp(-recency_weight * f). 0.35 ≈ the
                                  # newest fold counting ~2x the one two
                                  # folds older. See the [Fix 108] note in
                                  # the selection block before using it.
    gross_penalty=0.02,           # [Fix 38] composite penalty per 1.0x gross
    max_gross=None,               # [Fix 38] e.g. 3.0 → discard configs whose
                                  # avg CV gross exceeds $300 per $100 target
    winsorize_pct=0.01,           # [Fix 42] clip returns at 1%/99% before PCA
                                  # (candidate selection only; 0 = old numbers)
    robust_cov='none',            # [Fix 42] 'ledoit' → shrinkage-covariance PCA
    pit_size=True,                # [Fix 44] apply the size floor point-in-time
                                  # (as-of each window) using historical caps
    pit_members=True,             # [Fix 49] restrict index-sourced candidates
                                  # to names that were ACTUALLY members as-of
                                  # each window (needs pit data; off if absent)
    size_band=None,               # [Fix 45] (lo, hi) → keep candidates whose cap
                                  # is within lo–hi × the target's cap; overrides
                                  # the flat floor when set, e.g. (0.3, 3.0)
    size_override=None,           # [Fix 45] forward-looking cap views, e.g.
                                  # {'XYZ': 'large'} or {'XYZ': 8000} (USD mm)
    mask_stale_days=7,            # [Fix 56] mask prices that stay flat for
                                  # more than N consecutive sessions
                                  # (suspensions); None disables the guard
    notional_mm=None,             # [Fix 72] gross book notional in USD mm —
                                  # dollarizes the trade-ticket amounts
    rolling_oos=True,             # [Fix 71] walk-forward rolling OOS curve
                                  # (section [1c] + chart)
    rolling_refit_every=21,       # [Fix 71] rows between weight refits
    rolling_max_windows=12,       # [Fix 71] refit segments to walk back
    boot_n=2000,                  # [Fix 70] bootstrap draws for the CI on
                                  # pooled OOS R2 / TE; 0 disables
    hedge_horizon_days=None,      # [Fix 91/100] how long you intend to HOLD
                                  # this hedge, in trading days. Scales the
                                  # reported risk to that horizon, reports
                                  # what holds of that length actually
                                  # delivered, and compares the horizon with
                                  # the measured rebalance cadence.
    horizon_sets_windows=False,   # [Fix 100] re-cut the OOS windows to the
                                  # horizon. OFF by default — measured, it
                                  # does not improve realized horizon
                                  # performance and costs evidence. See the
                                  # note at the top of STEP 1.
    sample_window=None,           # [Fix 76] restrict the analysis window:
                                  # '3m' / '90d' / '26w' / '1y' (last N up to
                                  # the end of the data) or an absolute
                                  # (start, end) tuple, e.g. ('2026-07-13',
                                  # '2026-07-28') / ('13-Jul-2026', None).
                                  # Validation auto-adapts — see [Fix 77].
    report='compact',             # [Fix 79] 'compact' (default) or 'full'
                                  # (adds long-form explanations and per-
                                  # config weight tables)
    show_plots=True,
    _rehedge_of=None,             # [Fix 115] set by quick_rehedge so the
                                  # report can say what it inherited and what
                                  # it did not re-validate
):
    outer_window = outer_window or cv_test_window
    _kv_reset()             # [Fix 114] never inherit a failed run's rows
    transcript_reset()      # [Fix 115] this run's text copy starts empty
    # ── [Fix 100] WHY THE HORIZON DOES NOT RE-CUT THE WINDOWS ───────────────
    # [Fix 91] originally set the OOS windows to the holding period, on the
    # intuition that you should be validated on the period you will live.
    # That intuition does not survive measurement. A hedge ratio is a
    # property of the covariance between target and basket; estimated on
    # daily returns it targets the SAME quantity as one estimated on 63-day
    # returns, but with ~63x more observations and far less noise. Changing
    # the SCORING window only changes which config wins — and longer windows
    # mean FEWER of them, so the choice gets noisier, not sharper.
    #
    # Tested directly: train on 600 rows, then score the SAME weights on a
    # held-out tail in non-overlapping 63-row blocks, 8 paired trials per
    # regime.
    #     iid residuals   default 16.39%  vs  horizon-matched 16.43%
    #                     (t=+0.44, horizon better in 4/8)
    #     AR(1) 0.6       default 20.57%  vs  horizon-matched 20.58%
    #                     (t=+0.09, horizon better in 5/8)
    # No detectable benefit either way. So the horizon is used where it
    # genuinely carries information — scaling risk to the holding period
    # ([Fix 92]), reporting what holds of that length actually delivered
    # ([Fix 98]), and checking the horizon against the measured rebalance
    # cadence ([Fix 78/84]) — and NOT to shrink the evidence base. Pass
    # horizon_sets_windows=True to restore the old behaviour.
    #
    # ── [Fix 112] THE STRONGER REASON, AND THE ONE THAT SETTLES IT ─────────
    # [Fix 100] only tested one way of using the horizon (re-cutting the
    # windows). That left the obvious objection unanswered: hedging for a
    # MONTH and hedging for HALF A YEAR are different problems, so surely the
    # horizon belongs in the selection somehow?
    #
    # The mechanism by which it could is real. A hedge ratio is horizon-free
    # only if the residual is iid. When residuals are serially dependent,
    #     VR_h = Var(sum of h residuals) / (h x Var(residual))
    # and h-day tracking error is daily-TE x sqrt(VR_h). Two configs with the
    # SAME daily TE can therefore have very different 21-day or 126-day TE.
    # Crucially this needs NO extra data — VR is estimable from the same
    # daily returns — so it dodges the objection that killed [Fix 100].
    #
    # It was implemented and measured. Three selection rules, scored on
    # REALISED non-overlapping h-day tracking error on held-out data, 25
    # paired seeds x 40 candidate baskets, h = 5/21/63/126, two worlds (one
    # built specifically so residual persistence differs across configs):
    #     rule                                    h=21      h=63     h=126
    #     A  minimise daily TE      (current)   baseline  baseline  baseline
    #     B  minimise TE x sqrt(VR) (proposed)   -0.50pp   -0.54pp   -0.93pp
    #     C  minimise TE on h-blocks ([Fix 100]) -0.84pp   -0.82pp   -2.55pp
    # Both horizon-aware rules are WORSE, at every horizon, in both worlds,
    # and B wins in only 4-11 of 25 seeds. The reason is visible in the
    # estimates themselves:
    #     corr(VR_in-sample, VR_out-of-sample) = -0.32 .. +0.18
    #     corr(VR_out-of-sample, realised)     = +0.54 .. +0.86
    # The quantity MATTERS enormously and is NOT forecastable. Selecting on
    # it does not merely fail to help — it maximises an estimation error, so
    # it actively hurts, and hurts MORE at longer horizons because VR gets
    # harder to estimate as h grows.
    #
    # That is the answer to "shouldn't a 1-month and a 6-month hedge differ?"
    # They differ in the RISK you carry and in how often you must refit —
    # both of which are reported. They do not differ in which covariance you
    # should estimate, and the longer the horizon the LESS the data can say
    # about anything horizon-specific. So the horizon shapes the report and
    # the monitoring plan, and stays out of the fit.
    if hedge_horizon_days:
        _h = int(hedge_horizon_days)
        if _h < 1:
            raise ValueError('[Fix 91] hedge_horizon_days must be >= 1')
        if horizon_sets_windows:
            cv_test_window = outer_window = int(np.clip(_h, 10, 63))
            _P(f"  [Fix 100] horizon_sets_windows=True → OOS windows set "
                  f"to {cv_test_window} rows (measured as no better than the "
                  f"default; kept as an option)")
    # [Fix 46] single-name vs portfolio ---------------------------------------
    pf_legs, pf_w, is_portfolio, pf_label = _coerce_targets(target_ticker)
    _uniq = lambda seq: list(dict.fromkeys(seq))
    if is_portfolio:
        target_short = pf_label
        target_keepset = set(pf_legs)
        target_ticker_single = None
    else:
        target_ticker = pf_legs[0]            # canonical full ticker string
        target_ticker_single = target_ticker
        target_short = target_ticker.split(' ')[0]
        target_keepset = {target_ticker}

    # ── STEP 1+2: GET PRICES (pre-downloaded data or fresh download) ─────────
    if data is not None:
        # [Fix 66] The old [Fix 46] check demanded the book EQUAL the book
        # the data was downloaded for — which made quick_rehedge crash the
        # moment a churning book added or dropped a leg, the exact use case
        # it was built for. The real requirement is only that every leg has
        # downloaded prices: legs are excluded from the candidate pool
        # downstream, and ex-legs simply become candidates (legitimate — and
        # often the best hedges for a rebalance book). Universe DISCOVERY
        # ran for the original book, so after a MATERIAL change (new sector
        # or market) re-run CELL A; for leg churn this is exactly right.
        _have = set(data['prices'].columns)
        _missing = [l for l in pf_legs if l not in _have]
        if _missing:
            raise ValueError(
                f"[Fix 66] legs {_missing} are not in the downloaded data — "
                f"re-run download_hedge_data() for this book, or pass them "
                f"via force_include= there so they are fetched")
        if set(_coerce_targets(data.get('target'))[0]) != set(pf_legs):
            _P(f"  [Fix 66] note: data was downloaded for "
                  f"{data.get('target_label', data.get('target'))} — reusing "
                  f"its universe for {target_short}. Fine for leg churn on "
                  f"the same book; re-run CELL A after a material change.")
        prices = data['prices'].copy()
        _P(f"Using data already downloaded: {prices.shape[0]} days of "
              f"prices for {prices.shape[1]} names (nothing fetched)")
    else:
        data = download_hedge_data(target_ticker, candidates=candidates,
                                   session=session or get_session(),
                                   min_mktcap_usd_mm=min_mktcap_usd_mm,
                                   exclude=exclude,
                                   force_include=force_include)  # [Fix 40]
        prices = data['prices'].copy()

    # [Fix 40] Resolve the must-have securities against the downloaded
    # columns BEFORE any filter runs, so every filter below can exempt them.
    forced, fi_unres = _resolve_force_include(force_include, prices.columns,
                                              target_ticker)
    if fi_unres:
        _P(f"  ⚠ force_include: {fi_unres} not in the downloaded data — "
              f"re-run download_hedge_data(..., force_include=...) to fetch "
              f"them (they will be ignored this run)")
    if forced:
        _P(f"  Force-include active: {', '.join(forced)} will be held in "
              f"every basket (CV folds, nested windows, and final recipe)")

    # [Fix 36] Manual exclusions on COLUMNS — works on pre-downloaded data
    kept_cols, removed_cols = _apply_exclusions(
        [c for c in prices.columns if c not in target_keepset], exclude)
    readd = [c for c in removed_cols if c in forced]       # [Fix 40]
    if readd:
        _P(f"  ⚠ {', '.join(readd)} in both exclude= and force_include= — "
              f"force_include wins, keeping")
        kept_cols += readd
        removed_cols = [c for c in removed_cols if c not in forced]
    if removed_cols:
        _P(f"  Manual exclusions: dropped {len(removed_cols)} names "
              f"({', '.join(t.split(' ')[0] for t in removed_cols[:12])}"
              f"{', ...' if len(removed_cols) > 12 else ''})")
    if removed_cols or readd:
        prices = prices[_uniq(list(target_keepset) + kept_cols)]

    # [Fix 35] Market-cap floor on COLUMNS — the old code only consumed
    # min_mktcap_usd_mm on the fresh-download path, so with data= the
    # parameter was silently ignored no matter what value was passed.
    caps = data.get('mktcap_usd_mm') or {}
    caps_hist = (data.get('mktcap_hist')
                 if isinstance(data.get('mktcap_hist'), pd.DataFrame)
                 else pd.DataFrame())
    pit_active = bool(pit_size) and not caps_hist.empty        # [Fix 44]
    pm_snaps = (data.get('pit_members') or {}) if pit_members else {}  # [Fix 49]
    pm_dates = sorted(pm_snaps.keys())
    pm_union = set().union(*pm_snaps.values()) if pm_snaps else set()
    pm_exempt = set(data.get('non_index_shorts') or [])
    pm_active = bool(pm_snaps)
    if pm_active:
        _P(f"  [Fix 49] point-in-time index membership active — "
              f"{len(pm_dates)} snapshots, {len(pm_union)} names ever members")
    if min_mktcap_usd_mm and not pit_active:
        # static (current-cap) floor on COLUMNS — used only when point-in-time
        # caps are unavailable or pit_size=False. NOTE: this legacy path carries
        # the "big now / small then" look-ahead that [Fix 44] removes.
        if not caps:
            _P("  ⚠ market-cap floor requested but data has no "
                  "'mktcap_usd_mm' — re-run download_hedge_data() to enable it")
        else:
            kept, dropped, unknown = list(target_keepset), 0, 0
            for c in prices.columns:
                if c in target_keepset:
                    continue
                if c.split(' ')[0] in ETF_SHORT_SET or c in forced:
                    kept.append(c)
                    continue
                mc = caps.get(c)
                if mc is None:
                    unknown += 1
                    kept.append(c)       # unknown → keep (fail-open)
                    continue
                if mc >= min_mktcap_usd_mm:
                    kept.append(c)
                else:
                    dropped += 1
            if dropped:
                _P(f"  Market-cap floor {_fmt_mktcap(min_mktcap_usd_mm)}: "
                      f"dropped {dropped} names"
                      + (f" ({unknown} unknown caps kept)" if unknown else ""))
                prices = prices[_uniq(kept)]
    elif pit_active:
        _P(f"  the {_fmt_mktcap(min_mktcap_usd_mm)} size floor is applied "
              f"using what each name was worth AT THE TIME, not what it is "
              f"worth today — so a name that grew into the filter cannot "
              f"flatter the backtest")

    prices = prices.ffill()
    if mask_stale_days:                                     # [Fix 56]
        prices = _mask_stale(prices, mask_stale_days)
    if sample_window is not None:                           # [Fix 76]
        _sw0, _sw1 = _parse_sample_window(sample_window, prices.index)
        _avail = (prices.index[0], prices.index[-1])
        prices = prices.loc[_sw0:_sw1]
        if len(prices) < 10:
            raise ValueError(
                f"[Fix 76] sample window {_sw0:%Y-%m-%d} → {_sw1:%Y-%m-%d} "
                f"leaves {len(prices)} sessions (data spans "
                f"{_avail[0]:%Y-%m-%d} → {_avail[1]:%Y-%m-%d}) — need ≥ 10")
        _P(f"  [Fix 76] sample window: {prices.index[0]:%Y-%m-%d} → "
              f"{prices.index[-1]:%Y-%m-%d} ({len(prices)} sessions)")
    # completeness threshold applies WITHIN the analysis window
    prices = prices.dropna(axis=1, thresh=int(len(prices) * 0.85))
    if is_portfolio:                             # [Fix 46]
        _pl = [c for c in pf_legs if c in prices.columns]
        if len(_pl) < 2:
            raise ValueError(
                f"Fewer than 2 portfolio legs have ≥85% price history "
                f"({pf_legs}) — cannot build the book")
    elif target_ticker not in prices.columns:    # [Fix 21]
        raise ValueError(
            f"{target_ticker} has price data for less than 85% of the window — "
            "history too short for this configuration")
    # [Fix 40] a forced name with <85% history cannot be modelled — say so
    # loudly instead of silently forcing NaNs into every window
    lost = [c for c in forced if c not in prices.columns]
    if lost:
        _P(f"  ⚠ force_include: {', '.join(lost)} dropped — less than 85% "
              f"price history in the window; cannot be forced into baskets")
        forced = [c for c in forced if c in prices.columns]

    # [Fix 26/31/32/33] Timezone restriction by CLOSING-TIME GAP, applied to
    # COLUMNS so one global download supports both restricted and
    # unrestricted runs. Two listings count as synchronous when their cash
    # closes are within tz_tolerance_hours of each other (circular clock
    # distance). The old AMER/EMEA/APAC buckets meant restrict=True could
    # never separate e.g. Taipei from Mumbai (both 'APAC', closes 4.5h
    # apart), and unmapped exchanges ('AT' etc.) fell into 'OTHER'.
    # Note: with the default 'auto', restriction already fires whenever the
    # target's own close-time group has >= 30 candidates — so passing True
    # explicitly only changes anything when that group is small.
    if is_portfolio:                             # [Fix 46] book close = legs' if synchronous
        _lc = [_close_utc_of(c) for c in pf_legs]
        if all(x is not None for x in _lc) and (
                max((_close_gap_hours(a, b) for i, a in enumerate(_lc)
                     for b in _lc[i + 1:]), default=0.0) <= tz_tolerance_hours):
            tgt_close = float(np.mean(_lc))
        else:
            tgt_close = None                     # spans tz / unknown → conservative
    else:
        tgt_close = _close_utc_of(target_ticker)
    others = [c for c in prices.columns if c not in target_keepset]
    if tgt_close is None:
        _P(f"  ⚠ unknown exchange close for {target_ticker} — timezone "
              f"restriction unavailable; universe treated as cross-timezone")
        same_tz, do_restrict = [], False
    else:
        same_tz = [c for c in others
                   if _close_utc_of(c) is not None
                   and _close_gap_hours(_close_utc_of(c), tgt_close)
                   <= tz_tolerance_hours]
        do_restrict = (restrict_to_target_tz is True or
                       (restrict_to_target_tz == 'auto' and len(same_tz) >= 30))
    if do_restrict and len(same_tz) < len(others):
        # [Fix 40] forced names survive the restriction; if a forced name is
        # cross-timezone the async detection below flips to 2-day returns
        keep_tz = same_tz + [c for c in list(forced) + pf_legs
                             if c not in same_tz]   # [Fix 46] legs kept
        _P(f"  Restricting universe to closes within "
              f"{tz_tolerance_hours:.1f}h of the target's "
              f"({len(others)} → {len(keep_tz)} candidates"
              + (f"; {len(keep_tz) - len(same_tz)} kept by force_include"
                 if len(keep_tz) > len(same_tz) else "")
              + "; set restrict_to_target_tz=False to keep global names)")
        prices = prices[_uniq(list(target_keepset) + keep_tz)]
    elif restrict_to_target_tz is True and len(same_tz) == len(others):
        _P(f"  All candidates already close within "
              f"{tz_tolerance_hours:.1f}h of the target — nothing to drop")

    # ── Cross-timezone detection → 2-day overlapping returns [Fix 5/34] ─────
    # Trigger when any retained pair of closes is further apart than the
    # tolerance, or when any exchange close is unknown (conservative). The
    # old check compared bucket labels and EXCLUDED unknown ('OTHER')
    # exchanges, so a genuinely async pair could pass undetected.
    closes = [_close_utc_of(c) for c in prices.columns]
    known = sorted({c for c in closes if c is not None})
    has_unknown = any(c is None for c in closes)
    max_gap = max((_close_gap_hours(a, b)
                   for i, a in enumerate(known) for b in known[i + 1:]),
                  default=0.0)
    if async_adjust == 'auto':
        two_day = has_unknown or max_gap > tz_tolerance_hours
    else:
        two_day = bool(async_adjust)
    ann_sqrt = np.sqrt(252 / (2 if two_day else 1))
    if two_day:
        reason = ("unknown exchange close present"
                  if has_unknown and max_gap <= tz_tolerance_hours
                  else f"max close gap {max_gap:.1f}h > "
                       f"{tz_tolerance_hours:.1f}h tolerance")
        _P(f"  ⚠ Universe treated as cross-timezone ({reason}) → "
              f"switching to 2-day overlapping returns to correct "
              f"async-trading bias")

    rets = np.log(prices / prices.shift(1))
    if is_portfolio:                             # [Fix 46] synthetic book return
        legs_present = [c for c in pf_legs if c in rets.columns]
        if len(legs_present) < 2:
            raise ValueError(f"Only {len(legs_present)} portfolio legs survived "
                             f"filtering — cannot build the book ({pf_legs})")
        if len(legs_present) < len(pf_legs):
            miss = [c for c in pf_legs if c not in legs_present]
            _P(f"  ⚠ portfolio legs dropped by filtering: {miss} — weights "
                  f"renormalized over the {len(legs_present)} remaining legs")
        _wser = pd.Series([pf_w[c] for c in legs_present], index=legs_present)
        _wser = _wser / (_wser.abs().sum() or 1.0)
        rets[PORTFOLIO_COL] = rets[legs_present].mul(_wser, axis=1).sum(
            axis=1, skipna=False)
    if two_day:
        rets = rets.rolling(2).sum()
    rets = rets.iloc[2 if two_day else 1:]      # drop the leading NaN rows
    if is_portfolio:                             # [Fix 46] from here, hedge the book
        target_ticker = PORTFOLIO_COL

    # [Fix 39] Overlapping 2-day returns make ADJACENT rows share one daily
    # return, so every train/test boundary below gets a 1-row purge gap —
    # otherwise the last training row leaks half of the first test row and
    # every OOS statistic is slightly optimistic.
    purge = 1 if two_day else 0

    # ── [Fix 77/85] ADAPTIVE VALIDATION TIERS ───────────────────────────────
    # The protocol degrades explicitly rather than hard-failing on short
    # windows; _validation_tier holds the rules ([Fix 85]) so preview_window
    # and the GUI cannot promise a protocol other than the one that runs.
    rows = len(rets)
    _tier = _validation_tier(rows, lookback_options, cv_test_window,
                             cv_n_folds, outer_window, outer_folds, purge,
                             basket_size_options, rolling_refit_every)
    val_mode = _tier['mode']
    hold_rows = _tier['hold_rows']
    lookback_options = _tier['lookback_options']
    cv_test_window = _tier['cv_test_window']
    cv_n_folds = _tier['cv_n_folds']
    outer_window = _tier['outer_window']
    outer_folds = _tier['outer_folds']
    min_train_rows = _tier['min_train_rows']
    rolling_refit_every = _tier['rolling_refit_every']
    basket_size_options = _tier['basket_size_options']
    for _note in _tier['notes']:
        _P(_note)

    universe = [c for c in prices.columns
                if c not in set(pf_legs) and c != target_ticker]   # [Fix 46]
    forced = [c for c in forced if c in universe]           # [Fix 40]

    # [Fix 44/45] point-in-time eligibility: size floor / size band / overrides
    _sz_over = {}
    for k, v in (size_override or {}).items():
        base = str(k).upper().replace(' EQUITY', '').strip()
        _sz_over[_fmt_bbg_ticker(str(k)) or base] = v
        _sz_over[base] = v
        _sz_over[base.split(' ')[0]] = v

    def _cap_asof(c, asof):
        for key in (c, c.split(' ')[0]):
            if key in _sz_over:
                v = _sz_over[key]
                if isinstance(v, str):
                    return np.inf if v.lower().startswith('l') else 0.0
                return float(v)
        if pit_active and c in caps_hist.columns:
            s = caps_hist[c]
            s = s[s.index <= asof].dropna()
            if len(s):
                return float(s.iloc[-1])
        return caps.get(c)                        # fallback: current USD cap

    def _eligible(asof):
        if not min_mktcap_usd_mm and not size_band and not pm_active:
            return universe
        tgt_cap = None
        if size_band:
            if is_portfolio:
                # [Fix 74] compare candidates to the GROSS-WEIGHTED AVERAGE
                # leg cap, not the SUM of leg caps. For a book of three mega
                # caps the sum (~$8tn for AAPL+NVDA+MSFT) made lo x cap
                # exceed every listed stock, so the band silently kept only
                # ETFs and forced names.
                _num = _den = 0.0
                for _lc in pf_legs:
                    _v = _cap_asof(_lc, asof)
                    _w = abs(pf_w.get(_lc, 0.0))
                    if _v is not None and np.isfinite(_v) and _w > 0:
                        _num += _w * float(_v)
                        _den += _w
                tgt_cap = (_num / _den) if _den > 0 else None
            else:
                tgt_cap = _cap_asof(target_ticker_single, asof)
        out = []
        for c in universe:
            if c.split(' ')[0] in ETF_SHORT_SET or c in forced:
                out.append(c)
                continue
            if pm_active and c in pm_union \
                    and c.split(' ')[0] not in pm_exempt:      # [Fix 49]
                snap = pm_snaps[pm_dates[0]]
                for dtt in reversed(pm_dates):
                    if dtt <= asof:
                        snap = pm_snaps[dtt]
                        break
                if c not in snap:
                    continue          # not an index member as-of this window
            mc = _cap_asof(c, asof)
            if mc is None:
                out.append(c)                     # fail-open
                continue
            if size_band and tgt_cap:
                lo, hi = size_band
                if lo * tgt_cap <= mc <= hi * tgt_cap:
                    out.append(c)
            elif mc >= (min_mktcap_usd_mm or 0):
                out.append(c)
        return out

    def _inject_synth_loadings(lds):
        """[Fix 46] Give the synthetic book a loading row = weighted sum of its
        legs' loadings (portfolio factor exposure is linear in the legs)."""
        if not is_portfolio:
            return lds
        legs_in = [c for c in pf_legs if c in lds.index]
        if not legs_in:
            raise RuntimeError("no portfolio legs present in this window")
        w = np.array([pf_w[c] for c in legs_in], dtype=float)
        w = w / (np.abs(w).sum() or 1.0)
        lds.loc[PORTFOLIO_COL] = (lds.loc[legs_in].values * w[:, None]).sum(axis=0)
        return lds
    _P(f"  ✓ {len(rets)} days of returns across {len(universe) + 1} names")

    # ── Reusable building blocks ─────────────────────────────────────────────
    def _forced_basket(ranked, ok, bsize):
        """[Fix 40] Assemble a basket that always contains the forced names
        (those usable in this window), topped up with the best-ranked other
        candidates. Forced names OCCUPY SLOTS inside the same basket size,
        so CV metrics measure the basket that would actually be traded; if
        more names are forced than basket_size, the basket grows to hold
        them all."""
        f_ok = [c for c in forced if c in ok]
        rest = [t for t in ranked if t not in f_ok]
        return f_ok + rest[:max(0, bsize - len(f_ok))]

    def _select_and_fit(train_rets, pc_spec, bsize):
        """Select a basket via PCA ranking on train_rets and fit RidgeCV."""
        ok = train_rets.columns[(train_rets.notna().all()) & (train_rets.std() > 1e-6)]
        if target_ticker not in ok:
            raise RuntimeError(f"{target_short} has incomplete data or zero "
                               f"variance in this window — cannot fit.")
        tr = train_rets[ok]
        u = [c for c in _eligible(tr.index[-1]) if c in ok]     # [Fix 44/45] PIT
        pca_cols = [c for c in tr.columns if c != PORTFOLIO_COL]  # [Fix 46]
        tr_w = _winsorize_returns(tr[pca_cols], winsorize_pct)   # [Fix 42]
        sc = StandardScaler()
        tr_sc = pd.DataFrame(sc.fit_transform(tr_w), index=tr.index,
                             columns=pca_cols)
        cap = len(pca_cols) - 1
        n_pc = (_n_pc_auto(tr_sc, cap, robust_cov) if pc_spec == 'auto'
                else int(min(int(pc_spec), cap)))
        pca = _pca_fit(tr_sc, n_pc, robust_cov)                  # [Fix 42]
        # [Fix 23] scale loadings by sqrt(eigenvalue): similarity in factor-
        # exposure space, where PC1 counts far more than higher (noisier) PCs
        lds = pd.DataFrame(pca.components_.T * np.sqrt(pca.explained_variance_),
                           index=pca_cols,
                           columns=[f'PC{i+1}' for i in range(n_pc)])
        lds = _inject_synth_loadings(lds)                        # [Fix 46]
        rank = _rank_scores(lds, tr, target_ticker, u)
        basket = _forced_basket(rank.index.tolist(), ok, bsize)   # [Fix 40]
        if len(basket) < 2:                      # [Fix 21]
            raise RuntimeError("Fewer than 2 valid candidates in this window "
                               "— cannot build a basket.")
        reg, r2, resid = _fit_basket(tr, target_ticker, basket)
        return tr, rank, basket, reg, r2, resid, n_pc, pca, lds

    def _eval_on_window(reg, basket, test_rets):
        """Score a fitted basket on an unseen window → (r2, vol_red) or NaNs."""
        needed = basket + [target_ticker]
        if not all(c in test_rets.columns and test_rets[c].notna().all()
                   for c in needed):
            return np.nan, np.nan
        r2 = reg.score(test_rets[basket].values, test_rets[target_ticker].values)
        resid = (test_rets[target_ticker].values -
                 reg.predict(test_rets[basket].values))
        tgt_std = np.std(test_rets[target_ticker].values) * ann_sqrt
        vr = 1 - (np.std(resid) * ann_sqrt) / tgt_std if tgt_std > 1e-6 else 0.0
        return r2, vr

    pc_grid = ['auto'] if pc_options == 'auto' else list(pc_options)   # [Fix 4]

    def _walk_forward_cv(rets_sub, verbose=False):
        """Run the full walk-forward CV grid on rets_sub → summary DataFrame
        (or None if every fold failed)."""
        results = []
        for lookback in lookback_options:
            if lookback + cv_test_window * cv_n_folds > len(rets_sub):
                if verbose:
                    _P(f"  Skipping lookback={lookback} (insufficient data)")
                continue
            for fold in range(cv_n_folds):
                te = len(rets_sub) - fold * cv_test_window
                ts = te - cv_test_window
                tr_e = ts - purge                     # [Fix 39] purge gap
                tr_s = max(0, tr_e - lookback)
                if tr_s >= tr_e or ts >= te:
                    continue
                ret_tr_all = rets_sub.iloc[tr_s:tr_e]
                ret_te_all = rets_sub.iloc[ts:te]

                # [Fix 6] complete data + non-zero variance only, no fillna
                ok = ret_tr_all.columns[(ret_tr_all.notna().all()) &
                                        (ret_tr_all.std() > 1e-6)]
                ok = [c for c in ok if ret_te_all[c].notna().all()]
                if target_ticker not in ok:
                    if verbose:
                        _P(f"  ⚠ skipping fold {fold} (lookback {lookback}): "
                              f"{target_short} incomplete or zero variance")
                    continue
                ret_tr = ret_tr_all[ok]
                ret_te = ret_te_all[ok]
                if len(ret_tr) < min_train_rows or len(ret_te) < 10:
                    continue                              # [Fix 80]
                cur_uni = [c for c in _eligible(ret_tr.index[-1]) if c in ok]  # [Fix 44/45]

                pca_cols = [c for c in ret_tr.columns if c != PORTFOLIO_COL]   # [Fix 46]
                ret_tr_w = _winsorize_returns(ret_tr[pca_cols], winsorize_pct)  # [Fix 42]
                scaler = StandardScaler()
                ret_tr_sc = pd.DataFrame(scaler.fit_transform(ret_tr_w),
                                         index=ret_tr.index, columns=pca_cols)
                cap = len(pca_cols) - 1
                for pc_spec in pc_grid:
                    n_pc = (_n_pc_auto(ret_tr_sc, cap, robust_cov)
                            if pc_spec == 'auto' else int(min(pc_spec, cap)))
                    pca = _pca_fit(ret_tr_sc, n_pc, robust_cov)               # [Fix 42]
                    # [Fix 23] eigenvalue-weighted loadings (see _select_and_fit)
                    loadings = pd.DataFrame(
                        pca.components_.T * np.sqrt(pca.explained_variance_),
                        index=pca_cols,
                        columns=[f'PC{i+1}' for i in range(n_pc)])
                    loadings = _inject_synth_loadings(loadings)              # [Fix 46]
                    ranked = _rank_scores(loadings, ret_tr, target_ticker,
                                          cur_uni).index.tolist()
                    for bsize in basket_size_options:
                        basket = _forced_basket(ranked, ok, bsize)  # [Fix 40]
                        if len(basket) < 2:
                            continue
                        reg, r2_tr, _ = _fit_basket(ret_tr, target_ticker, basket)
                        n, p = len(ret_tr), len(basket)
                        adj_r2 = 1 - (1 - r2_tr) * (n - 1) / max(1, n - p - 1)
                        r2_te, vol_red = _eval_on_window(reg, basket, ret_te)
                        if np.isnan(r2_te):
                            continue
                        results.append({
                            'lookback': lookback, 'pc_spec': str(pc_spec),
                            'n_pcs': n_pc, 'basket_size': bsize, 'fold': fold,
                            'r2_train': r2_tr, 'adj_r2_train': adj_r2,
                            'r2_test': r2_te, 'vol_red_test': vol_red,
                            'gross': float(np.abs(reg.coef_).sum()),  # [Fix 38]
                        })
        if not results:
            return None, None
        cv = pd.DataFrame(results)
        summary = cv.groupby(['lookback', 'pc_spec', 'basket_size']).agg(
            avg_r2_oos=('r2_test', 'mean'), std_r2_oos=('r2_test', 'std'),
            avg_vol_red=('vol_red_test', 'mean'),
            avg_adj_r2=('adj_r2_train', 'mean'),
            avg_gross=('gross', 'mean'),
            n_folds=('fold', 'count')).reset_index()
        # Sanitize NaN-able metrics so idxmax() always finds something
        summary['std_r2_oos'] = summary['std_r2_oos'].fillna(0)
        summary['avg_r2_oos'] = summary['avg_r2_oos'].fillna(-1)
        summary['avg_vol_red'] = summary['avg_vol_red'].fillna(0)
        summary['avg_gross'] = summary['avg_gross'].fillna(0)
        # [Fix 28] lower-confidence-bound composite: penalize dispersion
        # instead of rewarding it. The old 1/std "stability" bonus let a
        # reliably-useless config outrank a noisy-but-useful one.
        # [Fix 38] gross-exposure penalty: ridge on near-collinear ETFs
        # happily builds $500+-gross long/short recipes whose extra R² is a
        # multicollinearity artifact (e.g. short QQQ / long XLK offsets).
        # Penalize by average gross, in units of the $100 target notional
        # (avg_gross=2.6 means $260 gross), and optionally hard-cap it.
        # ── [Fix 108] RECENCY — "if it deteriorates, why is it the best?" ──
        # The composite averages OOS R² over ALL folds with equal weight, so
        # a config that was excellent two years ago and mediocre last quarter
        # can still win on the mean. The old build only NOTICED that after
        # the fact (a WARN in the ticket: "R² deteriorates toward the most
        # recent window"); nothing in the selection ever acted on it. That is
        # a fair criticism of the logic, and there are two honest halves to
        # the answer:
        #   1. It is DELIBERATE that the default does not chase recency. A
        #      hedge ratio is a covariance estimate; with cv_test_window=42
        #      each fold is ~2 months, so "the last two folds" is a handful
        #      of independent observations. Selecting on those alone swaps a
        #      known bias (stale regime) for a much larger variance (picking
        #      noise). Measured across configs the fold-to-fold spread of R²
        #      is routinely ±0.15 — bigger than most real deterioration.
        #   2. But the user should SEE the trade-off rather than have it
        #      hidden. So the recency figures are now computed for every
        #      config, always, and section [1b] reports what a
        #      recent-folds-only selection WOULD have chosen and what it
        #      would have scored. recency_weight>0 acts on it.
        _nf = int(cv['fold'].max()) + 1 if len(cv) else 1
        _recent_k = max(1, int(np.ceil(_nf / 2.0)))     # newest half of folds
        _rec = (cv[cv['fold'] < _recent_k]
                .groupby(['lookback', 'pc_spec', 'basket_size'])
                .agg(recent_r2=('r2_test', 'mean'),
                     recent_vol_red=('vol_red_test', 'mean'),
                     n_recent=('fold', 'count')).reset_index())
        summary = summary.merge(_rec, on=['lookback', 'pc_spec',
                                          'basket_size'], how='left')
        # per-config slope of R² against fold index. fold 0 is the NEWEST, so
        # a POSITIVE slope means R² rises as you go back = deteriorating now.
        # [Fix 114] computed with an explicit loop rather than
        # groupby.apply(..., include_groups=False): that keyword only exists
        # from pandas 2.2, and without it 2.2+ emits a DeprecationWarning
        # while <2.2 raises TypeError outright. A purely diagnostic column
        # must not be able to take the whole selection down on an older
        # kernel, and the loop is over a handful of configs either way.
        _sl_rows = []
        for _key, _g in cv.groupby(['lookback', 'pc_spec', 'basket_size']):
            _sl_rows.append({
                'lookback': _key[0], 'pc_spec': _key[1],
                'basket_size': _key[2],
                'r2_slope_per_fold': (
                    float(np.polyfit(_g['fold'].values,
                                     _g['r2_test'].values, 1)[0])
                    if _g['fold'].nunique() >= 3 else np.nan)})
        _sl = pd.DataFrame(_sl_rows)
        summary = summary.merge(_sl, on=['lookback', 'pc_spec',
                                         'basket_size'], how='left')
        # the score the ranking actually uses
        if recency_weight and recency_weight > 0:
            _w = np.exp(-float(recency_weight) * cv['fold'].values)
            _cvw = cv.assign(_w=_w, _r2w=cv['r2_test'].values * _w,
                             _vrw=cv['vol_red_test'].values * _w)
            _wm = (_cvw.groupby(['lookback', 'pc_spec', 'basket_size'])
                   .agg(_sw=('_w', 'sum'), _sr=('_r2w', 'sum'),
                        _sv=('_vrw', 'sum')).reset_index())
            _wm['w_r2'] = _wm['_sr'] / _wm['_sw']
            _wm['w_vr'] = _wm['_sv'] / _wm['_sw']
            summary = summary.merge(_wm[['lookback', 'pc_spec', 'basket_size',
                                         'w_r2', 'w_vr']],
                                    on=['lookback', 'pc_spec', 'basket_size'],
                                    how='left')
            _r2_use = summary['w_r2'].fillna(summary['avg_r2_oos'])
            _vr_use = summary['w_vr'].fillna(summary['avg_vol_red'])
        else:
            _r2_use, _vr_use = summary['avg_r2_oos'], summary['avg_vol_red']
        summary['composite'] = ((_r2_use - 0.5 * summary['std_r2_oos']) * 0.6
                                + _vr_use * 0.4
                                - gross_penalty * summary['avg_gross'])
        # the same score computed on the RECENT folds only — never used for
        # the live pick, always reported, so the alternative is visible
        summary['composite_recent'] = (
            (summary['recent_r2'].fillna(-1)) * 0.6
            + summary['recent_vol_red'].fillna(0) * 0.4
            - gross_penalty * summary['avg_gross'])
        if max_gross is not None:
            within = summary[summary['avg_gross'] <= max_gross]
            if len(within):
                summary = within.reset_index(drop=True)
            else:
                _P(f"  ⚠ no config satisfies max_gross={max_gross:.1f} — "
                      f"cap ignored for this run")
        return summary, cv


    # ── [Fix 77] SHORT-SAMPLE FLOW (HOLDOUT / IN-SAMPLE) ────────────────────
    # Too little data for the walk-forward machinery: fit once, score the
    # held-out tail if there is one, print an honest mini-report in the
    # [Fix 79] layout, and return the same result-dict shape as a full run.
    if val_mode in ('HOLDOUT', 'IN-SAMPLE'):
        _ps0 = 'auto' if pc_options == 'auto' else int(list(pc_options)[0])
        _bs0 = int(max(basket_size_options))
        fit_rets = rets.iloc[:rows - hold_rows - (purge if hold_rows else 0)]
        try:
            (_, s_rank, s_basket, s_reg, s_r2_in, s_resid, s_npc,
             _, _) = _select_and_fit(fit_rets, _ps0, _bs0)
        except RuntimeError as e:
            raise ValueError(
                f"[Fix 77] cannot fit on {len(fit_rets)} rows: {e}")
        s_w = pd.Series(s_reg.coef_, index=s_basket)
        # [Fix 83] adjusted R² — with k regressors on a short window the raw
        # in-sample R² is inflated by construction, and this tier is exactly
        # where that bites. A NEGATIVE adjusted R² means the fit explains
        # less than its degrees of freedom cost.
        _n_fit, _k_fit = len(fit_rets), len(s_basket)
        s_r2_adj = (1 - (1 - s_r2_in) * (_n_fit - 1) / max(1, _n_fit - _k_fit - 1)
                    if _n_fit > _k_fit + 1 else np.nan)
        s_te_in = float(np.std(s_resid) * ann_sqrt)
        s_tvol = float(fit_rets[target_ticker].std() * ann_sqrt)
        s_vr_in = 1 - s_te_in / s_tvol if s_tvol > 1e-6 else 0.0
        # single-instrument comparison on the SAME basis
        _okc = fit_rets.columns[(fit_rets.notna().all())
                                & (fit_rets.std() > 1e-6)]
        _pool = [c for c in _okc if c != target_ticker
                 and c not in set(pf_legs)]
        _acorr = (fit_rets[_pool].corrwith(fit_rets[target_ticker]).abs()
                  if _pool else pd.Series(dtype=float))
        _etf_p = [c for c in _pool if c.split(' ')[0] in ETF_SHORT_SET]
        _etf_c = _acorr[_etf_p].dropna() if _etf_p else pd.Series(dtype=float)
        s_best_etf = _etf_c.idxmax() if len(_etf_c) else None  # [Fix 83]
        s_r2e = np.nan
        s_reg_e = None
        if s_best_etf is not None:
            s_reg_e, s_r2e, _ = _fit_basket(fit_rets, target_ticker,
                                            [s_best_etf])
        s_r2_h = s_vr_h = s_te_h = s_r2e_h = np.nan
        if hold_rows:
            s_test = rets.iloc[rows - hold_rows:]
            s_r2_h, s_vr_h = _eval_on_window(s_reg, s_basket, s_test)
            if not np.isnan(s_r2_h):
                _rh = (s_test[target_ticker].values
                       - s_reg.predict(s_test[s_basket].values))
                s_te_h = float(np.std(_rh) * ann_sqrt)
            if s_reg_e is not None:
                s_r2e_h, _ = _eval_on_window(s_reg_e, [s_best_etf], s_test)
        # verdict — honest about how thin the evidence is
        if val_mode == 'HOLDOUT':
            _gap_s = (s_r2_in - s_r2_h) if not np.isnan(s_r2_h) else np.nan
            grade, reasons = _ticket_verdict(
                s_r2_h, np.nan, _gap_s, None, False, s_r2_h, s_r2e_h, False)
            if grade.startswith('✓'):
                grade = '△ TRADE AT REDUCED SIZE'
            reasons.insert(0, f'single {hold_rows}-row holdout — thin '
                              f'evidence; widen sample_window for a real '
                              f'verdict')
        else:
            grade = '? IN-SAMPLE ONLY — DO NOT SIZE OFF THIS'
            reasons = [
                f'{rows} rows: no out-of-sample evidence is possible',
                f'{len(s_basket)} regressors on {len(fit_rets)} obs — the '
                f'in-sample R² is inflated by construction',
                'use for exposure identification; validate on a longer '
                'window before trading']
        # display names (built later in the full flow — needed here)
        _shorts = pd.Series({c: c.split(' ')[0] for c in fit_rets.columns})
        _dup = _shorts[_shorts.duplicated(keep=False)].index
        disp = {c: (c if c in _dup else _shorts[c])
                for c in fit_rets.columns}
        if is_portfolio:
            disp[target_ticker] = pf_label
        # mini-report in the [Fix 79] layout
        L = 78
        # [Fix 114] SHORT-SAMPLE PATH — same renderer as the full report.
        # This branch kept its own print()-based title/kv, so a user who set
        # sample_window='3m' (a documented workflow) still got the packed
        # monospace layout the whole [Fix 104] pass existed to remove. It is
        # also the tier where the reader most needs to be told what a number
        # means, because none of these numbers are properly validated.
        title = _section
        def kv(label, value, implication=''):
            _KVCTX['rows'].append((label, value, implication))
        kvo = _kv_open
        title(f"PCA HEDGE — {target_short}   ·   "
              f"{datetime.now():%Y-%m-%d %H:%M}")
        _P(f"  {val_mode} validation · {len(universe)} names considered · "
           f"{'2-day overlapping' if two_day else 'daily'} returns · "
           f"{rows} rows ({rets.index[0]:%d%b%y}→{rets.index[-1]:%d%b%y})")
        title("[0] TICKET")
        _verdict_banner(grade, reasons,                        # [Fix 87]
                        f"{val_mode} tier — {TIER_MEANING[val_mode]}")
        _P(f"  EXECUTE   per $100 of "
           f"{'gross book' if is_portfolio else 'long ' + target_short}"
           + (f"  ($MM column: ${notional_mm:g}mm gross)"
              if notional_mm else ""))
        _t_head = ['ACTION', '$/100', 'INSTRUMENT', 'FULL TICKER', '★']
        if notional_mm:
            _t_head.insert(2, '≈$MM')
        _t_rows = []
        for _tk, _wgt in s_w.sort_values(key=lambda s: -s.abs()).items():
            _row = ['SHORT' if _wgt > 0 else 'LONG',
                    '$' + format(abs(_wgt) * 100, '.1f'),
                    disp.get(_tk, _tk.split(' ')[0]), _tk,
                    '★' if _tk in forced else '']
            if notional_mm:
                _row.insert(2, '$' + format(abs(_wgt) * notional_mm, ',.2f'))
            _t_rows.append(_row)
        _print_table(_t_head, _t_rows)
        _net = s_w.sum() * 100
        _gr = s_w.abs().sum() * 100
        kvo('THE POSITION')
        kv('Exposure', f"gross ${_gr:.1f} · net "
                       f"{'short' if _net >= 0 else 'long'} ${abs(_net):.1f}",
           f"You put on ${_gr:.0f} of hedge per $100 of book. Financing and "
           f"borrow scale with this and are not modelled.")
        kv('Hedged (in-sample)', f"vol {s_tvol*100:.1f}% → "
           f"{s_te_in*100:.1f}%/yr ({s_vr_in*100:.0f}% cut) · "
           f"R² {s_r2_in:.2f}",
           'Fit and scored on the SAME rows. On a window this short that is '
           'close to meaningless as a forecast — it is the ceiling, and the '
           'gap to reality is unmeasured here.')
        if notional_mm:                                        # [Fix 86]
            _te_s = s_te_h if np.isfinite(s_te_h) else s_te_in
            kv('RISK ($)', f"±{_fmt_money_mm(_te_s*notional_mm)}/yr · "
                           f"±{_fmt_money_mm(_te_s*notional_mm/np.sqrt(252))}"
                           f"/day",
               f"1σ residual P&L on ${notional_mm:g}mm gross"
               + (' — priced off the holdout.' if np.isfinite(s_te_h)
                  else ' — priced IN-SAMPLE, so treat it as a FLOOR on the '
                       'real risk, not an estimate of it.'))
        title("[1] EVIDENCE — all of it")
        kvo('EVIDENCE', note=f"{val_mode} tier: "
            f"{TIER_MEANING[val_mode]} Everything here is descriptive; "
            f"nothing on this path has survived a walk-forward.")
        kv('In-sample R²', f"{s_r2_in:.3f} on {len(fit_rets)} rows "
           f"({len(s_basket)} names, PCs {s_npc})"
           + (f" · adj {s_r2_adj:.3f}" if np.isfinite(s_r2_adj) else ""),
           'The adjusted figure charges the fit for its own degrees of '
           'freedom — with this many regressors on this few rows, that is '
           'the honest one of the two.')
        if np.isfinite(s_r2_adj) and s_r2_adj < 0.1:
            kv('  ⚠ adjusted R² ≈ 0', f"{s_r2_adj:.3f}",
               'On this window the fit barely beats its own '
               'degrees-of-freedom cost. Treat the basket as a HYPOTHESIS, '
               'not a hedge — re-run on a longer window before trading it.')
        if hold_rows:
            kv('Holdout OOS R²',
               (f"{s_r2_h:.3f} · vol red {s_vr_h*100:.1f}% · TE "
                f"{s_te_h*100:.1f}%/yr"
                if not np.isnan(s_r2_h)
                else 'n/a (incomplete data in the holdout)'),
               f"The only number here the fit did not see: the last "
               f"{hold_rows} rows. One window, so it is an anecdote — but it "
               f"is the only honest anecdote on the page.")
        if s_best_etf is not None:
            kv('Best single ETF', f"{disp.get(s_best_etf, s_best_etf)}: "
               f"in-sample R² {s_r2e:.2f}"
               + (f" · holdout {s_r2e_h:.2f}"
                  if not np.isnan(s_r2e_h) else ''),
               'If this is close to the basket, take the ETF — one liquid '
               'leg beats a fitted basket that has not been validated.')
        kv('Next step', f"needs ≥ ~150 rows for the full protocol",
           "For CV, nested windows and the rolling curve, give the model "
           "more data — e.g. sample_window='1y', or drop sample_window "
           "entirely.")
        if show_plots:
            kv('  charts', 'skipped',
               'Charts are suppressed in short-sample modes — there is not '
               'enough data behind them to be worth drawing.')
        _kv_flush()
        return {
            'weights': s_w, 'basket': s_basket, 'ranking': s_rank,
            'cv_summary': None, 'cv_detail': None,
            'nested_validation': pd.DataFrame(),
            'nested_r2_mean': np.nan, 'nested_r2_std': np.nan,
            'nested_vol_red_mean': np.nan, 'nested_pooled_r2': np.nan,
            'selected_unseen_r2': s_r2_h, 'selected_unseen_r2_std': np.nan,
            'selected_unseen_pooled_r2': s_r2_h,
            'outer_window_detail': pd.DataFrame(),
            'bootstrap_ci': None, 'rolling_oos': pd.DataFrame(),
            'decay': {}, 'turnover': {},
            'ticket': {'grade': grade, 'reasons': reasons,
                       'quality_r2': (s_r2_h if hold_rows else np.nan),
                       'quality_source': ('holdout' if hold_rows
                                          else 'in-sample only'),
                       'worst_window_r2': np.nan, 'worst_window_te': np.nan,
                       'kill_switch_te': np.nan, 'overfit_gap': np.nan,
                       'rebalance_rows': None},
            'nowcast': {}, 'horizon': {},                  # [Fix 91/93]
            'target_label': target_short,              # [Fix 103] for export
            'analysis_span': (rets.index[0], rets.index[-1], rows),  # [Fix 99]
            'holdout_r2': s_r2_h, 'holdout_vol_red': s_vr_h,
            'holdout_te': s_te_h, 'in_sample_r2': s_r2_in,
            'in_sample_r2_adj': s_r2_adj,                  # [Fix 83]
            'validation_mode': val_mode,
            'etf_benchmark': pd.DataFrame(), 'etf_nested_r2_mean': np.nan,
            'name_benchmark': pd.DataFrame(), 'name_nested_r2_mean': np.nan,
            'best_single': None, 'best_etf': s_best_etf,
            'returns': fit_rets, 'prices': prices,
            'two_day_returns': two_day,
            'excluded': removed_cols, 'force_included': forced,
            'is_portfolio': is_portfolio,
            'portfolio_legs': pf_legs, 'portfolio_weights': pf_w,
            'factor_leakage': None,
            'pit_size_active': pit_active, 'pit_members_active': pm_active,
            'params': {'lookback': len(fit_rets), 'pc_spec': _ps0,
                       'basket_size': _bs0,
                       'tz_tolerance_hours': tz_tolerance_hours,
                       'min_mktcap_usd_mm': min_mktcap_usd_mm,
                       'gross_penalty': gross_penalty, 'max_gross': max_gross,
                       'force_include': force_include,
                       'winsorize_pct': winsorize_pct,
                       'robust_cov': robust_cov,
                       'pit_size': pit_size, 'pit_members': pit_members,
                       'size_band': size_band,
                       'size_override': size_override,
                       'mask_stale_days': mask_stale_days,
                       'notional_mm': notional_mm, 'boot_n': boot_n,
                       'sample_window': sample_window, 'report': report},
        }

    # ── STEP 3: NESTED WALK-FORWARD VALIDATION [Fix 20] ──────────────────────
    # For each of the last `outer_folds` windows: rerun CV + selection + refit
    # using ONLY data prior to that window, then score on it. This measures
    # the whole pipeline out-of-sample, across multiple periods — not a
    # single-window estimate.
    _P(f"Nested walk-forward validation — the whole pipeline on data it has "
          f"never seen: {outer_folds} unseen "
          f"window{'s' if outer_folds != 1 else ''} of {outer_window}d, with "
          f"the config re-chosen from scratch before each one on "
          f"{cv_n_folds} earlier fold{'s' if cv_n_folds != 1 else ''}...")
    nested_rows = []
    nested_y_all, nested_res_all = [], []                     # [Fix 69]
    N = len(rets)
    for j in range(1, outer_folds + 1):
        o_end = N - (j - 1) * outer_window
        o_start = o_end - outer_window
        inner = rets.iloc[:max(0, o_start - purge)]   # [Fix 39] purge gap
        outer_test = rets.iloc[o_start:o_end]
        summary_j, _ = _walk_forward_cv(inner)
        if summary_j is None:
            _P(f"  outer window {j}: inner CV failed — skipped")
            continue
        cfg = summary_j.loc[summary_j['composite'].idxmax()]
        lb_j = int(cfg['lookback'])
        try:
            _, _, bkt_j, reg_j, _, _, _, _, _ = _select_and_fit(
                inner.iloc[-lb_j:], cfg['pc_spec'], int(cfg['basket_size']))
            r2_j, vr_j = _eval_on_window(reg_j, bkt_j, outer_test)
        except RuntimeError as e:
            _P(f"  outer window {j}: refit failed ({e}) — skipped")
            continue
        if not np.isnan(r2_j):                                # [Fix 69]
            _yo = outer_test[target_ticker].values
            nested_y_all.append(_yo)
            nested_res_all.append(_yo - reg_j.predict(
                outer_test[bkt_j].values))
        nested_rows.append({
            'window': j, 'start': outer_test.index[0], 'end': outer_test.index[-1],
            'lookback': lb_j, 'pc_spec': cfg['pc_spec'],
            'basket_size': int(cfg['basket_size']),
            'r2_oos': r2_j, 'vol_red_oos': vr_j,
        })
        _P(f"  window {j} ({outer_test.index[0]:%d %b %Y} → "
              f"{outer_test.index[-1]:%d %b %Y}): OOS R² {r2_j:.3f} · "
              f"vol red {vr_j*100:.1f}%   [picked lb={lb_j}, "
              f"pc={cfg['pc_spec']}, k={int(cfg['basket_size'])}]")
    nested = pd.DataFrame(nested_rows)
    nested_valid = nested.dropna(subset=['r2_oos']) if len(nested) else nested
    if len(nested_valid):
        nested_r2_mean = nested_valid['r2_oos'].mean()
        nested_r2_std = nested_valid['r2_oos'].std(ddof=0)
        nested_vr_mean = nested_valid['vol_red_oos'].mean()
    else:
        nested_r2_mean = nested_r2_std = nested_vr_mean = np.nan
        _P("  ⚠ nested validation produced no valid windows")
    # [Fix 69] pooled R² across the CONCATENATED nested windows — the same
    # statistic as [Fix 65]'s pooled figure, but for the HONEST pipeline
    # (selection re-run before each window). This is the headline number.
    nested_pooled = np.nan
    if nested_y_all:
        _yc = np.concatenate(nested_y_all)
        _rc = np.concatenate(nested_res_all)
        _sst = float(np.sum((_yc - _yc.mean()) ** 2))
        if _sst > 1e-12:
            nested_pooled = 1.0 - float(np.sum(_rc ** 2)) / _sst

    # ── [Fix 27/53] HONESTY BENCHMARKS: simplest hedges, same windows ──────
    # [Fix 27] scored a single best-correlated ETF on the outer windows.
    # [Fix 53] extends it in two ways, because section [3] used to quote
    # IN-SAMPLE single-name and single-ETF R² next to out-of-sample basket
    # numbers — a comparison the basket could not lose:
    #   1. there is now a single-NAME benchmark, not only a single-ETF one;
    #   2. the instrument AND the lookback are both chosen using only data
    #      BEFORE each outer window (per lookback take the best-|corr|
    #      candidate, keep the pair with the highest TRAIN R²). The old code
    #      hard-wired lookback = min(lookback_options), which handed the
    #      simple hedge a worse config than the basket was allowed to search.
    # Both benchmarks are scored on the SAME unseen outer windows as the PCA
    # numbers, so section [3] can compare like with like.
    def _single_bench(pool_filter):
        rows = []
        for j in range(1, outer_folds + 1):
            o_end = N - (j - 1) * outer_window
            o_start = o_end - outer_window
            best = None                      # (train_r2, pick, lookback, reg)
            for lb in lookback_options:
                tr_all = rets.iloc[max(0, o_start - purge - lb):o_start - purge]
                ok = tr_all.columns[(tr_all.notna().all()) &
                                    (tr_all.std() > 1e-6)]
                if target_ticker not in ok:
                    continue
                tr = tr_all[ok]
                pool = [c for c in ok if c != target_ticker
                        and c not in set(pf_legs) and pool_filter(c)]
                if not pool:
                    continue
                pick = tr[pool].corrwith(tr[target_ticker]).abs().idxmax()
                reg_b, r2_tr, _ = _fit_basket(tr, target_ticker, [pick])
                if best is None or r2_tr > best[0]:
                    best = (r2_tr, pick, lb, reg_b)
            if best is None:
                continue
            _, pick, lb, reg_b = best
            r2_b, vr_b = _eval_on_window(reg_b, [pick],
                                         rets.iloc[o_start:o_end])
            if not np.isnan(r2_b):
                rows.append({'window': j, 'instrument': pick, 'lookback': lb,
                             'r2_oos': r2_b, 'vol_red_oos': vr_b})
        return pd.DataFrame(rows)

    etf_bench = _single_bench(lambda c: c.split(' ')[0] in ETF_SHORT_SET)
    name_bench = _single_bench(lambda c: c.split(' ')[0] not in ETF_SHORT_SET)
    etf_r2_mean = etf_bench['r2_oos'].mean() if len(etf_bench) else np.nan
    etf_vr_mean = etf_bench['vol_red_oos'].mean() if len(etf_bench) else np.nan
    name_r2_mean = name_bench['r2_oos'].mean() if len(name_bench) else np.nan
    name_vr_mean = name_bench['vol_red_oos'].mean() if len(name_bench) else np.nan

    # [Fix 30] evaluate ANY config on the outer windows (trained strictly on
    # data preceding each window) — used for the selected config in [1] and
    # for every basket in the top-5 section
    outer_windows = [(N - j * outer_window, N - (j - 1) * outer_window)
                     for j in range(1, outer_folds + 1)]

    def _config_outer_oos(lb, ps, bs, detail=False):
        """[Fix 30] mean OOS stats for a config on the outer windows.
        [Fix 65] detail=True additionally returns per-window rows (dates,
        R², vol red, target vol, fitted weights) plus pooled residual /
        target arrays, so section [1b] can diagnose WHY the windows
        disagree instead of hiding the spread inside a mean ± std."""
        r2s, vrs, rows, res_all, y_all = [], [], [], [], []
        for wi, (os_, oe_) in enumerate(outer_windows, start=1):
            try:
                _, _, bkt_w, reg_w, _, _, _, _, _ = _select_and_fit(
                    rets.iloc[max(0, os_ - purge - lb):os_ - purge],
                    ps, bs)                           # [Fix 39] purge gap
                test = rets.iloc[os_:oe_]
                r2_w, vr_w = _eval_on_window(reg_w, bkt_w, test)
                if not np.isnan(r2_w):
                    r2s.append(r2_w)
                    vrs.append(vr_w)
                    if detail:
                        y_w = test[target_ticker].values
                        res_w = y_w - reg_w.predict(test[bkt_w].values)
                        rows.append({
                            'window': wi,           # 1 = most recent
                            'start': test.index[0], 'end': test.index[-1],
                            'r2': r2_w, 'vol_red': vr_w,
                            'tgt_vol': float(np.std(y_w) * ann_sqrt),
                            'weights': pd.Series(reg_w.coef_, index=bkt_w)})
                        res_all.append(res_w)
                        y_all.append(y_w)
            except RuntimeError:
                continue
        if not r2s:
            return ((np.nan, np.nan, np.nan, 0, pd.DataFrame(),
                     np.nan, None, None)
                    if detail else (np.nan, np.nan, np.nan, 0))
        out = (float(np.mean(r2s)), float(np.std(r2s)),
               float(np.mean(vrs)), len(r2s))
        if not detail:
            return out
        # [Fix 65] pooled OOS R²: one R² over the CONCATENATED windows.
        # Each per-window R² rests on ~outer_window observations and is a
        # noisy statistic; the pooled figure rests on all of them at once
        # and is the more stable headline of the two.
        y_cat = np.concatenate(y_all)
        r_cat = np.concatenate(res_all)
        sst = float(np.sum((y_cat - y_cat.mean()) ** 2))
        pooled = 1.0 - float(np.sum(r_cat ** 2)) / sst if sst > 1e-12 else np.nan
        # [Fix 70] hand back the pooled arrays so the bootstrap CI can
        # resample the exact series the pooled R² was computed on
        return out + (pd.DataFrame(rows), pooled, y_cat, r_cat)

    def _roll_stats_for(lb, ps, bs, max_windows=None, step=None):
        """[Fix 109] Walk-forward rolling OOS for ANY config, not just the
        selected one.

        The rolling curve was computed inline for the winner only, so the
        [4] ALTERNATIVES table could show a config's average OOS R² but not
        how it BEHAVED through time — which is the number that decides
        whether a challenger is really better. Factored out here so the
        alternatives are judged on the same evidence as the pick.

        Returns dict(te_median, te_worst, te_latest, r2_median, r2_worst,
        r2_latest, r2var_median, n_neg, n_rows) or None.
        """
        _st = max(5, int(step or rolling_refit_every))
        _nw = int(max_windows or rolling_max_windows)
        _ys, _rs = [], []
        for _i in range(_nw, 0, -1):
            _e = len(rets) - (_i - 1) * _st
            _s = _e - _st
            if _s < 0 or _s - purge - lb < 0:
                continue
            try:
                _, _, _bk, _rg, _, _, _, _, _ = _select_and_fit(
                    rets.iloc[_s - purge - lb:_s - purge], ps, bs)
            except RuntimeError:
                continue
            _te_r = rets.iloc[_s:_e]
            _need = _bk + [target_ticker]
            if not all(c in _te_r.columns and _te_r[c].notna().all()
                       for c in _need):
                continue
            _yv = _te_r[target_ticker].values
            _ys.append(_yv)
            _rs.append(_yv - _rg.predict(_te_r[_bk].values))
        if not _ys:
            return None
        _df = pd.DataFrame({'y': np.concatenate(_ys),
                            'resid': np.concatenate(_rs)})
        _w = 21
        if len(_df) < _w + 1:
            return None
        _sse = _df['resid'].pow(2).rolling(_w).sum()
        _sst = (_df['y'].rolling(_w).var(ddof=0) * _w).where(
            lambda s: s > 1e-12)
        _r2 = (1.0 - _sse / _sst).dropna()
        _r2v = (1.0 - (_df['resid'].rolling(_w).var(ddof=0) * _w)
                / _sst).dropna()
        _te = (_df['resid'].rolling(_w).std() * ann_sqrt).dropna()
        if not len(_te):
            return None
        return {'te_median': float(_te.median()),
                'te_worst': float(_te.max()),
                'te_latest': float(_te.iloc[-1]),
                'r2_median': float(_r2.median()) if len(_r2) else np.nan,
                'r2_worst': float(_r2.min()) if len(_r2) else np.nan,
                'r2_latest': float(_r2.iloc[-1]) if len(_r2) else np.nan,
                'r2var_median': (float(_r2v.median()) if len(_r2v)
                                 else np.nan),
                'n_neg': int((_r2 < 0).sum()),
                'n_scored': int(len(_r2)),
                'n_rows': int(len(_df))}

    # ── STEP 4: FULL-DATA CV → SELECT LIVE CONFIG ────────────────────────────
    _P(f"Full-data walk-forward CV — choosing the config to trade...")
    summary, cv_df = _walk_forward_cv(rets, verbose=True)
    if summary is None:
        raise RuntimeError("All CV folds failed — the stock likely lacks "
                           "sufficient liquidity or price history.")
    best = summary.loc[summary['composite'].idxmax()]
    best_lookback = int(best['lookback'])
    best_pc_spec = best['pc_spec']
    best_bsize = int(best['basket_size'])

    # ── STEP 5: FINAL REFIT ON FULL DATA (live weights) ──────────────────────
    try:
        (returns_final, final_rank, basket_final, reg_f,
         r2_final, resid_final, n_pc_final, pca_f, loadings_f) = _select_and_fit(
            rets.iloc[-best_lookback:], best_pc_spec, best_bsize)  # [Fix 46]
    except RuntimeError as _e_final:                       # [Fix 81]
        _tail = rets.iloc[-best_lookback:][target_ticker]
        _nan = int(_tail.isna().sum())
        raise ValueError(
            f"[Fix 81] cannot fit the FINAL {best_lookback}d window: "
            f"{_e_final}\n"
            f"  {target_short} has {_nan}/{len(_tail)} missing returns there"
            + (" — most likely a trading halt masked by the stale-price "
               "guard.\n  Options: raise mask_stale_days (or set it to None "
               "to disable), pick a\n  sample_window that avoids the halt, or "
               "hedge a liquid proxy instead."
               if _nan else
               " — the name may have zero variance (limit-locked or "
               "untraded).\n  Try a longer lookback or a different "
               "sample_window."))

    # [Fix 40] if a forced name has bad data in the FINAL window it cannot be
    # held — flag it explicitly rather than let the user assume it is inside
    miss_final = [c for c in forced if c not in basket_final]
    if miss_final:
        _P(f"  ⚠ force_include: {', '.join(miss_final)} could not be held "
              f"in the final basket (incomplete data or zero variance in the "
              f"final {best_lookback}d window)")

    weights_final = pd.Series(reg_f.coef_, index=basket_final)
    te_final = np.std(resid_final) * ann_sqrt
    tgt_std_final = returns_final[target_ticker].std() * ann_sqrt
    vol_red_final = 1 - te_final / tgt_std_final if tgt_std_final > 1e-6 else 0.0
    try:
        # [Fix 39] overlapping 2-day residuals carry induced MA(1)
        # autocorrelation that makes the ADF p-value anti-conservative;
        # test every 2nd (non-overlapping) residual instead
        resid_adf = resid_final[::2] if two_day else resid_final
        adf_p = adfuller(resid_adf, maxlag=min(20, len(resid_adf) // 5))[1]
    except Exception:
        adf_p = np.nan

    # [Fix 29] single-instrument alternatives chosen by realized hedging
    # power (|corr| with the target in the final window), not by rank_score
    corr_final = returns_final.corr()[target_ticker].drop(
        [target_ticker] + pf_legs, errors='ignore')        # [Fix 46] no self/legs
    best_single = corr_final.abs().idxmax()
    etf_cands = [c for c in corr_final.index if c.split(' ')[0] in ETF_SHORT_SET]
    best_etf = (corr_final[etf_cands].abs().idxmax()
                if etf_cands else best_single)   # [Fix 14]
    reg_s, r2_single, _ = _fit_basket(returns_final, target_ticker, [best_single])
    reg_e, r2_etf, _ = _fit_basket(returns_final, target_ticker, [best_etf])

    # ── Display names: short name; fall back to full ticker on collision ────
    shorts = pd.Series({c: c.split(' ')[0] for c in returns_final.columns})
    dup = shorts[shorts.duplicated(keep=False)].index
    disp = {c: (c if c in dup else shorts[c]) for c in returns_final.columns}
    if is_portfolio:                                       # [Fix 46]
        disp[target_ticker] = pf_label

    # [Fix 41] shared table renderer for every basket printout: same values
    # as before (ACTION / AMOUNT / name / full ticker), now in a boxed table,
    # plus a ★ marker on force-included names [Fix 40]
    def _basket_rows(weights):
        return [['SHORT' if wgt > 0 else 'LONG',
                 '$' + format(abs(wgt) * 100, '.1f'),
                 disp.get(tk, tk.split(' ')[0]), tk,
                 '★' if tk in forced else '']
                for tk, wgt in weights.sort_values(key=lambda s: -s.abs()).items()]

    BASKET_HEADERS = ['BUY / SELL', 'PER $100', 'NAME', 'FULL TICKER',
                      'PINNED']

    # ── [Fix 65/70/71/72] PRECOMPUTE OOS EVIDENCE ────────────────────────────
    # Used by the trade ticket, the charts and sections [1]/[1b]/[1c] below.
    # Computed once, up here, so the ticket and the diagnostics can never
    # disagree about the same number.
    (sel_r2m, sel_r2s, sel_vrm, sel_n,
     sel_detail, sel_pooled, sel_y, sel_res) = _config_outer_oos(
        best_lookback, best_pc_spec, best_bsize, detail=True)   # [Fix 65]
    n_eff = max(4, int(outer_window / (2 if two_day else 1)))
    wdisp = _window_dispersion_stats(sel_detail, sel_pooled, n_eff)
    boot = None
    if boot_n and sel_y is not None and len(sel_y):             # [Fix 70]
        boot = _block_bootstrap_ci(sel_y, sel_res, n_boot=int(boot_n),
                                   block=(8 if two_day else None),
                                   ann=ann_sqrt)

    # ── [Fix 71] ROLLING OOS CURVE ───────────────────────────────────────────
    # Three outer windows are three noisy readings. This walks the SELECTED
    # config forward: refit weights every `rolling_refit_every` rows on data
    # available at that point, score the NEXT segment out-of-sample, and
    # concatenate the daily OOS residuals into one series — stability becomes
    # a curve you can see instead of three scalars. The CONFIG was chosen on
    # full data, so the row-(3) selection caveat from [1] applies here too;
    # the WEIGHTS, however, never see their scoring segment.
    rolling_df = pd.DataFrame()
    roll_stats, decay_stats, turnover_stats = {}, {}, {}      # [Fix 78]
    if rolling_oos:
        _step = max(5, int(rolling_refit_every))
        _segs_y, _segs_r, _segs_ix, _n_skip = [], [], [], 0
        _dec_map, _bskts = {}, []                             # [Fix 84]
        for _i in range(int(rolling_max_windows), 0, -1):
            _e = len(rets) - (_i - 1) * _step
            _s = _e - _step
            if _s < 0 or _s - purge - best_lookback < 0:
                _n_skip += 1
                continue
            try:
                _, _, _bkt_r, _reg_r, _, _, _, _, _ = _select_and_fit(
                    rets.iloc[_s - purge - best_lookback:_s - purge],
                    best_pc_spec, best_bsize)
            except RuntimeError:
                _n_skip += 1
                continue
            _test_r = rets.iloc[_s:_e]
            _need = _bkt_r + [target_ticker]
            if not all(c in _test_r.columns and _test_r[c].notna().all()
                       for c in _need):
                _n_skip += 1
                continue
            _y_sr = _test_r[target_ticker].values
            _segs_y.append(_y_sr)
            _segs_r.append(_y_sr - _reg_r.predict(_test_r[_bkt_r].values))
            _segs_ix.append(_test_r.index)
            # [Fix 84] PAIRED staleness measurement. Fixes 78/82, which
            # bucketed residuals by raw row-age: with a fixed refit step,
            # each age bucket always landed on the same calendar offsets
            # (mod step), so the buckets differed by WHICH DAYS they
            # covered as much as by weight age — an abrupt regime break
            # produced a perfectly flat curve. Now every calendar row is
            # scored by SEVERAL vintages of weights (0, 1, 2 refits stale)
            # and only rows carrying all three are compared, so the groups
            # are identical in calendar terms and differ only in staleness.
            # This also states the answer in the trader's own units: what
            # does TE become if I skip a refit?
            _bskts.append((set(_bkt_r),
                           pd.Series(_reg_r.coef_, index=_bkt_r)))
            _H = 3 * _step
            if _s + _H <= len(rets):
                _hz = rets.iloc[_s:_s + _H]
                _okh = _hz[_need].notna().all(axis=1).values
                if _okh.any():
                    _sub = _hz.loc[_okh]
                    _rv = (_sub[target_ticker].values
                           - _reg_r.predict(_sub[_bkt_r].values))
                    _pos = np.arange(_s, _s + _H)[_okh]
                    _ago = (_pos - _s) // _step        # 0 / 1 / 2 refits old
                    for _p, _a, _r1 in zip(_pos, _ago, _rv):
                        _dec_map.setdefault(int(_p), {})[int(_a)] = float(_r1)
        if _segs_y:
            rolling_df = pd.DataFrame(
                {'y': np.concatenate(_segs_y),
                 'resid': np.concatenate(_segs_r)},
                index=pd.DatetimeIndex(
                    np.concatenate([ix.values for ix in _segs_ix])))
            _w_roll = 21
            # ── [Fix 105] WHY ROLLING R² GOES DEEPLY NEGATIVE ──────────────
            # Reported case: an equal-weighted AAPL / TSLA / SK Hynix basket
            # showed rolling R² ≈ -0.4 for one stretch and ≈ +0.8 for the
            # next. That is not a bug, but the old single number could not
            # tell you WHY, and it disagreed with the TE printed beside it.
            #
            # The reason is an asymmetry in the standard R² definition (the
            # one sklearn's .score() uses, so it was consistent across this
            # file, just not self-explanatory):
            #     R² = 1 - Σ(resid)² / Σ(y - ȳ)²
            # The DENOMINATOR removes the window mean from the target. The
            # NUMERATOR does not remove anything from the residual. So if the
            # hedge drifts — the target grinds away from the basket over the
            # window, which is exactly what an idiosyncratic run in one leg
            # does — that drift is charged in full to the numerator and the
            # benchmark it is measured against never pays it. Writing the
            # residual as mean m and standard deviation s:
            #     R² = 1 - (s² + m²)/σ_y²
            # A hedge that HALVES the target's vol (s = 0.5σ_y, a good hedge)
            # but drifts one target-sigma over the window (m = σ_y) scores
            #     R² = 1 - (0.25 + 1.00) = -0.25.
            # Meanwhile TE = s·√252 is computed with .std(), which DOES
            # demean — so the same window prints a perfectly healthy TE. That
            # is the contradiction: R² was measuring level tracking and TE
            # was measuring variance tracking, and both were called "how good
            # is the hedge".
            #
            # Measured on the synthetic reproduction of the reported case, 9
            # of the 13 windows with negative R² had the hedge genuinely
            # CUTTING variance; the negative sign was pure drift.
            #
            # Neither number is wrong — they answer different questions — so
            # both are now computed and both are shown:
            #   r2_roll      unchanged, the strict definition (level + shape)
            #   r2_var_roll  demeaned BOTH sides = pure variance reduction,
            #                the question "did hedging reduce my P&L swing?"
            #   drift_pen    m²·n / Σ(y-ȳ)², the gap between them = the share
            #                of R² lost purely to directional drift
            _resid_r = rolling_df['resid']
            _sse_r = _resid_r.pow(2).rolling(_w_roll).sum()
            _sst_r = rolling_df['y'].rolling(_w_roll).var(ddof=0) * _w_roll
            _sstw = _sst_r.where(_sst_r > 1e-12)
            rolling_df['r2_roll'] = 1.0 - _sse_r / _sstw
            rolling_df['r2_var_roll'] = 1.0 - (
                _resid_r.rolling(_w_roll).var(ddof=0) * _w_roll) / _sstw
            rolling_df['drift_pen'] = (
                _resid_r.rolling(_w_roll).mean().pow(2) * _w_roll) / _sstw
            rolling_df['te_roll'] = (_resid_r
                                     .rolling(_w_roll).std() * ann_sqrt)
            rolling_df['tgt_vol_roll'] = (rolling_df['y']
                                          .rolling(_w_roll).std() * ann_sqrt)
            _te_f = rolling_df['te_roll'].dropna()
            _r2_f = rolling_df['r2_roll'].dropna()
            _r2v_f = rolling_df['r2_var_roll'].dropna()
            if len(_te_f):
                _neg = rolling_df.dropna(subset=['r2_roll'])
                _neg = _neg[_neg['r2_roll'] < 0]
                roll_stats = {
                    'n_days': int(len(rolling_df)),
                    'n_segments': len(_segs_y), 'n_skipped': _n_skip,
                    'te_median': float(_te_f.median()),
                    'te_worst': float(_te_f.max()),
                    'te_latest': float(_te_f.iloc[-1]),
                    'r2_median': (float(_r2_f.median())
                                  if len(_r2_f) else np.nan),
                    'r2_worst': (float(_r2_f.min())
                                 if len(_r2_f) else np.nan),
                    'r2_latest': (float(_r2_f.iloc[-1])
                                  if len(_r2_f) else np.nan),
                    # [Fix 105] the variance-only twin and the drift split
                    'r2var_median': (float(_r2v_f.median())
                                     if len(_r2v_f) else np.nan),
                    'r2var_worst': (float(_r2v_f.min())
                                    if len(_r2v_f) else np.nan),
                    'r2var_latest': (float(_r2v_f.iloc[-1])
                                     if len(_r2v_f) else np.nan),
                    'n_neg': int(len(_neg)),
                    'n_scored': int(len(_r2_f)),
                    # of the negative windows, how many still cut variance?
                    'n_neg_but_cut': int((_neg['r2_var_roll'] > 0).sum())
                    if len(_neg) else 0,
                    'drift_median': (float(rolling_df['drift_pen']
                                           .dropna().median())
                                     if rolling_df['drift_pen'].notna().any()
                                     else np.nan)}
        # [Fix 78] TE by weight age → recommended rebalance cadence. Each
        # refit's weights are scored on the NEXT 1..63 rows; pooling the
        # residuals by age answers "how fast do fitted weights go stale?".
        # The recommended cadence is the longest age bucket whose TE is
        # still within 15% of fresh-weight TE. Ages are rows (≈ trading
        # days; 2-day overlapping rows in cross-tz mode).
        # [Fix 84] compare vintages on the SAME calendar rows
        _common = sorted(p for p, d in _dec_map.items() if len(d) == 3)
        if len(_common) >= 30:
            _te_ago = []
            for _a in range(3):
                _v = np.array([_dec_map[p][_a] for p in _common])
                _te_ago.append(((_a + 1) * _step,
                                float(np.std(_v) * ann_sqrt)))
            _fresh = _te_ago[0][1]
            _rec = _te_ago[0][0]
            for _iv, _te in _te_ago[1:]:
                if _te <= 1.15 * _fresh:
                    _rec = _iv
                else:
                    break
            decay_stats = {'te_by_interval': _te_ago, 'fresh_te': _fresh,
                           'rec_rows': _rec, 'n_rows': len(_common)}
        if len(_bskts) >= 2:                                  # [Fix 78]
            _js, _cs = [], []
            for (_sa, _wa), (_sb, _wb) in zip(_bskts[:-1], _bskts[1:]):
                _js.append(len(_sa & _sb) / max(1, len(_sa | _sb)))
                _names = sorted(_sa | _sb)
                _va = _wa.reindex(_names).fillna(0.0).values
                _vb = _wb.reindex(_names).fillna(0.0).values
                _na, _nb = np.linalg.norm(_va), np.linalg.norm(_vb)
                if _na > 1e-12 and _nb > 1e-12:
                    _cs.append(float(_va.dot(_vb) / (_na * _nb)))
            turnover_stats = {'mean_jaccard': float(np.mean(_js)),
                              'mean_cosine': (float(np.mean(_cs))
                                              if _cs else np.nan),
                              'n_refits': len(_bskts)}

    # ── [Fix 72] TICKET INPUTS + VERDICT ─────────────────────────────────────
    # quality = the most honest pooled OOS figure available: nested pooled
    # ([Fix 69], selection re-run before each window) first, then the
    # selected config's pooled ([Fix 65], mild selection bias), then the
    # nested mean as a last resort.
    if not np.isnan(nested_pooled):
        _quality, _quality_src = nested_pooled, 'nested pooled — honest'
    elif sel_n and not np.isnan(sel_pooled):
        _quality, _quality_src = sel_pooled, 'selected-config pooled'
    elif not np.isnan(nested_r2_mean):
        _quality, _quality_src = nested_r2_mean, 'nested mean'
    else:
        _quality, _quality_src = np.nan, 'n/a'
    _ref_pca_t = sel_r2m if sel_n else nested_r2_mean
    _simple_pool = [x for x in (etf_r2_mean, name_r2_mean)
                    if not np.isnan(x)]
    _ref_simple_t = max(_simple_pool) if _simple_pool else np.nan
    _overfit_gap = ((r2_final - _quality)
                    if not np.isnan(_quality) else np.nan)
    kill_te = (1.5 * wdisp['worst_te']
               if np.isfinite(wdisp['worst_te']) else np.nan)
    _roll_breach = bool(roll_stats and np.isfinite(kill_te)
                        and roll_stats['te_latest'] > kill_te)
    # ── [Fix 93/97] NOW-CAST — "I need to hedge NOW; is it working lately?"
    # The pooled numbers say how the hedge has behaved ON AVERAGE. This asks
    # about the present. But reacting to "the latest reading" is itself a
    # bias: one window carries only ~outer_window observations and scatters
    # widely, so a threshold on the newest point cries wolf on hedges that
    # are perfectly fine. Measured on stable synthetic hedges, the [Fix 93]
    # rule (latest TE > 1.25x its median, or ANY single signal) downgraded
    # 17% of them — one good hedge in six, for nothing.
    #
    # [Fix 97] makes it evidence-weighted instead:
    #   · judge a BLOCK of recent readings (the last ~horizon rows), not one
    #     point, so a single noisy day cannot move the verdict;
    #   · score that block against the hedge's OWN historical distribution
    #     of the same statistic (an empirical percentile), instead of an
    #     arbitrary ratio — "unusual for this hedge", not "above 1.25";
    #   · require PERSISTENCE (most of the block elevated), and either two
    #     independent signals or one extreme (>=95th pct) one;
    #   · it can only ever LOWER a tick from ✓ to △, never raise a grade,
    #     and it never replaces the pooled evidence.
    nowcast = {}
    if roll_stats or (sel_n and len(sel_detail)):
        _sig, _why = [], []
        # signal 1 — recent block of rolling TE vs its own history
        _pct = _persist = np.nan
        _blk = np.nan
        if len(rolling_df) and rolling_df['te_roll'].notna().sum() >= 12:
            _te_s = rolling_df['te_roll'].dropna()
            _k = int(min(max(10, outer_window), max(5, len(_te_s) // 3)))
            _recent, _hist = _te_s.iloc[-_k:], _te_s.iloc[:-_k]
            if len(_hist) >= 8:
                _blk = float(_recent.median())
                _pct = float((_hist <= _blk).mean())
                _persist = float((_recent > _hist.median()).mean())
                if _pct >= 0.85 and _persist >= 0.6:
                    _sig.append('te_block')
                    _why.append(
                        f"recent tracking error sits in the top "
                        f"{100*(1-_pct):.0f}% of this hedge's own history "
                        f"({_blk*100:.1f}%/yr) and stays there")
        # signal 2 — newest unseen window below the sampling-noise band
        _lw_r2 = _lw_flag = None
        if sel_n and len(sel_detail):
            _lw_r2 = float(sel_detail.iloc[0]['r2'])       # window 1 = latest
            _lw_flag = wdisp['flags'][0] if wdisp['flags'] else ''
            if _lw_flag == 'BELOW band':
                _sig.append('window')
                _why.append(f"newest unseen window R² {_lw_r2:.2f} is below "
                            f"the range sampling noise explains")
        # signal 3 — monotone deterioration toward the present
        if wdisp['trend']:
            _sig.append('trend')
            _why.append('R² has deteriorated steadily toward the present')
        # decide: two independent signals, or one extreme TE block
        _strong = (np.isfinite(_pct) and _pct >= 0.95 and _persist >= 0.6)
        if len(_sig) >= 2 or _strong:
            _cond = 'DEGRADED'
        elif (np.isfinite(_pct) and _pct <= 0.40 and not _sig):
            _cond = 'GOOD'
        else:
            _cond = 'NORMAL'
            if _sig:                       # one weak signal: note, don't act
                _why = [w + " (single signal — not acted on)" for w in _why]
        # sizing reference: the more conservative of the worst validated
        # window and current conditions — conservative by construction,
        # NOT "the latest reading is the truth"
        _te_now = max([x for x in (
            (_blk if np.isfinite(_blk) else None),
            (wdisp['worst_te'] if np.isfinite(wdisp['worst_te']) else None))
            if x is not None] or [np.nan])
        nowcast = {'condition': _cond, 'why': _why, 'signals': _sig,
                   'latest_window_r2': _lw_r2, 'latest_flag': _lw_flag,
                   'te_block': (float(_blk) if np.isfinite(_blk) else np.nan),
                   'te_pctile': (float(_pct) if np.isfinite(_pct) else np.nan),
                   'persistence': (float(_persist) if np.isfinite(_persist)
                                   else np.nan),
                   'sizing_te_now': _te_now}
    grade, reasons = _ticket_verdict(
        _quality, wdisp['worst_r2'], _overfit_gap, wdisp['min_sim'],
        wdisp['trend'], _ref_pca_t, _ref_simple_t, _roll_breach)
    if nowcast.get('condition') == 'DEGRADED':            # [Fix 93]
        if grade.startswith('✓'):
            grade = '△ TRADE AT REDUCED SIZE'
        for _w93 in nowcast['why']:
            if _w93 not in reasons:
                reasons.append(_w93)

    # ── [Fix 91/92] HORIZON — risk stated over the holding period ───────────
    horizon_info = {}
    if hedge_horizon_days:
        _h = int(hedge_horizon_days)
        _te_src = (nowcast.get('sizing_te_now')
                   if nowcast and np.isfinite(nowcast.get('sizing_te_now',
                                                          np.nan))
                   else (wdisp['worst_te'] if np.isfinite(wdisp['worst_te'])
                         else te_final))
        _res_src = (sel_res if sel_res is not None and len(sel_res)
                    else resid_final)
        _sig_h, _vr = _horizon_sigma(_res_src, _h, float(_te_src), two_day)
        _refits = (int(np.ceil(_h / decay_stats['rec_rows'])) - 1
                   if decay_stats else None)
        # [Fix 98] What did a hold of THIS length actually deliver? Split
        # the walk-forward OOS residuals into consecutive NON-OVERLAPPING
        # blocks of the horizon and report the realized spread. This is the
        # direct answer to "make it best over my 3-month horizon": it is the
        # empirical distribution of 3-month outcomes, not an annualized
        # abstraction, and it is what the config search was scored on.
        _blocks = None
        if len(rolling_df) >= 2 * _h:
            _res_b = rolling_df['resid'].values
            _y_b = rolling_df['y'].values
            _nb = len(_res_b) // _h
            _tes, _helped = [], 0
            for _i98 in range(_nb):
                _sl = slice(_i98 * _h, (_i98 + 1) * _h)
                _rt = float(np.std(_res_b[_sl]) * ann_sqrt)
                _yt = float(np.std(_y_b[_sl]) * ann_sqrt)
                _tes.append(_rt)
                if _yt > 1e-9 and _rt < _yt:
                    _helped += 1
            if _nb >= 2:
                _blocks = {'n': _nb, 'te_median': float(np.median(_tes)),
                           'te_worst': float(np.max(_tes)),
                           'te_best': float(np.min(_tes)),
                           'share_helped': _helped / _nb}
        horizon_info = {'days': _h, 'sigma_pct': float(_sig_h),
                        'vr': _vr, 'te_source': float(_te_src),
                        'refits_needed': _refits, 'blocks': _blocks}
    ticket_info = {'grade': grade, 'reasons': reasons,
                   'quality_r2': _quality, 'quality_source': _quality_src,
                   'worst_window_r2': wdisp['worst_r2'],
                   'worst_window_te': wdisp['worst_te'],
                   'kill_switch_te': kill_te,
                   'overfit_gap': _overfit_gap,
                   'rebalance_rows': decay_stats.get('rec_rows'),  # [Fix 78]
                   'nowcast': nowcast,                    # [Fix 93]
                   'horizon': horizon_info}               # [Fix 91]

    # ── STEP 6: CHARTS ───────────────────────────────────────────────────────
    # [Fix 117] Charts lead the output, as they always did. [Fix 115] moved them
    # below the report on the theory that the verdict should come first; that
    # was wrong for how this is actually used — the pictures are what you scan
    # to decide whether the ticket below is worth reading. They keep the
    # [Fix 115] restyling (report type scale, readable date ticks) and lose the
    # relocation.
    def _draw_charts():
        if not show_plots:
            return
        _CLR = {'tgt': '#c0392b', 'hedged': '#1b6b3a', 'pick': '#1f3f66',
                'other': '#9fb3c8', 'warn': '#b7791f', 'kill': '#8c1c1c'}
        _rc = {'font.size': 9, 'axes.titlesize': 11, 'axes.titleweight': 'bold',
               'axes.labelsize': 9, 'axes.labelcolor': _RF_MUTE,
               'axes.edgecolor': '#c8d2dd', 'axes.titlecolor': _RF_HEAD,
               'axes.grid': True, 'grid.color': '#eef1f5',
               'grid.linewidth': .8, 'axes.axisbelow': True,
               'xtick.color': _RF_MUTE, 'ytick.color': _RF_MUTE,
               'xtick.labelsize': 8, 'ytick.labelsize': 8,
               'legend.fontsize': 8, 'legend.frameon': False,
               'figure.facecolor': 'white', 'axes.spines.top': False,
               'axes.spines.right': False, 'lines.linewidth': 1.4}
        with plt.rc_context(_rc):
            fig, axes = plt.subplots(2, 3, figsize=(16, 8.5))
            ev = pca_f.explained_variance_ratio_
            axes[0, 0].bar(range(1, n_pc_final + 1), ev * 100,
                           color=_CLR['pick'], alpha=.75)
            axes[0, 0].plot(range(1, n_pc_final + 1), np.cumsum(ev) * 100,
                            'o-', color=_CLR['tgt'], ms=4)
            axes[0, 0].set_title(f'Scree — {n_pc_final} PCs, '
                                 f'{np.sum(ev)*100:.0f}% of variance')
            axes[0, 0].set_xlabel('principal component')
            axes[0, 0].set_ylabel('% of variance')
            # [Fix 62] label only what matters. Annotating every name is
            # unreadable past ~100 candidates: the target, the basket and the 30
            # best-ranked candidates carry all the information anyone reads this
            # panel for; the rest stay as unlabelled dots.
            lab = set(basket_final) | {target_ticker} | set(final_rank.index[:30])
            for tk in loadings_f.index:
                c = (_CLR['tgt'] if tk == target_ticker
                     else (_CLR['hedged'] if tk in basket_final
                           else _CLR['other']))
                s = 130 if tk == target_ticker else 34
                axes[0, 1].scatter(loadings_f.loc[tk, 'PC1'],
                                   loadings_f.loc[tk, 'PC2'], c=c, s=s,
                                   edgecolors='white', linewidths=.5, zorder=3)
                if tk in lab:
                    axes[0, 1].annotate(disp[tk],
                                        (loadings_f.loc[tk, 'PC1'],
                                         loadings_f.loc[tk, 'PC2']),
                                        fontsize=6.5, color=_RF_MUTE,
                                        xytext=(3, 3),
                                        textcoords='offset points')
            axes[0, 1].set_title('PC1 vs PC2 — red = target, green = basket')
            axes[0, 1].set_xlabel('PC1'); axes[0, 1].set_ylabel('PC2')
            corrs = (returns_final.corr()[target_ticker].drop(target_ticker)
                     .sort_values(ascending=False).head(10))
            colors = [_CLR['hedged'] if t in basket_final else _CLR['other']
                      for t in corrs.index]
            axes[0, 2].barh(range(len(corrs)), corrs.values, color=colors)
            axes[0, 2].invert_yaxis()
            axes[0, 2].set_yticks(range(len(corrs)))
            axes[0, 2].set_yticklabels([disp[t] for t in corrs.index],
                                       fontsize=7.5)
            axes[0, 2].set_title(f'Top correlations with {target_short} '
                                 f'— green = in the basket')
            axes[0, 2].set_xlabel('correlation')
            spread = pd.Series(resid_final, index=returns_final.index)
            # [Fix 15] true cumulative % = expm1(cumulative log return)
            axes[1, 0].plot(np.expm1(returns_final[target_ticker].cumsum()) * 100,
                            color=_CLR['tgt'], label='unhedged')
            axes[1, 0].plot(np.expm1(spread.cumsum()) * 100,
                            color=_CLR['hedged'], label='hedged')
            axes[1, 0].axhline(0, color='#b0bcc8', ls='--', lw=.9)
            axes[1, 0].legend(); axes[1, 0].set_ylabel('%')
            axes[1, 0].set_title('Cumulative return % — unhedged vs hedged')
            w = 20
            axes[1, 1].plot(returns_final[target_ticker].rolling(w).std()
                            * ann_sqrt * 100, color=_CLR['tgt'],
                            label='unhedged')
            axes[1, 1].plot(spread.rolling(w).std() * ann_sqrt * 100,
                            color=_CLR['hedged'], label='hedged')
            axes[1, 1].legend(); axes[1, 1].set_ylabel('%/yr')
            axes[1, 1].set_title(f'Rolling {w}-obs annualized vol %')
            for lb in summary['lookback'].unique():
                sub = (summary[summary['lookback'] == lb]
                       .groupby('basket_size')['avg_r2_oos'].mean())
                axes[1, 2].plot(sub.index, sub.values, 'o-', ms=4,
                                label=f'lookback {int(lb)}d')
            axes[1, 2].set_xlabel('basket size')
            axes[1, 2].set_ylabel('OOS R²')
            axes[1, 2].legend()
            axes[1, 2].set_title('Sensitivity — does a bigger basket help?')
            # date axes: matplotlib's default tick density overlapped the
            # labels into an unreadable smear on the two time-series panels
            for _ax in (axes[1, 0], axes[1, 1]):
                _ax.xaxis.set_major_locator(_mdates.AutoDateLocator(maxticks=7))
                _ax.xaxis.set_major_formatter(_mdates.DateFormatter('%b %y'))
            fig.tight_layout(pad=1.6)
            plt.show()
            # [Fix 71] rolling OOS curve — TE panel with the kill-switch line,
            # R² panel with the pooled level
            if len(rolling_df) and rolling_df['te_roll'].notna().any():
                fig2, ax2 = plt.subplots(2, 1, figsize=(13, 6), sharex=True)
                ax2[0].plot(rolling_df.index, rolling_df['te_roll'] * 100,
                            color=_CLR['hedged'])
                if np.isfinite(wdisp['worst_te']):
                    ax2[0].axhline(wdisp['worst_te'] * 100, color=_CLR['warn'],
                                   ls='--', label='worst validated window')
                    ax2[0].axhline(1.5 * wdisp['worst_te'] * 100,
                                   color=_CLR['kill'], ls='--',
                                   label='kill-switch (1.5×)')
                    # headroom + an opaque box, or the legend sits on top of
                    # the kill-switch line it is labelling
                    ax2[0].margins(y=.22)
                    ax2[0].legend(loc='lower left', frameon=True,
                                  facecolor='white', edgecolor='#dfe4ea',
                                  framealpha=.95)
                ax2[0].set_ylabel('%/yr')
                ax2[0].set_title('Rolling 21-row OOS tracking error — above '
                                 'the red line, stop')
                ax2[1].plot(rolling_df.index,
                            rolling_df['r2_roll'].clip(-0.5, 1.0),
                            color=_CLR['pick'])
                if not np.isnan(sel_pooled):
                    ax2[1].axhline(sel_pooled, color='#8a97a6', ls='--',
                                   label='pooled OOS R²')
                    ax2[1].legend(loc='lower left')
                ax2[1].set_title('Rolling 21-row OOS R² (clipped at −0.5)')
                ax2[1].xaxis.set_major_locator(
                    _mdates.AutoDateLocator(maxticks=9))
                ax2[1].xaxis.set_major_formatter(
                    _mdates.DateFormatter('%b %y'))
                fig2.tight_layout(pad=1.4)
                plt.show()

    _draw_charts()                                          # [Fix 117] up top

    # ── STEP 7: REPORT [Fix 79] ─────────────────────────────────────────────
    # One visual grammar: banner → [0] TICKET → [1] VALIDATION →
    # [2] COMPARISON → [3] DIAGNOSTICS → [4] ALTERNATIVES → [A] APPENDIX.
    # Sections are ordered by decision relevance; every diagnostic ends in a
    # one-line "►" takeaway. report='compact' (default) keeps every number
    # that changes the trade; report='full' adds the long-form explanations
    # and per-config weight tables.
    full_rep = (str(report).lower() == 'full')
    L = 78
    def hr(c='─'):
        _kv_flush()
        if not _IN_JUPYTER:
            _P(' ' + c * (L - 2))
    # [Fix 104] title/kv/print now route through the shared renderers. kv()
    # BUFFERS into the current block so a run of kv() calls comes out as ONE
    # table instead of 20 loose prints; _section/_print_table/_P flush it.
    # NOTE: `print` is NOT rebound here — find_best_hedge prints from the CV
    # and nested-validation loops long before this point, and a local rebind
    # would make every one of those an UnboundLocalError. The report body's
    # own print() call sites are written as _P() instead.
    title = _section
    def kv(label, value, implication=''):
        _KVCTX['rows'].append((label, value, implication))
    kvo = _kv_open

    title(f"PCA HEDGE — {target_short}"
          + ("  ·  RE-HEDGE" if _rehedge_of else '')
          + f"   ·   {datetime.now():%Y-%m-%d %H:%M}")
    if _rehedge_of:                                        # [Fix 115]
        _pp = _rehedge_of.get('params') or {}
        _P(f"  ► RE-HEDGE: weights refit on today's data using the config "
           f"already validated by the last full run "
           f"(lookback {_pp.get('lookback')}d · PCs {_pp.get('pc_spec')} · "
           f"basket {_pp.get('basket_size')}). The config was NOT re-searched.")
        _P(f"  the validation below re-runs on {outer_folds} outer window"
           f"{'s' if outer_folds != 1 else ''} rather than the full 3 — it is "
           f"a sanity check on the refit, not a fresh validation. The "
           f"config's credentials come from the run it inherits from. Press "
           f"RUN HEDGE (or pass full_validation=True) for the full protocol.")
    _P(f"  {val_mode} validation · {len(universe)} names considered · "
          f"{len(rets)} days of history "
          f"({rets.index[0]:%d %b %Y} → {rets.index[-1]:%d %b %Y})"
          + ("  ·  returns paired over 2 days because the markets close at "
             "different times" if two_day else ""))
    if sample_window is not None:                          # [Fix 99]
        _P(f"  period locked: every number below uses ONLY "
              f"{rets.index[0]:%d %b %Y} → {rets.index[-1]:%d %b %Y}. "
              f"Nothing outside it is fitted on, chosen on or scored on.")
    if hedge_horizon_days:                             # [Fix 91/100]
        _P(f"  you plan to hold this for {int(hedge_horizon_days)} trading "
              f"days, so risk is quoted over that period and [0] shows what "
              f"holds of that length actually delivered")
    if forced:
        _P(f"  names you pinned into every basket (★): "
              f"{', '.join(disp.get(c, c) for c in forced)}")

    # ── [0] TICKET ──────────────────────────────────────────────────────────
    # [Fix 116] PLAIN ENGLISH FIRST. Every label in this report used to be the
    # NAME OF A STATISTIC — 'Evidence', 'Exposure', 'Nested pooled OOS R²'. A
    # reader who does not already know the method cannot decode any of those,
    # and this section is the one everybody reads. Labels are now written as
    # two lines: what the row TELLS you, then the technical term underneath in
    # small type for anyone who wants it. Nothing computed changed — only the
    # words the numbers are filed under.
    title("[0] TICKET", "what to trade, how big, and when to stop")
    _sub = ''
    if not np.isnan(ticket_info['quality_r2']):
        # [Fix 86] plain English: R² is "share of the target's moves the
        # hedge tracks", vol reduction is the risk actually removed.
        _sub = (f"tracks {ticket_info['quality_r2']*100:.0f}% of "
                f"{target_short} out-of-sample")
        if sel_n and np.isfinite(sel_vrm):
            _sub += f", cutting {sel_vrm*100:.0f}% of its risk"
    # ── [Fix 117] THE TICKET, DE-DUPLICATED ────────────────────────────────
    # Measured on the previous build's own output: 73% appeared 7 times, 72%
    # five times, and 21.5% / 18.5% / 32.2% / 17.4% three times each — inside
    # ONE section. Three blocks were restating each other:
    #   QUALITY          repeated the verdict banner's subtitle verbatim, and
    #                    [1]'s headline row.
    #   AT A n-DAY HOLD  restated four rows that HOW TO RUN IT had already
    #                    printed two rows above, plus one from [3].
    #   THE POSITION     quoted the in-sample TE that SIZING then quoted again
    #                    as its "best case" comparison.
    # Both redundant blocks are gone. What is left is two blocks with no
    # overlap at all: WHAT YOU ARE PUTTING ON, then HOW TO RUN IT. Every
    # number in the ticket now appears exactly once, which is the only way the
    # eye can find the one it came for.
    #
    # Vocabulary: standard desk terms (R², TE, gross/net, OOS, bp) are used
    # plainly as labels — a trader reads them faster than any paraphrase. The
    # second, smaller label line is reserved for the things that are specific
    # to THIS tool and genuinely need saying: which window a number came from,
    # what "condition" means, what the stop is measured on.
    _verdict_banner(grade, reasons, _sub)                     # [Fix 87]
    _P(f"  THE TRADE — per $100 of "
          f"{'gross book' if is_portfolio else 'long ' + target_short}"
          + (f", and on the ${notional_mm:g}mm entered" if notional_mm else ""))
    _t_head = ['SIDE', '$/100', 'NAME', 'FULL TICKER', '★']
    if notional_mm:
        _t_head.insert(2, '≈$MM')
    _t_rows = []
    for _tk, _wgt in weights_final.sort_values(key=lambda s: -s.abs()).items():
        _row = ['SHORT' if _wgt > 0 else 'LONG',
                '$' + format(abs(_wgt) * 100, '.1f'),
                disp.get(_tk, _tk.split(' ')[0]), _tk,
                '★' if _tk in forced else '']
        if notional_mm:
            _row.insert(2, '$' + format(abs(_wgt) * notional_mm, ',.2f'))
        _t_rows.append(_row)
    _print_table(_t_head, _t_rows)
    _net = weights_final.sum() * 100
    _te_bp = te_final * 1e4 / np.sqrt(252)
    _gross = weights_final.abs().sum() * 100

    # ── what you are putting on ────────────────────────────────────────────
    kvo('WHAT YOU ARE PUTTING ON')
    kv('Gross / net',
       f"${_gross:.1f} gross · ${abs(_net):.1f} net "
       f"{'short' if _net >= 0 else 'long'} per $100",
       (f"Above $200 gross usually means the ridge paired off near-duplicate "
        f"names rather than hedging real risk — see [4] for a cheaper config."
        if _gross > 200 else
        f"A normal single-layer hedge. Borrow and financing scale with gross "
        f"and are NOT priced here."))
    _qbits = []
    if not np.isnan(ticket_info['quality_r2']):
        _qbits.append(f"{ticket_info['quality_r2']:.2f} pooled")
    if boot:
        _qbits.append(f"CI {boot['r2_lo']:.2f}–{boot['r2_hi']:.2f}")
    if not np.isnan(wdisp['worst_r2']):
        _qbits.append(f"{wdisp['worst_r2']:.2f} worst window")
    if _qbits:
        kv('OOS R²\nnested pooled — the honest one', ' · '.join(_qbits),
           'Judge the hedge on the worst window, not the pooled average. '
           'Full breakdown in [1].')
    kv('Vol cut\nIN-SAMPLE — the ceiling',
       f"{tgt_std_final*100:.1f}% → {te_final*100:.1f}%/yr · "
       f"{vol_red_final*100:.0f}% cut · {_te_bp:.0f} bp/day",
       'Scored on its own fitting data, so it is the best this hedge ever '
       'looks. Size off the row below, not this one.')

    # ── how to run it ──────────────────────────────────────────────────────
    kvo('HOW TO RUN IT', note='Every row is an instruction. Nothing here is '
        'repeated from the block above.')
    if np.isfinite(wdisp['worst_te']):
        kv('SIZE OFF\nworst validated window TE',
           f"{wdisp['worst_te']*100:.1f}%/yr",
           f"Not the {te_final*100:.1f}% in-sample figure above. This is what "
           f"the hedge did when it was working hardest, and it is what you "
           f"have to survive.")
    if notional_mm:                                           # [Fix 86]
        # dollarize the risk: a %/yr tracking error means little at the desk
        # until it is the 1-sigma P&L swing on the actual ticket size.
        _te_use = (wdisp['worst_te'] if np.isfinite(wdisp['worst_te'])
                   else te_final)
        _rl = (f"±{_fmt_money_mm(_te_use*notional_mm)}/yr · "
               f"±{_fmt_money_mm(_te_use*notional_mm/np.sqrt(252))}/day")
        if horizon_info:
            _rl += (f" · ±{_fmt_money_mm(horizon_info['sigma_pct']*notional_mm)}"
                    f" over the {horizon_info['days']}d hold "
                    f"(±{horizon_info['sigma_pct']*100:.1f}%)")
        kv('RISK ($)\n1σ residual P&L, priced off the worst window', _rl,
           f"On ${notional_mm:g}mm gross. Two-thirds of the time you land "
           f"inside these; roughly 1 in 20 is more than twice.")
    elif horizon_info:
        kv(f"RISK over the {horizon_info['days']}d hold\n1σ residual",
           f"±{horizon_info['sigma_pct']*100:.2f}%",
           'Two-thirds of holds land inside it; about 1 in 20 is more than '
           'twice it. This is the number for a risk limit.')
    if nowcast:                                        # [Fix 93/97]
        _nc = nowcast['condition']
        # [Fix 115] '·' is the separator the renderer splits readings on, so a
        # '·' used as a status glyph came out as a fact of its own.
        _mark = {'GOOD': '✓', 'NORMAL': '•', 'DEGRADED': '⚠'}[_nc]
        _bits = []
        if np.isfinite(nowcast.get('te_pctile', np.nan)):
            _bits.append(f"TE {nowcast['te_block']*100:.1f}%/yr "
                         f"({_ordinal(nowcast['te_pctile']*100)} pctile of its "
                         f"own history)")
        if nowcast.get('latest_window_r2') is not None:
            _bits.append(f"newest window R² "
                         f"{nowcast['latest_window_r2']:.2f}"
                         + (f" ({nowcast['latest_flag']})"
                            if nowcast.get('latest_flag') else ''))
        kv('CONDITION NOW\nlatest readings vs this hedge’s own history',
           f"{_mark} {_nc} · " + ' · '.join(_bits),
           ('Behaving as it always has — trade it at full size.'
            if _nc == 'GOOD' else
            ('Nothing unusual — trade it and keep the stop below.'
             if _nc == 'NORMAL' else
             'NOT behaving as validated — cut size, or refit before adding '
             'risk.')))
    if horizon_info and horizon_info.get('blocks'):        # [Fix 98]
        _hb = horizon_info['blocks']
        kv(f"{horizon_info['days']}-DAY HOLDS\nthe only evidence on your "
           f"actual horizon",
           f"{_hb['n']} holds · TE {_hb['te_median']*100:.1f}% median / "
           f"{_hb['te_worst']*100:.1f}% worst · cut risk in "
           f"{_hb['share_helped']*100:.0f}%",
           ('Below ~70% means the hedge is a coin-flip over a hold this long.'
            if _hb['share_helped'] < 0.7 else
            'Plan for the worst of those, not the median.'))
    if horizon_info and np.isfinite(horizon_info['vr']):
        _hh = horizon_info
        if _hh['vr'] > 1.15:
            kv('  horizon risk is padded\n  residuals trended',
               f"×{np.sqrt(_hh['vr']):.2f}",
               'A daily figure would understate a hold this long. The uplift '
               'is already in the RISK row. Treat it as a floor.')
        elif _hh['vr'] < 0.85:
            kv('  no discount taken\n  residuals mean-reverted',
               f"×{np.sqrt(_hh['vr']):.2f} available",
               'Would justify a smaller RISK number. Deliberately not taken — '
               'the effect does not repeat reliably enough to size on.')
    if np.isfinite(kill_te):
        _mon = f"kill at {kill_te*100:.1f}%/yr"
        if roll_stats:
            _mon += (f" · {roll_stats['te_latest']*100:.1f}% now"
                     + (' ⚠ ABOVE' if _roll_breach else ' ✓'))
        if decay_stats:
            _mon += f" · refit every {decay_stats['rec_rows']} rows"
        if horizon_info and horizon_info.get('refits_needed'):
            _mon += f" · ~{horizon_info['refits_needed']} refit(s) inside the hold"
        kv('STOP / REFIT\nrolling 21-row TE vs the kill level', _mon,
           (f"Above {kill_te*100:.1f}%/yr the hedge is no longer doing what it "
            f"was validated to do — refit or cut."
            + (' It is ALREADY above: do not add risk today.'
               if _roll_breach else '')
            + (f" Separately, refit every ~{decay_stats['rec_rows']} rows — "
               f"past that the weights decay faster than a refit costs."
               if decay_stats else '')))
    # [Fix 117] the horizon block that used to sit here restated four rows
    # printed immediately above it. Its one real contribution was telling you
    # WHICH of them dominates at your holding period — that is one line, not a
    # table.
    if hedge_horizon_days:
        _P("  ► over a "
           + (f"{int(hedge_horizon_days)}-day hold, a small persistent tilt "
              f"compounds faster than tracking noise — read the ADF and "
              f"FACTOR EXPOSURE rows in [3] before CONDITION NOW, which will "
              f"have turned over long before you are out"
              if int(hedge_horizon_days) >= 63 else
              f"{int(hedge_horizon_days)}-day hold, variance dominates — "
              f"CONDITION NOW and the STOP row are the lines that matter; "
              f"drift in [3] is secondary"))

    # ── [1] VALIDATION ──────────────────────────────────────────────────────
    title("[1] VALIDATION", "how good is this hedge, really?")
    # [Fix 115] this block used to open with a bare kv() and inherit whatever
    # headers the previous block had left in _KVCTX — most often the horizon
    # block's 'WATCH / THIS RUN / WHY AT THIS HORIZON', so the validation
    # numbers came out under a promise of a third column that was never filled.
    # [Fix 117] R² is R² — a trader reads the term faster than a paraphrase of
    # it, and the four flavours below only make sense NAMED, because the whole
    # point of the block is that they answer four different questions. The
    # small second label line says which question each one answers; the
    # appendix spells the four out side by side.
    kvo(headers=('MEASURE', 'THIS RUN', ''))
    if not np.isnan(nested_pooled):
        kv('► Nested pooled OOS R²\nthe headline — whole pipeline re-run',
           f"{nested_pooled:.3f}   (selection re-run before each window, so "
           f"nothing downstream ever saw the scoring data)")
    if sel_n:
        kv('This config, unseen OOS\nweights honest, config not',
           f"{sel_r2m:.3f} ± {sel_r2s:.3f} over {sel_n}×{outer_window}d · "
           f"pooled {sel_pooled:.3f}"
           + (f" · CI {boot['r2_lo']:.2f}–{boot['r2_hi']:.2f}" if boot else ''))
        kv('OOS vol reduction\nthe risk actually removed',
           f"{sel_vrm*100:.1f}%"
           + (f" · TE CI {boot['te_lo']*100:.1f}–{boot['te_hi']*100:.1f}%/yr"
              if boot else ''))
    kv('Full-data CV R²\nwhat the selection optimised',
       f"{best['avg_r2_oos']:.3f} ± {best['std_r2_oos']:.3f} "
       f"({int(best['n_folds'])} folds) — flattered by construction; never "
       f"rank configs on it")
    _ref_oos = sel_r2m if sel_n else nested_r2_mean
    if not np.isnan(_ref_oos):
        _gap = r2_final - _ref_oos
        _gv = ('✓ generalizes' if _gap < 0.15 else
               '△ some overfit' if _gap < 0.35 else '✗ overfit')
        kv('In-sample vs unseen\nthe overfit check',
           f"{r2_final:.3f} vs {_ref_oos:.3f} (gap {_gap:+.2f}) → {_gv}")
    kv('Selected config', f"lookback {best_lookback}d · PCs {best_pc_spec}"
       f"{f' ({n_pc_final})' if best_pc_spec == 'auto' else ''} · "
       f"basket {best_bsize}")
    if len(nested_valid):
        _picks = nested_valid.apply(
            lambda r: f"lb={int(r['lookback'])}, pc={r['pc_spec']}, "
                      f"size={int(r['basket_size'])}", axis=1)
        _modal = _picks.mode().iloc[0]
        _live = f"lb={best_lookback}, pc={best_pc_spec}, size={best_bsize}"
        if _modal != _live:
            _P(f"  ⚠ the pipeline picked [{_modal}] in "
                  f"{int((_picks == _modal).sum())}/{len(_picks)} unseen "
                  f"windows, not the live [{_live}] — see [4]")
    if sel_n >= 2 and len(sel_detail):
        _P()
        w_rows = []
        for _wi, (_, r) in enumerate(sel_detail.iterrows()):
            w_rows.append([f"{int(r['window'])}"
                           f"{' (latest)' if r['window'] == 1 else ''}",
                           f"{r['start']:%d%b}–{r['end']:%d%b%y}",
                           format(r['r2'], '.3f'),
                           format(r['vol_red'] * 100, '.1f') + '%',
                           format(r['tgt_vol'] * 100, '.1f') + '%',
                           wdisp['flags'][_wi]])
        _print_table(['WIN', 'DATES', 'OOS R²', 'VOL RED', 'TGT VOL',
                      'VS NOISE BAND'], w_rows,
                     title='THE THREE UNSEEN WINDOWS, ONE BY ONE',
                     note='TGT VOL is the target\'s own volatility in that '
                          'window — R² is mechanically depressed in calm '
                          'windows, so judge those by TE. NOISE BAND is the '
                          'spread a perfectly stable hedge would show on '
                          'windows this short.')
        _band = wdisp['band']
        _below = [str(int(r['window'])) for _wi, (_, r)
                  in enumerate(sel_detail.iterrows())
                  if wdisp['flags'][_wi] == 'BELOW band']
        if _band is not None:
            if not _below:
                _P(f"  ► the spread between those stretches is no more than "
                      f"chance would produce ({_band[0]*100:.0f}–"
                      f"{_band[1]*100:.0f}%) — it is not instability, so "
                      f"trust the combined number")
            else:
                _P(f"  ► stretch(es) {', '.join(_below)} came in BELOW what "
                      f"chance alone explains ({_band[0]*100:.0f}–"
                      f"{_band[1]*100:.0f}%) — something real went wrong "
                      f"there, not just a short sample")
        if np.isfinite(wdisp['reg_corr']) and wdisp['reg_corr'] > 0.5:
            _P(f"  ► the weak stretches are the CALM ones — a tracking score "
                  f"is mechanically lower when there is less to track. Judge "
                  f"those stretches by how far the hedge drifted instead")
        if wdisp['min_sim'] is not None:
            _stab = ('the same basket every time' if wdisp['min_sim'] > 0.8
                     else 'a drifting basket' if wdisp['min_sim'] > 0.5
                     else 'a DIFFERENT basket each time')
            _msg = (f"  ► re-picked on each stretch, the model chose "
                    f"{_stab} ({wdisp['min_sim']*100:+.0f}% alike at worst)")
            if wdisp['min_sim'] <= 0.5:
                _msg += (" — that churn is the real source of the spread "
                         "above. Try a longer lookback, a smaller basket, or "
                         "pin names with force_include")
            _P(_msg)
        if wdisp['trend']:
            _P("  ► ⚠ it tracked WORSE in the most recent stretches. Do not "
               "average that away — the newest reading is the best guess at "
               "tomorrow")
        # ── [Fix 108] IS THE WINNER STILL THE WINNER *NOW*? ────────────────
        # The selection maximises the composite over ALL CV folds. If the
        # hedge is decaying, that average is partly historical. This block
        # states the trade-off explicitly instead of leaving the user to
        # infer it: how fast this config is decaying, what a recent-folds-
        # only selection would have picked instead, and what that
        # alternative scored on the SAME recent folds.
        try:
            _bkey = ['lookback', 'pc_spec', 'basket_size']
            _srt = summary.sort_values('composite', ascending=False)
            _cur = _srt.iloc[0]
            _rsrt = summary.sort_values('composite_recent', ascending=False)
            _alt = _rsrt.iloc[0]
            _same = all(_cur[k] == _alt[k] for k in _bkey)
            _slope = _cur.get('r2_slope_per_fold', np.nan)
            kvo('RECENCY — is this still the right config TODAY?',
                note='The config was picked on ALL folds, weighting three '
                     'years ago the same as last month. These rows say what '
                     'that costs, and what a recency-only pick would be.')
            if np.isfinite(_slope):
                kv('Decay of the pick\nR² slope per fold',
                   f"{-_slope:+.3f} per {cv_test_window}-row fold",
                   ('HOLDING UP — its OOS R² is flat or improving as the folds '
                    'get newer.' if _slope <= 0.005 else
                    f"Loses ~{_slope:.3f} of R² per fold going forward, or "
                    f"~{_slope * max(int(_cur.get('n_folds', 1)) - 1, 1):.2f} "
                    f"across the {int(_cur.get('n_folds', 0))} folds measured. "
                    f"The headline average is flattered by the older folds; "
                    f"the newest reading is the better estimate of tomorrow."))
            if np.isfinite(_cur.get('recent_r2', np.nan)):
                kv('  all folds vs recent',
                   f"{_cur['avg_r2_oos']:.3f} all · {_cur['recent_r2']:.3f} "
                   f"newest {int(_cur.get('n_recent', 0))} folds",
                   'The gap between these two IS the question. Small means the '
                   'average is a fair summary; large means you are trading on '
                   'mostly historical evidence.')
            if _same:
                kv('  recency-only pick', 'SAME config',
                   'Selecting on the newest folds alone lands on this same '
                   'config, so the ranking is not an artifact of averaging '
                   'over old regimes.')
            else:
                kv('  recency-only pick',
                   f"lb {int(_alt['lookback'])} · pc {_alt['pc_spec']} · "
                   f"k {int(_alt['basket_size'])}",
                   f"A DIFFERENT config leads on the newest "
                   f"{int(_alt.get('n_recent', 0))} folds "
                   f"({_alt.get('recent_r2', float('nan')):.3f} vs "
                   f"{_cur.get('recent_r2', float('nan')):.3f}). Not chosen, "
                   f"on purpose — a handful of recent folds usually buys "
                   f"noise. Treat it as a challenger: run it from [4], or set "
                   f"recency_weight=0.35 and re-run.")
            if recency_weight:
                kv('  recency_weight', f"{recency_weight:g} (ACTIVE)",
                   f"The pick above already discounts older folds. The scores "
                   f"in the table at the top of [1] are unaffected — they "
                   f"never reweight anything.")
        except Exception as _e108:
            # [Fix 114] never swallow a diagnostic silently — a section that
            # just is not there reads as "nothing to report", which is the
            # opposite of the truth. Also flush whatever rows were already
            # buffered, or they surface inside the NEXT block.
            kv('  recency check', 'unavailable',
               f"could not be computed ({type(_e108).__name__}: {_e108}) — "
               f"the rest of the report is unaffected")
        _kv_flush()
        if full_rep:
            _P("    (noise band = Fisher-z scatter a STABLE hedge with "
                  "this pooled R² would produce on")
            _P("    windows of this length. It is derived for a single "
                  "predictor, so for a multi-name")
            _P("    basket the true band is somewhat WIDER — a marginal "
                  "'BELOW band' flag is weak")
            _P("    evidence. The CI conditions on the fitted weights.)")
    if roll_stats:
        kvo('ROLLING OOS — 21-row walk-forward',
            note='Every 21 trading days the weights are refit and scored on '
                 'the days that follow. The first three rows answer different '
                 'questions — when they disagree, the "cost of one-way drift" '
                 'row is the whole difference.')
        kv('Tracking error', f"median {roll_stats['te_median']*100:.1f}% · "
           f"worst {roll_stats['te_worst']*100:.1f}% · latest "
           f"{roll_stats['te_latest']*100:.1f}%/yr",
           'Sets the risk budget and the kill level.')
        kv('R² (strict)\ndrift charged', f"median "
           f"{roll_stats['r2_median']:+.2f} · worst "
           f"{roll_stats['r2_worst']:+.2f} · latest "
           f"{roll_stats['r2_latest']:+.2f}",
           'Charges BOTH mistracking and directional drift, so it can go '
           'negative while the hedge is still cutting risk — the next two '
           'rows say why.')
        if np.isfinite(roll_stats.get('r2var_median', np.nan)):
            kv('R² (variance only)\ndrift not charged', f"median "
               f"{roll_stats['r2var_median']:+.2f} · worst "
               f"{roll_stats['r2var_worst']:+.2f} · latest "
               f"{roll_stats['r2var_latest']:+.2f}",
               'Same windows, mean removed from both sides: purely "did '
               'hedging shrink the swing?". The twin of the TE row.')
        if roll_stats.get('n_neg'):
            _nn, _nc2 = roll_stats['n_neg'], roll_stats['n_neg_but_cut']
            kv('  negative-R² windows',
               f"{_nn} of {roll_stats['n_scored']}"
               + (f" · {_nc2} still CUT variance" if _nn else ''),
               (f"In {_nc2} of those {_nn}, the hedge genuinely did shrink the "
                f"swing — the negative score is entirely the target grinding "
                f"away from the basket, not a tracking failure. Judge those by "
                f"drift instead. The other {_nn - _nc2} are real: the hedge "
                f"made things worse."
                if _nc2 else
                f"In all {_nn} the hedge genuinely made the swing BIGGER. "
                f"Those are real failures, not an artifact of the measure."))
        if np.isfinite(roll_stats.get('drift_median', np.nan)):
            kv('  drift penalty',
               f"median {roll_stats['drift_median']:+.2f} R²",
               'The gap between the two rows above. Big here means the hedge '
               'tracks the SHAPE of the target but slides against its LEVEL — '
               'usually one leg running on its own story.')
        kv('  coverage', f"{roll_stats['n_days']} OOS rows across "
           f"{roll_stats['n_segments']} refits, "
           f"{rolling_df.index[0]:%d%b%y}→{rolling_df.index[-1]:%d%b%y}",
           'Read this block as STEADINESS, not as a clean score — the weights '
           'are honest but the recipe saw all the data. The clean score is in '
           'the table at the top of [1].')
    if decay_stats:                                            # [Fix 84]
        _dstr = ' · '.join(f"every {_iv}r {_te*100:.1f}%"
                           for _iv, _te in decay_stats['te_by_interval'])
        kv('TE vs refit interval', _dstr)
        _pen = (decay_stats['te_by_interval'][-1][1]
                / decay_stats['fresh_te'] - 1) * 100
        kv('Rebalance cadence',
           f"refit every ~{decay_stats['rec_rows']} rows · stretching to "
           f"{decay_stats['te_by_interval'][-1][0]}r "
           + (f"costs {_pen:.0f}% more TE" if _pen >= 0.5 else
              "costs nothing measurable"))
        kv('  basis', f"{decay_stats['n_rows']} rows, each scored by fresh, "
           f"one-refit-old and two-refit-old weights — paired, same days")
    if turnover_stats:
        kv('Basket turnover',
           f"consecutive refits share {turnover_stats['mean_jaccard']*100:.0f}%"
           f" of names · weight cosine {turnover_stats['mean_cosine']:+.2f} "
           f"({turnover_stats['n_refits']} refits)")

    # ── [2] COMPARISON ──────────────────────────────────────────────────────
    title("[2] COMPARISON",
          f"could one simple instrument have done the same job?")
    _etf_dir = 'short' if reg_e.coef_[0] > 0 else 'long'
    _P(f"  the simplest thing you could do instead: {_etf_dir} "
          f"${abs(reg_e.coef_[0])*100:.0f} of {disp[best_etf]} per $100 "
          f"of book")
    cmp_rows = []
    if sel_n:
        cmp_rows.append(['PCA basket — this config',
                         f"{sel_r2m:.3f} ± {sel_r2s:.3f}",
                         f"{sel_vrm*100:.1f}%", str(sel_n)])
    if not np.isnan(nested_r2_mean):
        cmp_rows.append(['PCA pipeline — nested',
                         f"{nested_r2_mean:.3f} ± {nested_r2_std:.3f}",
                         f"{nested_vr_mean*100:.1f}%",
                         str(len(nested_valid))])
    if len(etf_bench):
        _p = '/'.join(sorted({e.split(' ')[0]
                              for e in etf_bench['instrument']}))
        cmp_rows.append([f'Best single ETF [{_p}]', f"{etf_r2_mean:.3f}",
                         f"{etf_vr_mean*100:.1f}%", str(len(etf_bench))])
    if len(name_bench):
        _p = '/'.join(sorted({e.split(' ')[0]
                              for e in name_bench['instrument']}))
        cmp_rows.append([f'Best single name [{_p}]', f"{name_r2_mean:.3f}",
                         f"{name_vr_mean*100:.1f}%", str(len(name_bench))])
    if cmp_rows:
        _print_table(['HEDGE', 'OOS R²', 'VOL RED', 'WINDOWS'], cmp_rows,
                     note='All scored on the SAME unseen windows, so this is '
                          'a fair fight. If a single instrument gets close it '
                          'is usually the better trade — one leg, one borrow, '
                          'nothing to refit.')
    _ref_pca = sel_r2m if sel_n else nested_r2_mean
    _ref_simple = np.nanmax([etf_r2_mean, name_r2_mean])
    if not (np.isnan(_ref_pca) or np.isnan(_ref_simple)):
        _P("  ► " + ("the basket beats every single instrument on the same "
                     "unseen stretches — the extra legs are earning their keep"
                     if _ref_pca > _ref_simple else
                     "a single instrument held up at least as well — prefer "
                     "the simple hedge"))

    # ── [3] DIAGNOSTICS ─────────────────────────────────────────────────────
    title("[3] DIAGNOSTICS", "what is left after the hedge, and why")
    kvo()                     # [Fix 115] no 'DIAGNOSTICS' caption under a
    kv('Residual ADF\ndoes the leftover risk drift?',  # section is the caption
       f"p={adf_p:.4f} → "
       + ('stationary ✓ (oscillates around zero)' if adf_p < 0.05
          else 'MAY DRIFT ⚠ (can grind one way)'),
       'Drift is what pushes the rolling R² negative while TE still looks '
       'fine — the hedge tracks the SHAPE but slides against the LEVEL.')
    # ── [Fix 105] HOW MUCH OF THIS TARGET IS HEDGEABLE AT ALL ──────────────
    # The reported AAPL/TSLA/SK-Hynix case raised the right question: does
    # basket hedging even make sense for a concentrated book? A hedge built
    # from other listed instruments can only remove risk the universe SPANS.
    # Whatever is left — one name's own story — is unhedgeable by
    # construction, and it sets a CEILING on R² that no amount of config
    # search can beat. Reporting that ceiling turns "the hedge is bad" into
    # "the hedge is at 0.62 of a possible 0.68 — the gap is not the model".
    #
    # The ceiling is estimated as the in-sample R² of the target on ALL
    # principal components of the candidate universe (not just the k the
    # basket uses). That is generous — it is fitted, so it is an upper bound
    # in the strict sense — which is the right side to err on for a ceiling.
    try:
        _uni_cols = [c for c in returns_final.columns
                     if c != target_ticker and c not in set(pf_legs)]
        if len(_uni_cols) >= 5:
            _Xc = returns_final[_uni_cols].values
            _Xc = (_Xc - _Xc.mean(0)) / np.where(_Xc.std(0) < 1e-12, 1.0,
                                                 _Xc.std(0))
            _npc_ceil = int(min(len(_uni_cols), max(3, len(rets) // 20), 20))
            # _pca_fit returns a _PCAShim (components only, no .transform),
            # so project by hand: scores = X · Wᵀ
            _shim = _pca_fit(_Xc, _npc_ceil, robust_cov)
            _pcs = _Xc @ np.asarray(_shim.components_).T
            _yc = returns_final[target_ticker].values
            _reg_ceil = RidgeCV(alphas=RIDGE_ALPHAS).fit(_pcs, _yc)
            _ceil = float(max(0.0, _reg_ceil.score(_pcs, _yc)))
            _got = float(ticket_info['quality_r2']) if np.isfinite(
                ticket_info.get('quality_r2', np.nan)) else np.nan
            _idio = 1.0 - _ceil
            _msg = (f"About {_idio*100:.0f}% of what {target_short} does is "
                    f"its own story — nothing among the "
                    f"{len(_uni_cols)} candidates moves with it, at any "
                    f"weighting. That part cannot be hedged by anyone. ")
            if np.isfinite(_got):
                _msg += (f"This hedge got {_got*100:.0f} of the "
                         f"{_ceil*100:.0f} points available, so it captured "
                         f"{_got/_ceil*100 if _ceil > 0.05 else 0:.0f}% of "
                         f"what was there to take. ")
                _msg += ('The rest of the gap is structural, not a modelling '
                         'failure — no recipe can close it.'
                         if _got > 0.75 * _ceil else
                         'There is real room left — the alternatives in [4] '
                         'are worth comparing.')
            if _ceil < 0.45:
                _msg += (f" A ceiling this low is the honest verdict that this "
                         f"position is mostly a single-name bet. Expect the "
                         f"tracking score to jump around, and size off how far "
                         f"the hedge drifts instead.")
            kv('Hedgeable ceiling\nR² no config can beat here',
               f"R² ≤ ~{_ceil:.2f} · idio {_idio*100:.0f}%", _msg)
    except Exception as _e105:
        kv('Hedgeable ceiling', 'unavailable',              # [Fix 114]
           f"could not be estimated ({type(_e105).__name__}: {_e105}). "
           f"Without it there is no reading on how much of this target is "
           f"hedgeable at all — judge the hedge by how far it drifts, not by "
           f"the share of moves it tracks.")
    if is_portfolio:
        hedge_ret = returns_final[basket_final].values.dot(
            weights_final.values)
        hvar = float(np.var(hedge_ret))
        leg_rows = []
        for leg in [c for c in pf_legs if c in returns_final.columns]:
            lr = returns_final[leg].values
            b_h = (float(np.cov(lr, hedge_ret)[0, 1] / hvar)
                   if hvar > 1e-12 else np.nan)
            rc = float(np.corrcoef(resid_final, lr)[0, 1])
            leg_rows.append([disp.get(leg, leg),
                             format(pf_w.get(leg, 0.0), '+.2f'),
                             format(np.std(lr) * ann_sqrt * 100, '.1f') + '%',
                             format(b_h, '+.2f'), format(rc, '+.2f')])
        _P()
        _P("  PER-LEG — what the net-book hedge hides:")
        _print_table(['LEG', 'BOOK W', 'ANN VOL', 'β vs HEDGE',
                      'CORR w/ RESID'], leg_rows)
        _P("  ► the basket hedges the NET book; the leg with the largest "
              "|corr w/ resid| is the")
        _P("    one whose idio risk survives — hedge it separately, trim, "
              "or pair it off")
    _stock_cols = [c for c in returns_final.columns
                   if c != target_ticker and c not in set(pf_legs)
                   and c.split(' ')[0] not in ETF_SHORT_SET]      # [Fix 47]
    _caps_map = {c: _cap_asof(c, returns_final.index[-1])
                 for c in _stock_cols}
    fl = _factor_leakage(returns_final, target_ticker, weights_final,
                         caps_map=_caps_map, stock_cols=_stock_cols)
    if fl is None:
        _P("  (factor diagnostics need ~6+ non-ETF names in the "
              "universe)")
    else:
        _print_table(['FACTOR', 'TARGET β', 'RESIDUAL β'],
                     [[r['factor'], format(r['target_beta'], '+.3f'),
                       format(r['residual_beta'], '+.3f')]
                      for _, r in fl.iterrows()],
                     title='FACTOR EXPOSURE — target vs hedged residual',
                     note='Factors are sequentially orthogonalized. Residual '
                          'β ≈ 0 means the hedge neutralized that factor; a '
                          'large one is a tilt you are still carrying.')
        _lo = fl.loc[fl['residual_beta'].abs().idxmax()]
        _P(f"  ► largest surviving tilt: {_lo['factor']} "
              f"(residual β {_lo['residual_beta']:+.3f})")

    # ── [4] ALTERNATIVES ────────────────────────────────────────────────────
    # [Fix 109] Every alternative now carries the SAME evidence as the pick:
    # exact executable weights, unseen-window OOS, rolling TE / R² through
    # time, the recency split, and gross. Previously an alternative showed a
    # CV average and a truncated holdings string, so there was no way to tell
    # whether #2 was genuinely comparable or merely close on one statistic —
    # which is the only thing this section exists to answer.
    title("[4] ALTERNATIVES",
          "the next-best recipes, judged on exactly the same evidence")
    top5 = summary.sort_values('composite', ascending=False).head(5)
    _alt_cache = []
    for _i5, (_, cfg) in enumerate(top5.iterrows(), 1):
        _lb5, _ps5 = int(cfg['lookback']), cfg['pc_spec']
        _bs5 = int(cfg['basket_size'])
        _is_sel = (_lb5 == best_lookback and _ps5 == best_pc_spec
                   and _bs5 == best_bsize)
        _c_r2m, _c_r2s, _c_vrm, _c_n = _config_outer_oos(_lb5, _ps5, _bs5)
        _rs5 = (roll_stats if (_is_sel and roll_stats)
                else (_roll_stats_for(_lb5, _ps5, _bs5) if rolling_oos
                      else None))
        try:
            (_, _, _bkt5, _reg5, _r25, _, _, _, _) = _select_and_fit(
                rets.iloc[-_lb5:], _ps5, _bs5)
            _w5 = pd.Series(_reg5.coef_, index=_bkt5)
        except RuntimeError as _e5:
            _w5, _r25 = None, np.nan
        _alt_cache.append(dict(i=_i5, cfg=cfg, lb=_lb5, ps=_ps5, bs=_bs5,
                               sel=_is_sel, r2m=_c_r2m, r2s=_c_r2s,
                               vrm=_c_vrm, n=_c_n, roll=_rs5, w=_w5,
                               r2_in=_r25))
    # ── summary table: one row per config, every headline side by side ─────
    _hd = ['', 'CONFIG', 'CV R²', 'UNSEEN R²', 'ROLL TE med/worst',
           'ROLL R² med', 'RECENT R²', 'VOL RED', 'GROSS', '']
    _rows5 = []
    for _a in _alt_cache:
        _rr = _a['roll']
        _rows5.append([
            f"#{_a['i']}",
            f"lb {_a['lb']} · pc {_a['ps']} · k {_a['bs']}",
            f"{_a['cfg']['avg_r2_oos']:.3f}",
            (f"{_a['r2m']:.3f} ± {_a['r2s']:.3f}" if _a['n'] else 'n/a'),
            (f"{_rr['te_median']*100:.1f}% / {_rr['te_worst']*100:.1f}%"
             if _rr else 'n/a'),
            (f"{_rr['r2_median']:+.2f}" if _rr else 'n/a'),
            (f"{_a['cfg']['recent_r2']:.3f}"
             if np.isfinite(_a['cfg'].get('recent_r2', np.nan)) else 'n/a'),
            f"{_a['cfg']['avg_vol_red']*100:.0f}%",
            f"${_a['cfg']['avg_gross']*100:.0f}",
            ('◄ chosen' if _a['sel'] else '')])
    _print_table(_hd, _rows5,
                 note='Compare on UNSEEN R² and ROLL TE — never on CV R², '
                      'which is what the selection already optimised and is '
                      'therefore guaranteed to favour the pick. Config = '
                      'lookback days · PCs · basket size.')
    # ── per-config detail: the exact ticket each alternative would produce ──
    # [Fix 115] This used to repeat the SAME four explanation rows under every
    # config — five configs × four paragraphs of identical prose, ~20 rows of
    # boilerplate for four distinct sentences. The explanations are a property
    # of the COLUMNS, not of any one config, so they are stated once here and
    # each config below shows only its own numbers.
    _P("  ► each block below is the ticket that recipe would produce, and "
       "how it scored on exactly the same evidence as the one chosen")
    _kv_open('WHAT THE FOUR MEASURES BELOW MEAN')  # stated once, not per row
    kv('gross / in-sample', 'what you would put on, and its ceiling',
       'The second figure is the fit scored on its own data. Never size off it.')
    kv('unseen windows', 'refit before each window, scored on it',
       'The honest comparison between configs.')
    kv('rolling OOS', '21-row walk-forward TE and R²',
       'Read stability here — a config with a similar median but a much worse '
       'worst case is the more dangerous one.')
    kv('recency', 'all folds vs the newest folds',
       'Recent above all-fold means improving; the reverse means decaying.')
    _kv_flush()
    for _a in _alt_cache:
        _tagsel = '   ◄ THIS IS THE ONE IN THE TICKET' if _a['sel'] else ''
        if _a['w'] is None:
            _P(f"  #{_a['i']}  fits on {_a['lb']} days, {_a['bs']} names "
               f"— refit failed, no weights available")
            continue
        _P(f"  #{_a['i']}   fits on {_a['lb']} days, {_a['bs']} names in "
           f"the basket{_tagsel}")
        _print_table(BASKET_HEADERS, _basket_rows(_a['w']), indent='      ')
        _rr = _a['roll']
        _det = [('gross / in-sample',
                 f"${_a['w'].abs().sum()*100:.0f} per $100 · in-sample R² "
                 f"{_a['r2_in']:.2f}")]
        if _a['n']:
            _det.append(('unseen windows',
                         f"R² {_a['r2m']:.3f} ± {_a['r2s']:.3f} · vol red "
                         f"{_a['vrm']*100:.0f}% · {_a['n']} windows"))
        if _rr:
            _det.append((
                'rolling OOS',
                f"TE {_rr['te_median']*100:.1f}% med / "
                f"{_rr['te_worst']*100:.1f}% worst / "
                f"{_rr['te_latest']*100:.1f}% latest · R² "
                f"{_rr['r2_median']:+.2f} med / {_rr['r2_worst']:+.2f} worst "
                f"/ {_rr['r2_latest']:+.2f} latest · "
                + (f"{_rr['n_neg']} of {_rr.get('n_scored', 0)} windows "
                   f"NEGATIVE" if _rr.get('n_neg') else
                   f"0 of {_rr.get('n_scored', 0)} windows negative")))
        if np.isfinite(_a['cfg'].get('recent_r2', np.nan)):
            _det.append((
                'recency',
                f"all folds {_a['cfg']['avg_r2_oos']:.3f} · recent "
                f"{_a['cfg']['recent_r2']:.3f}"
                + (f" · slope {-_a['cfg']['r2_slope_per_fold']:+.3f}/fold"
                   if np.isfinite(_a['cfg'].get('r2_slope_per_fold', np.nan))
                   else '')))
        _kv_block(None, _det, headers=('MEASURE', 'THIS CONFIG'))

    # ── [A] APPENDIX ────────────────────────────────────────────────────────
    title("[A] APPENDIX", "settings, definitions and caveats")
    # ── [Fix 110] WHICH KNOBS ACTUALLY MOVED ANYTHING ──────────────────────
    # Two questions kept coming back — "does the hedge duration change the
    # analysis?" and "does the book size?" — and the report gave no way to
    # answer either without reading the source. Every setting that can change
    # a NUMBER is now listed with what it touched, and, just as importantly,
    # the ones that touch nothing are listed as touching nothing.
    _rows_a = []
    _rows_a.append((
        'Book size\nnotional_mm',
        f"${notional_mm:g}mm" if notional_mm else 'not set',
        'DISPLAY ONLY — it does not affect the model. A hedge ratio is a '
        'ratio: the weights, the recipe search, the testing and every score '
        'are identical at $1mm and $100mm. All this does is turn them into '
        'dollars. (What book size DOES change in real life — market impact '
        'and whether you can borrow the stock — is not modelled here at all; '
        'see the notes at the end.)'))
    _rows_a.append((
        'Hedge duration\nhedge_horizon_days',
        f"{int(hedge_horizon_days)} trading days" if hedge_horizon_days
        else 'not set',
        ('DOES affect the report, NOT the weights. It scales the risk to '
         'your holding period, reports what holds of that length actually '
         'delivered, and checks whether the weights go stale before the hold '
         'ends. It deliberately does NOT change which stretches are held back '
         'for testing — that was measured and found to buy nothing while '
         'costing evidence. horizon_sets_windows=True restores the old '
         'behaviour and WILL change which recipe wins.'
         if hedge_horizon_days else
         'Not set, so risk is quoted annualised and the "holds" evidence is '
         'skipped. Set it to the number of trading days you intend to hold '
         '— it costs nothing and adds the only figures measured on your '
         'actual horizon.')))
    # ── [Fix 112] HOW MUCH OF THE HISTORY IS ACTUALLY READ ─────────────────
    # The walk-forward consumes a FIXED budget of rows, set by the fold
    # counts and the longest lookback — not by how much data you downloaded:
    #     max(lookback) + cv_test_window x cv_n_folds
    #                   + outer_window x outer_folds + purge
    # Everything older than that is never touched. Measured directly
    # (synthetic, 10 seeds, default settings): 2.5y / 3y / 4y / 5y of history
    # produce BIT-IDENTICAL results. So "download more history" is not a way
    # to stabilise the selection — it does nothing at all. Raising the fold
    # counts IS the way to consume more, and measured that did not help
    # either (cv_n_folds 5->10 and outer_folds 3->6 at 5y: paired
    # d = -0.011 nested R², t = -1.29, better in 4 of 10 — the extra folds
    # reach back into older, less relevant regimes).
    _need = (max(lookback_options) + cv_test_window * cv_n_folds
             + outer_window * outer_folds + purge * 2)
    _spare = len(rets) - _need
    _rows_a.append((
        'How much history was read\nof what you downloaded',
        f"≈{min(_need, len(rets))} of {len(rets)} rows"
        + (f" · {_spare} unused" if _spare > 0 else ''),
        (f"The SELECTION and the VALIDATION read about {_need} rows and "
         f"stop; the oldest {_spare} ({_spare/252:.1f}y) never enter either. "
         f"Measured: 2.5 and 5 years of history give identical picks and "
         f"identical scores at these settings. So downloading more history "
         f"does not steady the choice — only testing on more stretches "
         f"consumes more, and that was measured not to help either (it "
         f"reaches back into older, less relevant markets). Longer history IS "
         f"worth having for one thing: it allows a longer lookback, which "
         f"needs its own training run before every test stretch."
         if _spare > 20 else
         f"TIGHT — the testing wants about {_need} days and has "
         f"{len(rets)}. Test stretches get dropped when there is not enough "
         f"history in front of them to train on, which is why the count can "
         f"come back below {cv_n_folds}. Here — and only here — more history "
         f"genuinely helps.")))
    _rows_a.append((
        'Basket sizes tried\nnames per basket', str(tuple(basket_size_options)),
        'The grid the selection searched. 3/5/7 is a deliberate '
        'bias-variance compromise: fewer names means less overfitting but '
        'less spanning; more means the ridge starts pairing off '
        'near-duplicates and gross climbs. k=1 is covered separately by [2], '
        'which tests the single best ETF and the single best name on the '
        'same unseen windows.'))
    # ── [Fix 113] IS THERE ROOM FOR A LONGER LOOKBACK? ─────────────────────
    # A lookback is only worth offering if the NESTED validation can still
    # train it — the oldest outer window's inner CV needs
    #     lb + cv_test_window*cv_n_folds  rows of data before it.
    # Below that, the option is either dropped or (worse) available to the
    # live pick but not to the validation, so you could select something you
    # never validated.
    # Measured (synthetic, 20 paired seeds, nested pooled OOS R²):
    #     adding 504 with 5y of data   concentrated +0.0146 (t=+3.19)
    #                                  well-spanned  +0.0031 (t=+2.28)
    #     adding 504 with 3y of data   504 is chosen 0/20 times — no effect,
    #                                  because there is not enough history to
    #                                  train it before each window.
    # So the default (126, 252) is right for the default 3-year download, and
    # a longer lookback is worth adding ONLY together with more history. That
    # is the one thing extra history genuinely buys — see 'History READ'.
    _lb_room = (len(rets) - outer_folds * outer_window - purge
                - cv_test_window * cv_n_folds)
    _lb_hint = ''
    if 504 > max(lookback_options):
        if 504 <= _lb_room:
            _lb_hint = (" You have enough history to VALIDATE a 504-row (2y) "
                        "lookback and are not trying one. Measured on 20 "
                        "paired seeds, adding it moved nested OOS R² by "
                        "+0.015 (t=+3.2) on a concentrated target and +0.003 "
                        "(t=+2.3) on a well-spanned one. Worth an A/B: "
                        "lookback_options=("
                        + ', '.join(str(x) for x in
                                    list(lookback_options) + [504]) + ").")
        else:
            _lb_hint = (f" A 2-year (504-row) lookback would need about "
                        f"{504 - _lb_room} more rows of history before it "
                        f"could be validated, so it is correctly not offered "
                        f"here — measured at 3 years of data it is chosen 0 "
                        f"times out of 20 and buys nothing.")
    _rows_a.append((
        'Lookbacks tried\ndays of history per refit', str(tuple(lookback_options)),
        # [Fix 115] this row was a 120-word essay. The measured finding stays,
        # because it answers the question this row actually gets asked; the
        # recital of the experiment does not.
        'How much history each refit sees. Short adapts faster and is noisier; '
        'long is steadier and staler. The right lookback is set by how STABLE '
        'the relationships are, NOT by how long you intend to hold — measured '
        'across four kinds of market, the winner was identical at a 10-day and '
        'a 120-day hold, but dropped from 504 days to 126 when the '
        'relationships broke. If [1] says the recipe is fading, a shorter '
        'lookback is the first thing to try.' + _lb_hint))
    _rows_a.append((
        'Outlier clipping\nwinsorize', f"{winsorize_pct:.1%}" if winsorize_pct else 'off',
        'Clips returns before the PCA that SELECTS candidates — never '
        'before the ridge that sets the hedge ratios, so no fitted number '
        'is computed on doctored data. 1% keeps a handful of observations '
        'per tail from steering the eigenvectors. Raising it much further '
        'starts deleting genuine crisis co-movement, which is exactly the '
        'behaviour a hedge exists to capture.'))
    _rows_a.append((
        'Penalty on a big position\ngross_penalty',
        f"{gross_penalty:g} per 1.0x gross"
        + (f" · hard cap {max_gross:g}x" if max_gross else ' · no hard cap'),
        f"Costs a recipe {gross_penalty*100:.0f} points in the ranking per "
        f"$100 of position. This chosen basket runs "
        f"${weights_final.abs().sum()*100:.0f} per $100, so it paid "
        f"{gross_penalty*weights_final.abs().sum():.3f}. The penalty is a "
        f"proxy, not a cost model — borrow, financing and impact are NOT "
        f"priced anywhere in this tool. If gross matters to you "
        f"economically, set max_gross (3.0 is a common desk limit) rather "
        f"than relying on the penalty to hold it down."))
    _rows_a.append((
        'How much recent data counts extra\nrecency_weight',
        f"{recency_weight:g}" + (' (ACTIVE)' if recency_weight else
                                 ' — every test stretch counts the same'),
        'At 0 the search treats a test stretch from three years ago exactly '
        'like one from last month. See "IS THIS STILL THE RIGHT RECIPE '
        'TODAY?" in [1] for what that cost on this run, and what choosing on '
        'recent data alone would have picked instead.'))
    _kv_block('SETTINGS — what each knob actually did', _rows_a,
              headers=('SETTING', 'THIS RUN', 'WHAT IT AFFECTS'),
              note='Reference material — nothing here changes unless you '
                   'change it. The rows that say DISPLAY ONLY genuinely do '
                   'not move a single number in the model.')
    _kv_block('WORDS USED IN THIS REPORT', WORDS_USED,   # [Fix 116]
              headers=('TERM', 'IN ONE LINE', 'A BIT MORE'))
    _kv_block('FOUR WAYS OF SCORING A HEDGE — and why they differ',
              R2_GLOSSARY_ROWS, note=R2_GLOSSARY_NOTE,   # [Fix 52]
              headers=('HOW IT WAS MEASURED', 'WHAT IT IS WORTH',
                       'WHAT IT ACTUALLY TELLS YOU'))
    hr()
    _P("  Read before trading:")
    _P("  · the candidate list comes from today's index members and peers"
          + (" — but membership is checked as of each date, so delisted names "
             "are included where they belong ✓" if pm_active else
             ". Names that have since delisted are missing, which flatters "
             "every backtest here"))
    _P("  · the size floor uses "
          + ("what each name was worth at the time ✓" if pit_active else
             "what each name is worth TODAY — a name that grew into the "
             "filter can flatter the backtest"))
    _P("  · these are statistical ratios only. Borrow cost, slippage and FX "
          "are NOT priced anywhere in this tool, and returns are in each "
          "name's own currency")
    if is_portfolio:
        _P("  · your book is hedged as ONE combined position, not leg by "
              "leg. The names you hold are excluded from the candidate list")
    if forced:
        _P("  · ★ names were pinned in by you, not chosen by the model — "
              "run it once without them to see what they cost")
    _P('═' * L)

    return {
        'weights': weights_final, 'basket': basket_final,
        'ranking': final_rank, 'cv_summary': summary, 'cv_detail': cv_df,
        'nested_validation': nested,
        'nested_r2_mean': nested_r2_mean, 'nested_r2_std': nested_r2_std,
        'nested_vol_red_mean': nested_vr_mean,
        'selected_unseen_r2': sel_r2m, 'selected_unseen_r2_std': sel_r2s,
        'selected_unseen_pooled_r2': sel_pooled,       # [Fix 65]
        'outer_window_detail': sel_detail,             # [Fix 65]
        'nested_pooled_r2': nested_pooled,             # [Fix 69]
        'bootstrap_ci': boot,                          # [Fix 70]
        'rolling_oos': rolling_df,                     # [Fix 71]
        'ticket': ticket_info,                         # [Fix 72]
        'nowcast': nowcast,                            # [Fix 93]
        'target_label': target_short,                  # [Fix 103] for export
        'analysis_span': (rets.index[0], rets.index[-1], len(rets)),  # [Fix 99]
        'horizon': horizon_info,                       # [Fix 91]
        'validation_mode': val_mode,                   # [Fix 77]
        'decay': decay_stats,                          # [Fix 78]
        'turnover': turnover_stats,                    # [Fix 78]
        'etf_benchmark': etf_bench,
        'etf_nested_r2_mean': etf_r2_mean,
        'name_benchmark': name_bench,                  # [Fix 53]
        'name_nested_r2_mean': name_r2_mean,           # [Fix 53]
        'best_single': best_single, 'best_etf': best_etf,
        'returns': returns_final, 'prices': prices,
        'two_day_returns': two_day,
        'excluded': removed_cols,                      # [Fix 36]
        'force_included': forced,                      # [Fix 40]
        'is_portfolio': is_portfolio,                  # [Fix 46]
        'portfolio_legs': pf_legs, 'portfolio_weights': pf_w,
        'factor_leakage': fl,                          # [Fix 43]
        'pit_size_active': pit_active,                 # [Fix 44]
        'pit_members_active': pm_active,               # [Fix 49]
        'params': {'lookback': best_lookback, 'pc_spec': best_pc_spec,
                   'basket_size': best_bsize,
                   'tz_tolerance_hours': tz_tolerance_hours,
                   'min_mktcap_usd_mm': min_mktcap_usd_mm,
                   'gross_penalty': gross_penalty, 'max_gross': max_gross,
                   'force_include': force_include,
                   'winsorize_pct': winsorize_pct, 'robust_cov': robust_cov,
                   'pit_size': pit_size, 'pit_members': pit_members,
                   'size_band': size_band,
                   'size_override': size_override,
                   'mask_stale_days': mask_stale_days,    # [Fix 56]
                   'notional_mm': notional_mm,            # [Fix 72]
                   'boot_n': boot_n,                      # [Fix 70]
                   'sample_window': sample_window,        # [Fix 76]
                   'hedge_horizon_days': hedge_horizon_days,  # [Fix 91]
                   'horizon_sets_windows': horizon_sets_windows,  # [Fix 100]
                   'report': report},                     # [Fix 79]
    }




# ── [Fix 63] FAST RE-HEDGE — for books that change daily ────────────────────
# An index-rebalance book churns as predictions update, and re-running the
# full grid plus nested validation for every intraday change is overkill: the
# expensive part is CONFIG SELECTION, and the config does not change from one
# day to the next. quick_rehedge pins the config that was already validated
# and refits the weights only — seconds, and zero Bloomberg calls. Re-run the
# full find_best_hedge weekly, or whenever the book changes materially.
def quick_rehedge(book, data, prev_result, show_plots=True, rolling_oos=True,
                  full_validation=False):
    """Refit hedge weights for an updated book using a previously VALIDATED
    config. `prev_result` is the dict returned by a full find_best_hedge run.

    The config is NOT re-searched — that is the point of this path, and it is
    what makes it seconds instead of minutes. What you can choose:

      show_plots       draw the charts (default True)
      rolling_oos      compute the rolling walk-forward curve (default True)
      full_validation  re-run the same number of validation windows as a full
                       run (3 outer × 5 inner) instead of the fast 1 × 2

    [Fix 115] show_plots and rolling_oos used to be hard-wired OFF here, so a
    re-hedge silently dropped every chart the full run had drawn and reported
    a single outer window with no explanation of why. Both are now the
    caller's choice, defaulting ON, and the report says in [0] exactly what
    was and was not re-validated.
    """
    p = prev_result['params']
    return find_best_hedge(
        book, data=data,
        lookback_options=(p['lookback'],),
        pc_options='auto' if p['pc_spec'] == 'auto' else [int(p['pc_spec'])],
        basket_size_options=(p['basket_size'],),
        min_mktcap_usd_mm=p['min_mktcap_usd_mm'],
        tz_tolerance_hours=p['tz_tolerance_hours'],
        gross_penalty=p['gross_penalty'], max_gross=p['max_gross'],
        force_include=p['force_include'],
        winsorize_pct=p['winsorize_pct'], robust_cov=p['robust_cov'],
        pit_size=p['pit_size'], pit_members=p['pit_members'],
        size_band=p['size_band'], size_override=p['size_override'],
        mask_stale_days=p.get('mask_stale_days', 7),
        notional_mm=p.get('notional_mm'),              # [Fix 72]
        boot_n=p.get('boot_n', 2000),                  # [Fix 70]
        rolling_oos=rolling_oos,                       # [Fix 71]
        sample_window=p.get('sample_window'),          # [Fix 76]
        hedge_horizon_days=p.get('hedge_horizon_days'),  # [Fix 91]
        horizon_sets_windows=p.get('horizon_sets_windows', False),
        report=p.get('report', 'compact'),             # [Fix 79]
        cv_n_folds=5 if full_validation else 2,
        outer_folds=3 if full_validation else 1,
        show_plots=show_plots,
        _rehedge_of=prev_result)                       # [Fix 115] for the note


# =============================================================================
# CELL 3: FRONT END — [Fix 89/95/96] hedge_ui()
# =============================================================================
# A point-and-click form covering EVERY engine option, so the tool is driven
# without reading or editing code. It is a THIN layer: each button calls the
# same download_hedge_data / find_best_hedge / quick_rehedge a code user
# would call, with the same arguments — the model, validation and report are
# byte-for-byte identical to the scripted path. The form only assembles
# arguments; it never re-implements logic.
#
# What it gives you beyond the scripted path:
#   · live preview of the validation tier for the chosen window BEFORE a run
#     (engine's own _validation_tier, [Fix 85] — cannot drift);
#   · cached downloads with a live fit check ([Fix 88/90]) — it refuses to
#     run when a leg was never downloaded, and flags stale prices;
#   · one-click EXPORT of weights (CSV), the full report (TXT) and the run
#     settings (JSON) under hedge_reports/;
#   · RE-HEDGE (fast): quick_rehedge pinned to the last validated config;
#   · "Show code": the exact scripted equivalent of any run, so nothing done
#     in the GUI is unreproducible;
#   · results always in UI_STATE ('result', 'report_text', 'code').
#
# If ipywidgets is missing it prints the install command and a code template
# instead of failing.
# =============================================================================

UI_STATE = {'data': None, 'result': None, 'code': '', 'report_text': '',
            'loaded_from': None}

_UI_CSS = """
<style>
/* [Fix 96/102] typography + spacing. ipywidgets defaults are cramped and
   mix fonts; this gives one type scale, roomy rows and grouped sections. */
.hedge-ui, .hedge-ui .widget-label, .hedge-ui input, .hedge-ui select,
.hedge-ui textarea, .hedge-ui .widget-button, .hedge-ui .widget-dropdown {
  font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Inter',
               'Helvetica Neue', Arial, sans-serif !important;
  -webkit-font-smoothing: antialiased;
}
.hedge-ui { color: #1d2b3a; line-height: 1.45; }
.hedge-ui .widget-label {
  font-size: 12.5px !important; color: #33506d !important;
  font-weight: 600 !important; letter-spacing: .01em;
}
.hedge-ui input, .hedge-ui select, .hedge-ui textarea {
  font-family: 'SF Mono', Menlo, Consolas, monospace !important;
  font-size: 12.5px !important; color: #16222e !important;
  border: 1px solid #cfd8e3 !important; border-radius: 5px !important;
  padding: 3px 7px !important; background: #fdfefe !important;
}
.hedge-ui input:focus, .hedge-ui textarea:focus, .hedge-ui select:focus {
  border-color: #4a7fb5 !important; outline: none !important;
  box-shadow: 0 0 0 2px rgba(74,127,181,.15) !important;
}
.hedge-ui .widget-hbox, .hedge-ui .widget-vbox { margin: 0 !important; }
.hedge-ui .widget-text, .hedge-ui .widget-dropdown,
.hedge-ui .widget-textarea, .hedge-ui .widget-checkbox,
.hedge-ui .widget-datepicker { margin: 3px 0 !important; }
.hedge-ui .widget-button {
  font-weight: 600 !important; font-size: 12.5px !important;
  border-radius: 5px !important; letter-spacing: .01em;
  margin: 0 8px 0 0 !important;
}
.hedge-ui .p-TabBar-tabLabel, .hedge-ui .lm-TabBar-tabLabel {
  font-size: 13px; font-weight: 650; letter-spacing: .02em;
}
.hedge-ui .p-TabPanel-stackedPanel, .hedge-ui .lm-TabPanel-stackedPanel {
  padding: 6px 10px 10px 10px;
}
.hedge-ui .hui-h {
  font-size: 11px; font-weight: 700; color: #33506d;
  border-bottom: 1px solid #e2e8f0; margin: 16px 0 8px 0;
  padding-bottom: 4px; text-transform: uppercase; letter-spacing: .07em;
}
.hedge-ui .hui-h:first-child { margin-top: 4px; }
.hedge-ui .hui-help {
  font-size: 11.5px; color: #6b7a8c; margin: 0 0 8px 2px;
  line-height: 1.45; max-width: 560px;
}
.hedge-ui .hui-title {
  font-size: 18px; font-weight: 700; color: #16324b;
  letter-spacing: -.01em; margin: 2px 0 1px 0;
}
.hedge-ui .hui-sub {
  font-size: 12px; color: #6b7a8c; margin-bottom: 10px;
}
.hedge-ui .widget-output { margin-top: 8px; }
</style>"""

def _ui_parse_book(text):
    """Parse the target box into a str / list / dict for _coerce_targets.

    One entry per line or comma-separated. 'TICKER' or 'TICKER, weight'.
    A single line with no weight → single-name hedge. Several lines with no
    weights → equal-weight long book. Any weight present → weighted book
    (missing weights default to 1.0)."""
    entries = []
    for raw in str(text).replace(';', '\n').splitlines():
        raw = raw.strip()
        if not raw:
            continue
        if ',' in raw:
            parts = [p.strip() for p in raw.split(',')]
            if len(parts) >= 2 and _ui_isnum(parts[-1]):
                entries.append((' '.join(parts[:-1]).strip(),
                                float(parts[-1])))
                continue
            for p in parts:                      # 'AAPL US, MSFT US'
                if p:
                    entries.append((p, None))
            continue
        bits = raw.split()
        if len(bits) >= 2 and _ui_isnum(bits[-1]):
            entries.append((' '.join(bits[:-1]), float(bits[-1])))
        else:
            entries.append((raw, None))
    if not entries:
        raise ValueError("no ticker entered")
    if len(entries) == 1 and entries[0][1] is None:
        return entries[0][0]
    if all(w is None for _, w in entries):
        return [t for t, _ in entries]
    return {t: (1.0 if w is None else w) for t, w in entries}


def _ui_isnum(s):
    try:
        float(s)
        return True
    except (TypeError, ValueError):
        return False


def _ui_ints(text, default):
    """'3,5,7' → (3, 5, 7); blank → default."""
    vals = [int(float(v)) for v in str(text).replace(' ', '').split(',')
            if v not in ('', None)]
    return tuple(vals) if vals else tuple(default)


def _ui_list(text):
    """'2330 TT, 0050 TT' → ['2330 TT', '0050 TT']; blank → None."""
    vals = [v.strip() for v in str(text).replace('\n', ',').split(',')
            if v.strip()]
    return vals or None


def _ui_code(target, window, kw, loaded_from=None):
    """Render the equivalent Python so any GUI run can be reproduced."""
    def fmt(v):
        return f"'{v}'" if isinstance(v, str) else repr(v)
    lines = ["# equivalent code for this run:"]
    if loaded_from:
        lines.append(f"data = load_hedge_data({fmt(loaded_from)})")
    else:
        lines.append(f"data = download_hedge_data({fmt(target)},"
                     f" min_mktcap_usd_mm={kw.get('min_mktcap_usd_mm', 5000)})")
    args = [f"    {k}={fmt(v)}," for k, v in sorted(kw.items())
            if v is not None]
    if window is not None:
        args.insert(0, f"    sample_window={fmt(window)},")
    lines.append(f"result = find_best_hedge({fmt(target)}, data=data,")
    lines += args
    lines.append(")")
    return '\n'.join(lines)


def _export_folder(out_dir=None):
    """[Fix 103] Resolve the export folder, creating it if possible. Falls
    back to ./hedge_reports when the configured drive is not mapped on this
    machine (e.g. running off-network) — never fails the export."""
    import os as _o
    import re as _r
    want = out_dir or HEDGE_EXPORT_DIR
    # A Windows path such as G:\FIN_COMM\... is a perfectly LEGAL single
    # filename on macOS/Linux, so makedirs would happily create a junk
    # folder whose name contains backslashes instead of failing. Detect the
    # drive-letter form and refuse it off Windows.
    if _o.name != 'nt' and _r.match(r'^[A-Za-z]:[\\/]', want):
        fb = _o.path.join(_o.getcwd(), 'hedge_reports')
        _o.makedirs(fb, exist_ok=True)
        return fb, (f"{want} is a Windows path and this machine is not "
                    f"Windows — wrote to {fb} instead")
    try:
        _o.makedirs(want, exist_ok=True)
        if _o.path.isdir(want):
            return want, None
    except Exception as e:
        err = str(e)
    else:
        err = 'not a directory'
    fb = _o.path.join(_o.getcwd(), 'hedge_reports')
    _o.makedirs(fb, exist_ok=True)
    return fb, (f"{want} is not reachable ({err}) — wrote to {fb} instead; "
                f"map the drive or set HEDGE_EXPORT_DIR")


def _docx_hedge_report(result, report_text, path):
    """[Fix 103] Write the hedge as a Word document a desk can file or send:
    verdict, the trade as a real table, risk and monitoring, the evidence,
    the exact settings, and the full text report as an appendix."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY = RGBColor(0x16, 0x32, 0x4B)
    SLATE = RGBColor(0x55, 0x66, 0x77)
    COL = {'✓': RGBColor(0x1B, 0x6B, 0x3A), '△': RGBColor(0x8A, 0x5A, 0x00),
           '✗': RGBColor(0x8C, 0x1C, 0x1C), '?': RGBColor(0x3D, 0x46, 0x50)}

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)
    base = doc.styles['Normal']
    base.font.name = 'Calibri'
    base.font.size = Pt(10.5)
    base.paragraph_format.space_after = Pt(5)

    def H(txt, size=15, before=14):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(txt)
        r.font.size = Pt(size); r.font.bold = True; r.font.color.rgb = NAVY
        return p

    def KV(label, value, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.12)
        r = p.add_run(f"{label}   ")
        r.font.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = color or NAVY
        r2 = p.add_run(str(value)); r2.font.size = Pt(10)

    tk = result.get('ticket') or {}
    label = result.get('target_label') or 'hedge'
    params = result.get('params') or {}
    notional = params.get('notional_mm')

    # ── header ──────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"PCA Hedge — {label}")
    r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = NAVY
    span = result.get('analysis_span')
    sub = f"{datetime.now():%d %b %Y  %H:%M}"
    if span:
        sub += (f"   ·   analysed {span[0]:%d %b %Y} → {span[1]:%d %b %Y} "
                f"({span[2]} rows)")
    sub += f"   ·   {result.get('validation_mode', '')} validation"
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(sub); r.font.size = Pt(9.5); r.font.color.rgb = SLATE

    # ── verdict ─────────────────────────────────────────────────────────
    grade = tk.get('grade', 'n/a')
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(grade)
    r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = COL.get(str(grade)[:1], NAVY)
    for why in (tk.get('reasons') or [])[:4]:
        b = doc.add_paragraph(style='List Bullet')
        b.paragraph_format.left_indent = Inches(0.3)
        b.paragraph_format.space_after = Pt(1)
        b.add_run(str(why)).font.size = Pt(9.5)
    if tk.get('quality_r2') is not None and np.isfinite(tk['quality_r2']):
        boot = result.get('bootstrap_ci') or {}
        q = (f"pooled out-of-sample R² {tk['quality_r2']:.2f} "
             f"({tk.get('quality_source', '')})")
        if boot:
            q += f"   ·   95% CI [{boot['r2_lo']:.2f}, {boot['r2_hi']:.2f}]"
        if np.isfinite(tk.get('worst_window_r2', np.nan)):
            q += f"   ·   worst window {tk['worst_window_r2']:.2f}"
        KV('QUALITY', q)

    # ── the trade ───────────────────────────────────────────────────────
    H('The trade')
    w = result['weights'].sort_values(key=lambda s: -s.abs())
    cols = ['ACTION', 'PER $100', 'INSTRUMENT', 'FULL TICKER']
    if notional:
        cols.insert(2, f'${notional:g}mm')
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = 'Light Grid Accent 1'
    for i, c in enumerate(cols):
        cell = t.rows[0].cells[i]
        cell.text = c
        for run in cell.paragraphs[0].runs:
            run.font.bold = True; run.font.size = Pt(9)
    forced = set(result.get('force_included') or [])
    for tkr, val in w.items():
        row = t.add_row().cells
        vals = ['SHORT' if val > 0 else 'LONG',
                f"${abs(val)*100:,.1f}",
                tkr.split(' ')[0] + (' ★' if tkr in forced else ''),
                tkr]
        if notional:
            vals.insert(2, f"${abs(val)*notional:,.2f}mm")
        for i, v in enumerate(vals):
            row[i].text = str(v)
            for run in row[i].paragraphs[0].runs:
                run.font.size = Pt(9)
    gross, net = w.abs().sum()*100, w.sum()*100
    KV('Exposure', f"gross ${gross:,.1f} · net "
                   f"{'short' if net >= 0 else 'long'} ${abs(net):,.1f} "
                   f"per $100 of "
                   f"{'gross book' if result.get('is_portfolio') else label}")

    # ── risk & monitoring ───────────────────────────────────────────────
    H('Risk and monitoring')
    nc = result.get('nowcast') or {}
    hz = result.get('horizon') or {}
    if np.isfinite(tk.get('worst_window_te', np.nan)):
        line = f"hedged vol {tk['worst_window_te']*100:.1f}%/yr (worst " \
               f"validated window — size off this, not the average)"
        if notional:
            line += (f"  ≈ ±{_fmt_money_mm(tk['worst_window_te']*notional)}"
                     f"/yr 1σ")
        KV('SIZING', line)
    if nc:
        KV('RIGHT NOW', f"{nc.get('condition')} — "
                        + '; '.join(nc.get('why') or ['no adverse signals']))
    if hz:
        line = (f"1σ slippage over {hz['days']} trading days ≈ "
                f"±{hz['sigma_pct']*100:.2f}%")
        if notional:
            line += f" ≈ ±{_fmt_money_mm(hz['sigma_pct']*notional)}"
        KV('HORIZON', line)
        blk = hz.get('blocks')
        if blk:
            KV(f"{hz['days']}-DAY HOLDS",
               f"{blk['n']} non-overlapping holds: median TE "
               f"{blk['te_median']*100:.1f}%, worst {blk['te_worst']*100:.1f}"
               f"%/yr; risk cut in {blk['share_helped']*100:.0f}% of them")
        if hz.get('refits_needed') is not None:
            KV('REFITS', 'none needed within the horizon'
               if hz['refits_needed'] <= 0
               else f"plan ~{hz['refits_needed']} re-hedge(s) during the hold")
    if np.isfinite(tk.get('kill_switch_te', np.nan)):
        KV('MONITOR', f"refit or cut if rolling 21-row tracking error "
                      f"exceeds {tk['kill_switch_te']*100:.1f}%/yr")

    # ── evidence ────────────────────────────────────────────────────────
    det = result.get('outer_window_detail')
    if det is not None and len(det):
        H('Out-of-sample windows')
        t2 = doc.add_table(rows=1, cols=5)
        t2.style = 'Light Grid Accent 1'
        for i, c in enumerate(['WINDOW', 'DATES', 'OOS R²', 'VOL RED',
                               'TARGET VOL']):
            t2.rows[0].cells[i].text = c
            for run in t2.rows[0].cells[i].paragraphs[0].runs:
                run.font.bold = True; run.font.size = Pt(9)
        for _, rw in det.iterrows():
            cells = t2.add_row().cells
            for i, v in enumerate([
                    f"{int(rw['window'])}"
                    + (' (latest)' if rw['window'] == 1 else ''),
                    f"{rw['start']:%d %b} – {rw['end']:%d %b %y}",
                    f"{rw['r2']:.3f}", f"{rw['vol_red']*100:.1f}%",
                    f"{rw['tgt_vol']*100:.1f}%"]):
                cells[i].text = v
                for run in cells[i].paragraphs[0].runs:
                    run.font.size = Pt(9)
    bits = []
    if np.isfinite(result.get('nested_pooled_r2', np.nan)):
        bits.append(f"nested pooled OOS R² {result['nested_pooled_r2']:.3f}")
    for nm, key in (('best single ETF', 'etf_nested_r2_mean'),
                    ('best single name', 'name_nested_r2_mean')):
        v = result.get(key)
        if v is not None and np.isfinite(v):
            bits.append(f"{nm} {v:.3f}")
    if bits:
        KV('VS SIMPLE HEDGES', ' · '.join(bits))
    dec = result.get('decay') or {}
    if dec.get('rec_rows'):
        KV('REBALANCE', f"refit about every {dec['rec_rows']} rows "
                        f"(measured tracking-error decay)")

    # ── settings ────────────────────────────────────────────────────────
    H('Settings used')
    keep = ('lookback', 'pc_spec', 'basket_size', 'hedge_horizon_days',
            'sample_window', 'min_mktcap_usd_mm', 'notional_mm',
            'robust_cov', 'winsorize_pct', 'gross_penalty', 'max_gross',
            'restrict_to_target_tz', 'tz_tolerance_hours', 'mask_stale_days',
            'force_include', 'size_band')
    txt = ' · '.join(f"{k}={params[k]}" for k in keep
                     if params.get(k) is not None)
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.12)
    r = p.add_run(txt); r.font.size = Pt(8.5); r.font.name = 'Consolas'
    r.font.color.rgb = SLATE

    # ── appendix ────────────────────────────────────────────────────────
    if report_text:
        doc.add_page_break()
        H('Appendix — full report', size=13, before=0)
        for line in report_text.split('\n'):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(0); pf.space_before = Pt(0)
            pf.line_spacing = 1.0
            r = p.add_run(line)
            r.font.name = 'Consolas'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            r.font.size = Pt(7.2)
    doc.save(path)
    return path


def export_result(result, report_text='', out_dir=None, also_csv=True):
    r"""[Fix 95/103] Export a finished hedge as a WORD DOCUMENT.

    Writes PCA_hedge_<label>_<timestamp>.docx into HEDGE_EXPORT_DIR (default
    G:\FIN_COMM\DeltaOne\Kenny\PCA), or into `out_dir` when given. The
    document carries the verdict, the trade as a table, risk and monitoring
    levels, the out-of-sample evidence, the exact settings, and the full
    text report as an appendix. `also_csv` additionally drops weights.csv
    next to it for execution/upload.

    Needs python-docx (pip install python-docx); without it the report is
    written as .txt instead so an export never silently fails. Returns the
    path actually written."""
    import os as _o
    # [Fix 115] In a notebook the report renders as HTML and never reaches
    # stdout, so a caller that captured stdout handed us only the progress
    # lines. The renderers keep their own text copy; prefer it.
    report_text = report_text or transcript_text()
    folder, warn = _export_folder(out_dir)
    if warn:
        print(f"  ⚠ {warn}")
    label = _slug(result.get('target_label') or 'hedge')
    stamp = datetime.now().strftime('%Y%m%d-%H%M')
    stem = _o.path.join(folder, f"PCA_hedge_{label}_{stamp}")
    written = None
    try:
        written = _docx_hedge_report(result, report_text, stem + '.docx')
    except ImportError:
        print("  ⚠ python-docx not installed (pip install python-docx) — "
              "writing a text report instead")
    except Exception as e:
        print(f"  ⚠ Word export failed ({e}) — writing a text report instead")
    if written is None:
        written = stem + '.txt'
        with open(written, 'w', encoding='utf-8') as fh:
            fh.write(report_text or 'no report text captured')
    if also_csv:
        try:
            w = result['weights']
            pd.DataFrame({
                'ticker': w.index,
                'action': ['SHORT' if x > 0 else 'LONG' for x in w.values],
                'weight_per_100_gross': w.values.round(6),
                'notional_mm': ([abs(x) * result['params']['notional_mm']
                                 for x in w.values]
                                if (result.get('params') or {}).get('notional_mm')
                                else [None] * len(w)),
                'forced': [t in (result.get('force_included') or [])
                           for t in w.index],
            }).to_csv(stem + '.csv', index=False)
        except Exception as e:
            print(f"  ⚠ weights CSV not written ({e})")
    print(f"  ✓ exported → {written}")
    return written


def hedge_ui(data=None):
    """[Fix 89/95] Point-and-click front end. Run this cell, use the form.

    `data` optionally seeds the form with an already-downloaded dataset;
    otherwise use the Data tab (download once → cached to disk forever)."""
    try:
        import ipywidgets as W
        from IPython.display import display, clear_output, HTML as _H
    except Exception:
        print("ipywidgets is not installed — the form cannot be shown.\n"
              "  install:  pip install ipywidgets    (then restart the "
              "kernel)\n\nScripted path meanwhile:\n"
              "  data   = download_hedge_data('AAPL US Equity', "
              "min_mktcap_usd_mm=1000)\n"
              "  result = find_best_hedge('AAPL US Equity', data=data, "
              "notional_mm=10, hedge_horizon_days=21)")
        return None

    if data is not None:
        UI_STATE['data'] = data

    LBL = {'description_width': '148px'}
    FULL = W.Layout(width='460px')
    HALF = W.Layout(width='225px')

    def H(txt):
        return W.HTML(f"<div class='hui-h'>{txt}</div>")

    def HELP(txt):
        return W.HTML(f"<div class='hui-help'>{txt}</div>")

    # ── Trade tab ───────────────────────────────────────────────────────────
    w_target = W.Textarea(
        value='AAPL US Equity',
        placeholder=("single name:   AAPL US Equity\n"
                     "book (one leg per line, weight optional):\n"
                     "AAPL US Equity, 50\nNVDA US Equity, 30\n"
                     "MSFT US Equity, -20"),
        description='Target / book', style=LBL,
        layout=W.Layout(width='460px', height='96px'))
    w_parsed = W.HTML()
    # ── [Fix 115] CUSTOM VALUES YOU CAN ACTUALLY TYPE INTO ──────────────────
    # Reported: "I can select custom for hedge duration and analysis window
    # but at least please let me enter WHAT custom? After running once the
    # custom box pops up — this is not good."
    # The boxes were created with layout.display='none' and only revealed by an
    # observer, so picking 'Custom…' left nothing to type into until some later
    # event forced a re-layout — by which time the run had already gone off
    # with the wrong horizon. Hiding an input to express "not applicable" is
    # the bug: it hides the ONE control the user just asked for.
    # They are now ALWAYS visible and always show the value that will actually
    # be used. Pick a preset and the box mirrors it (read-only); pick 'Custom…'
    # and the same box becomes editable. Nothing appears, disappears or moves.
    w_horizon = W.Dropdown(
        options=[('— not set —', 0), ('1 week (5d)', 5), ('2 weeks (10d)', 10),
                 ('1 month (21d)', 21), ('2 months (42d)', 42),
                 ('3 months (63d)', 63), ('Custom…', -1)],
        value=21, description='Hedge duration', style=LBL, layout=FULL)
    w_horizon_n = W.BoundedIntText(
        value=21, min=1, max=756, step=1, description='trading days',
        style=LBL, layout=HALF)
    w_horizon_hint = W.HTML()
    w_win = W.Dropdown(
        options=[('Full downloaded history', 'none'), ('Last 3 months', '3m'),
                 ('Last 6 months', '6m'), ('Last 1 year', '1y'),
                 ('Last 2 years', '2y'), ('Custom dates…', 'custom')],
        value='none', description='Analysis window', style=LBL, layout=FULL)
    w_from = W.DatePicker(description='from', style=LBL, layout=HALF)
    w_to = W.DatePicker(description='to', style=LBL, layout=HALF)
    w_dates = W.HBox([w_from, w_to])
    w_dates_hint = W.HTML()
    w_tier = W.HTML()
    w_notional = W.FloatText(value=10.0, description='Book size ($mm)',
                             style=LBL, layout=FULL)

    tab_trade = W.VBox([
        H('What to hedge'), w_target, w_parsed,
        H('How long you will hold the hedge'),
        HELP('risk is quoted over this hold, and the ticket reports what holds '
             'of that length actually delivered'),
        w_horizon, w_horizon_n, w_horizon_hint,
        H('Which period to analyse'),
        HELP('short windows can only support weaker validation — the line '
             'below tells you what the chosen window buys BEFORE you run'),
        w_win, w_dates, w_dates_hint, w_tier,
        H('Sizing'), w_notional])

    # ── Universe tab ────────────────────────────────────────────────────────
    w_mktcap = W.FloatText(value=1000.0, description='Min mkt cap ($mm)',
                           style=LBL, layout=FULL)
    w_basket = W.Text(value='3,5,7', description='Basket sizes', style=LBL,
                      layout=FULL)
    w_look = W.Text(value='126,252', description='Lookbacks (days)',
                    style=LBL, layout=FULL)
    w_force = W.Text(value='', placeholder='0050 TT, 2330 TT',
                     description='Force include', style=LBL, layout=FULL)
    w_excl = W.Text(value='', placeholder='002230, 3156 JP Equity',
                    description='Exclude', style=LBL, layout=FULL)
    w_tz = W.Dropdown(options=[('Auto', 'auto'),
                               ('Force same timezone', True),
                               ('Allow global', False)],
                      value='auto', description='Timezone', style=LBL,
                      layout=FULL)
    w_stale = W.IntText(value=7, description='Stale-price days', style=LBL,
                        layout=FULL)
    tab_uni = W.VBox([
        H('Candidate filters'), w_mktcap, w_tz, w_stale,
        H('Basket construction'), w_basket, w_look,
        H('Manual overrides'),
        HELP('force include = held in every basket (★); exclude = never a '
             'candidate. Short names or full tickers.'),
        w_force, w_excl])

    # ── Model tab (factor / fitting knobs) ─────────────────────────────────
    w_cov = W.Dropdown(options=[('Ledoit-Wolf shrinkage', 'ledoit'),
                                ('Ordinary PCA', 'none')],
                       value='ledoit', description='PCA covariance',
                       style=LBL, layout=FULL)
    w_pcs = W.Text(value='auto', description='PCs (auto or list)',
                   style=LBL, layout=FULL)
    w_wins = W.FloatText(value=0.01, step=0.005,
                         description='Winsorize (0=off)', style=LBL,
                         layout=FULL)
    w_gross = W.FloatText(value=0.02, step=0.01,
                          description='Gross penalty', style=LBL, layout=FULL)
    w_maxgross = W.FloatText(value=0.0, description='Max gross (0=off)',
                             style=LBL, layout=FULL)
    w_boot = W.IntText(value=2000, description='Bootstrap draws',
                       style=LBL, layout=FULL)
    w_report = W.Dropdown(options=[('Compact', 'compact'), ('Full', 'full')],
                          value='compact', description='Report', style=LBL,
                          layout=FULL)
    w_plots = W.Checkbox(value=True, description='Charts', indent=False)
    w_roll = W.Checkbox(value=True, description='Rolling OOS + cadence',
                        indent=False)
    # [Fix 115] "after rehedge... only one outer window?" — that was hard-wired
    # and unexplained. It is a real trade-off (seconds vs a full re-validation),
    # so it becomes a choice, and the report states which one ran.
    w_requick = W.Dropdown(
        options=[('Fast — 1 validation window', 'fast'),
                 ('Full — same 3 windows as RUN HEDGE', 'full')],
        value='fast', description='RE-HEDGE depth', style=LBL, layout=FULL)
    tab_model = W.VBox([
        H('Factor model'),
        HELP('these shape candidate SELECTION only; hedge weights always '
             'come from ridge on real returns'),
        w_cov, w_pcs, w_wins,
        H('Penalties'), w_gross, w_maxgross,
        H('Output'), w_boot, w_report, W.HBox([w_plots, w_roll]),
        H('Re-hedging'),
        HELP('RE-HEDGE never re-searches the config — that is what makes it '
             'fast. This only sets how many windows it re-validates the refit '
             'on. Charts and Rolling OOS above apply to RE-HEDGE too.'),
        w_requick])

    # ── Data tab ────────────────────────────────────────────────────────────
    w_cache = W.Dropdown(options=[('(no cached downloads)', None)],
                         description='Cached data', style=LBL,
                         layout=W.Layout(width='620px'))
    b_load = W.Button(description='Load selected', icon='folder-open')
    b_refresh = W.Button(description='Refresh list', icon='refresh')
    w_years = W.IntText(value=3, description='Years of history', style=LBL,
                        layout=FULL)
    w_pit1 = W.Checkbox(value=True, indent=False,
                        description='Point-in-time market caps')
    w_pit2 = W.Checkbox(value=True, indent=False,
                        description='Point-in-time index membership')
    b_dl = W.Button(description='Download from Bloomberg',
                    button_style='warning', icon='cloud-download',
                    layout=W.Layout(width='240px'))
    w_expdir = W.Text(value=HEDGE_EXPORT_DIR, description='Export folder',
                      style=LBL, layout=W.Layout(width='620px'))
    w_expcsv = W.Checkbox(value=True, indent=False,
                          description='also write weights .csv')
    tab_data = W.VBox([
        H('Cached downloads — reload with zero API calls'),
        HELP('✓ fits the current book · △ stale prices · ✗ a leg is missing'),
        W.HBox([w_cache]), W.HBox([b_load, b_refresh]),
        H('New download — the only action that touches Bloomberg'),
        HELP('discovers the universe for the book in the Trade tab, pulls '
             'prices and caps, and caches everything to disk automatically'),
        w_years, W.HBox([w_pit1, w_pit2]), b_dl,
        H('Where Export writes'),
        HELP('the Export button writes a Word report (plus weights .csv) '
             'here. If the drive is not mapped on this machine it falls '
             'back to a local hedge_reports folder and tells you.'),
        w_expdir, w_expcsv])

    tabs = W.Tab(children=[tab_trade, tab_uni, tab_model, tab_data])
    for i, t in enumerate(('Trade', 'Universe', 'Model', 'Data')):
        tabs.set_title(i, t)

    w_fit = W.HTML()
    w_status = W.HTML()
    b_run = W.Button(description='RUN HEDGE', button_style='success',
                     icon='play',
                     layout=W.Layout(width='190px', height='42px'))
    b_quick = W.Button(description='RE-HEDGE (fast)', icon='bolt',
                       tooltip='refit weights only, pinned to the last '
                               'validated config — seconds, no API calls',
                       layout=W.Layout(width='170px', height='42px'))
    b_export = W.Button(description='Export Word', icon='file-word-o',
                        tooltip='writes a Word report of this hedge to the '
                                'Export folder on the Data tab',
                        layout=W.Layout(width='150px', height='42px'))
    b_code = W.Button(description='Show code', icon='code',
                      layout=W.Layout(width='130px', height='42px'))
    out = W.Output()

    # ── helpers ─────────────────────────────────────────────────────────────
    # [Fix 115] The number box is the single source of truth for the horizon:
    # a preset writes into it, Custom… lets you edit it, and _horizon_arg only
    # ever reads it. That is what makes the displayed value and the value the
    # engine receives impossible to disagree about.
    _SYNC = {'busy': False}          # guard: preset → box → preset feedback

    def _horizon_arg():
        if w_horizon.value == 0:
            return None
        return int(w_horizon_n.value) or None

    def _window_arg():
        if w_win.value == 'none':
            return None
        if w_win.value == 'custom':
            if not (w_from.value or w_to.value):
                return None
            return (str(w_from.value) if w_from.value else None,
                    str(w_to.value) if w_to.value else None)
        return w_win.value

    def _on_horizon_preset(_=None):
        """Preset picked → mirror it into the box and lock the box."""
        if _SYNC['busy']:
            return
        _SYNC['busy'] = True
        try:
            custom = (w_horizon.value == -1)
            w_horizon_n.disabled = not custom
            if w_horizon.value > 0:
                w_horizon_n.value = int(w_horizon.value)
        finally:
            _SYNC['busy'] = False
        _horizon_hint()

    def _on_horizon_typed(_=None):
        """Typed a number → switch the dropdown to Custom… unless it matches
        a preset. Typing is an intent, so it wins."""
        if _SYNC['busy'] or w_horizon.value == 0:
            return
        _SYNC['busy'] = True
        try:
            _preset = {5, 10, 21, 42, 63}
            w_horizon.value = (int(w_horizon_n.value)
                               if int(w_horizon_n.value) in _preset else -1)
            w_horizon_n.disabled = (w_horizon.value != -1)
        finally:
            _SYNC['busy'] = False
        _horizon_hint()

    def _horizon_hint(_=None):
        h = _horizon_arg()
        if not h:
            w_horizon_hint.value = (
                "<div class='hui-help'>no horizon set — risk is quoted "
                "annualised and the per-hold evidence is skipped</div>")
            return
        w_horizon_hint.value = (
            f"<div class='hui-help'>using <b>{h} trading days</b> "
            f"(≈{h/21:.1f} months)"
            + ('' if w_horizon.value == -1 else
               ' — choose <b>Custom…</b> above to type your own')
            + "</div>")

    def _dates_hint(_=None):
        custom = (w_win.value == 'custom')
        w_from.disabled = w_to.disabled = not custom
        if not custom:
            w_dates_hint.value = (
                "<div class='hui-help'>choose <b>Custom dates…</b> above to "
                "enable these and type an exact range</div>")
            return
        if not (w_from.value or w_to.value):
            w_dates_hint.value = (
                "<div class='hui-help' style='color:#8a5a00'>pick a "
                "<b>from</b> and/or <b>to</b> date — until you do, the full "
                "downloaded history is used</div>")
            return
        w_dates_hint.value = (
            f"<div class='hui-help'>window: "
            f"<b>{w_from.value or 'start of data'}</b> → "
            f"<b>{w_to.value or 'end of data'}</b></div>")

    def _seed_dates():
        """Pre-fill the pickers with the loaded data's real span, so 'Custom
        dates…' starts from the range that actually exists rather than empty
        boxes with no clue what is valid."""
        d = UI_STATE.get('data')
        if d is None or w_from.value or w_to.value:
            return
        try:
            idx = d['prices'].index
            w_from.value, w_to.value = idx[0].date(), idx[-1].date()
        except Exception:
            pass

    def _status(msg, kind='info'):
        col = {'info': '#24435f', 'ok': '#1b6b3a', 'warn': '#8a5a00',
               'err': '#8c1c1c', 'busy': '#8a5a00'}[kind]
        _spin = ('<span style="display:inline-block;margin-right:6px">⏳</span>'
                 if kind == 'busy' else '')
        w_status.value = (f"<div style='color:{col};font-family:{_RF_UI};"
                          f"font-size:12px;margin-top:4px;font-weight:"
                          f"{700 if kind in ('busy', 'err') else 400}'>"
                          f"{_spin}{msg}</div>")

    # ── [Fix 106] ONE ACTION AT A TIME ──────────────────────────────────────
    # Reported: "I clicked several times of running or downloading, it re-ran
    # and re-downloaded — I thought it was not running so it kept running."
    # Two separate causes, both fixed here:
    #   1. NOTHING WAS DISABLED. ipywidgets fires on_click once per click and
    #      queues the callbacks, so a second click during a 4-minute run
    #      scheduled a SECOND full run (and, on the Data tab, a second
    #      Bloomberg download — the expensive one).
    #   2. NOTHING MOVED WHILE IT RAN. _run_engine redirected the engine's
    #      stdout into a StringIO and printed the whole thing at the END, so
    #      between "running…" and the report there were minutes of a blank
    #      cell. There was no way to tell a running job from a dead one, so
    #      clicking again was the reasonable thing to do.
    # Now: a busy flag rejects re-entry, every action button greys out for
    # the duration, and the engine's output streams line by line as it is
    # produced with an elapsed-time counter.
    _ACTION_BTNS = []                     # filled in after the buttons exist

    def _set_busy(on, label=''):
        UI_STATE['busy'] = bool(on)
        for _b in _ACTION_BTNS:
            try:
                _b.disabled = bool(on)
            except Exception:
                pass
        if on:
            _status(f"{label} — running… (buttons disabled until it "
                    f"finishes; do NOT click again)", 'busy')

    def _guard(fn, label, name):
        """Wrap a button handler: reject re-entry, disable, always restore."""
        def _cb(_b):
            if UI_STATE.get('busy'):
                with out:
                    print(f"\n[Fix 106] '{name}' ignored — {UI_STATE.get('busy_what', 'a job')} "
                          f"is still running. Watch the elapsed-time counter; "
                          f"the buttons re-enable when it finishes.")
                return
            UI_STATE['busy_what'] = name
            _set_busy(True, label)
            try:
                fn(_b)
            finally:
                _set_busy(False)
                UI_STATE['busy_what'] = None
        return _cb

    def _best_cache_for(target):
        try:
            cat = list_hedge_data()
        except Exception:
            return None
        for _, r in cat.iterrows():
            fp = r.get('_fp')
            if not fp:
                continue
            try:
                if check_cache_fit(target,
                                   fingerprint=fp)['status'] != 'mismatch':
                    return r['path']
            except Exception:
                continue
        return None

    def _refresh_cache(_=None):
        try:
            cat = list_hedge_data()
        except Exception:
            cat = pd.DataFrame()
        if len(cat):
            try:
                tgt = _ui_parse_book(w_target.value)
            except Exception:
                tgt = None
            opts = []
            for _, r in cat.iterrows():
                mark = ''
                if tgt is not None and r.get('_fp'):
                    try:
                        st = check_cache_fit(
                            tgt, fingerprint=r['_fp'])['status']
                        mark = {'exact': '✓ ', 'compatible': '✓ ',
                                'stale': '△ ', 'mismatch': '✗ '}[st]
                    except Exception:
                        pass
                opts.append((f"{mark}{r['label']}  ·  to "
                             f"{r['last_date'] or r['saved']}  ·  "
                             f"{r['n_tickers'] or '?'} tickers  ·  "
                             f"{r['MB']}MB", r['path']))
            w_cache.options = opts
        else:
            w_cache.options = [('(no cached downloads)', None)]

    def _refresh_parsed(_=None):
        try:
            tgt = _ui_parse_book(w_target.value)
            legs, wts, isp, lab = _coerce_targets(tgt)
            if isp:
                body = ' · '.join(f"{k.split(' ')[0]} {v:+.0%}"
                                  for k, v in wts.items())
                w_parsed.value = (f"<span style='color:#1b6b3a;font-size:"
                                  f"12px'>book {lab}: {body} "
                                  f"(gross-normalized)</span>")
            else:
                w_parsed.value = (f"<span style='color:#1b6b3a;font-size:"
                                  f"12px'>single name: {legs[0]}</span>")
        except Exception as e:
            w_parsed.value = (f"<span style='color:#8c1c1c;font-size:12px'>"
                              f"{e}</span>")
        _refresh_fit()

    def _refresh_fit(_=None):
        if UI_STATE['data'] is None:
            w_fit.value = ("<div style='color:#8a5a00;font-size:12px'>"
                           "no data loaded — open the <b>Data</b> tab</div>")
            return
        try:
            tgt = _ui_parse_book(w_target.value)
        except Exception:
            w_fit.value = ''
            return
        try:
            fit = check_cache_fit(tgt, data=UI_STATE['data'])
        except Exception as e:
            w_fit.value = (f"<div style='color:#8c1c1c;font-size:12px'>{e}"
                           f"</div>")
            return
        col = {'exact': '#1b6b3a', 'compatible': '#1b6b3a',
               'stale': '#8a5a00', 'mismatch': '#8c1c1c'}[fit['status']]
        mark = {'exact': '✓', 'compatible': '✓', 'stale': '△',
                'mismatch': '✗'}[fit['status']]
        extra = ''
        if fit['status'] == 'mismatch':
            hit = _best_cache_for(tgt)
            extra = ('<br>a cached set that fits: <b>'
                     + _os.path.basename(hit) + '</b> — pick it in the Data '
                     'tab' if hit else
                     '<br>press <b>Download from Bloomberg</b> in the Data '
                     'tab')
        w_fit.value = (f"<div style='color:{col};font-size:12px;"
                       f"font-family:Menlo,monospace;margin-top:4px'>{mark} "
                       f"data: {fit['message']}{extra}</div>")

    def _refresh_tier(_=None):
        # [Fix 115] no widget is shown or hidden here any more — the custom
        # inputs are permanent and _on_horizon_preset / _dates_hint enable or
        # disable them in place.
        if UI_STATE['data'] is None:
            w_tier.value = ("<span style='color:#8a5a00;font-size:12px'>"
                            "load or download data to preview the "
                            "validation tier</span>")
            return
        try:
            import io as _io
            from contextlib import redirect_stdout as _rso
            h = _horizon_arg()
            cvw = int(np.clip(h, 10, 63)) if h else 42
            with _rso(_io.StringIO()):
                pv = preview_window(
                    UI_STATE['data'], _window_arg(),
                    lookback_options=_ui_ints(w_look.value, (126, 252)),
                    cv_test_window=cvw,
                    basket_size_options=_ui_ints(w_basket.value, (3, 5, 7)))
            mode = pv.get('mode')
            if mode is None:
                w_tier.value = ("<span style='color:#8c1c1c;font-size:12px'>"
                                f"only {pv['sessions']} sessions — need ≥10"
                                "</span>")
                return
            col = {'FULL': '#1b6b3a', 'REDUCED': '#1b6b3a',
                   'HOLDOUT': '#8a5a00', 'IN-SAMPLE': '#8c1c1c'}[mode]
            w_tier.value = (
                f"<div style='color:{col};font-size:12px;font-family:Menlo,"
                f"monospace'>{pv['sessions']} sessions "
                f"({pv['start']:%d %b %Y} → {pv['end']:%d %b %Y}) → "
                f"<b>{mode}</b><br><span style='color:#555'>"
                f"{TIER_MEANING[mode]}</span></div>")
        except Exception as e:
            w_tier.value = (f"<span style='color:#8c1c1c;font-size:12px'>"
                            f"{e}</span>")

    def _kwargs():
        mg = w_maxgross.value or None
        pcs = str(w_pcs.value).strip().lower()
        return {'min_mktcap_usd_mm': w_mktcap.value,
                'notional_mm': (w_notional.value or None),
                'hedge_horizon_days': _horizon_arg(),
                'basket_size_options': _ui_ints(w_basket.value, (3, 5, 7)),
                'lookback_options': _ui_ints(w_look.value, (126, 252)),
                'pc_options': ('auto' if pcs in ('', 'auto')
                               else list(_ui_ints(pcs, (3,)))),
                'winsorize_pct': w_wins.value,
                'force_include': _ui_list(w_force.value),
                'exclude': _ui_list(w_excl.value),
                'restrict_to_target_tz': w_tz.value,
                'robust_cov': w_cov.value,
                'gross_penalty': w_gross.value,
                'max_gross': (mg if mg and mg > 0 else None),
                'mask_stale_days': (w_stale.value or None),
                'boot_n': int(w_boot.value),
                'report': w_report.value,
                'rolling_oos': w_roll.value,
                'show_plots': w_plots.value}

    def _run_engine(fn, *args, **kw):
        """[Fix 106] Run the engine with its output STREAMING to the cell.

        The old version redirected stdout into a StringIO for the whole run
        and printed it at the end, purely so the text could be captured for
        export. On a multi-minute run that produced a silent cell — the
        direct cause of the repeated clicking. A real tee keeps the capture
        AND shows progress as it happens; the HTML blocks the report emits
        via _ipy_display are unaffected (they were never on stdout).
        """
        import io as _io, sys as _sys, time as _time
        buf = _io.StringIO()

        class _Tee:
            def __init__(self, *sinks):
                self._sinks = sinks
            def write(self, s):
                for _s in self._sinks:
                    try:
                        _s.write(s)
                    except Exception:
                        pass
                return len(s)
            def flush(self):
                for _s in self._sinks:
                    try:
                        _s.flush()
                    except Exception:
                        pass

        _t0 = _time.time()
        _old = _sys.stdout
        _sys.stdout = _Tee(_old, buf)
        try:
            res = fn(*args, **kw)
        finally:
            _sys.stdout = _old
        # [Fix 115] in a notebook the report is displayed as HTML and never
        # goes near stdout, so the tee only ever caught stray print()s. The
        # renderers keep a faithful text copy — use it, and keep the tee as
        # the fallback for anything printed outside them.
        UI_STATE['report_text'] = transcript_text() or buf.getvalue()
        UI_STATE['last_seconds'] = _time.time() - _t0
        return res

    # ── actions ─────────────────────────────────────────────────────────────
    def on_load(_):
        with out:
            clear_output()
            if not w_cache.value:
                _status('nothing cached yet — download first', 'warn')
                return
            try:
                UI_STATE['data'] = load_hedge_data(w_cache.value)
                UI_STATE['loaded_from'] = w_cache.value
                _status(f"loaded · {UI_STATE['data']['prices'].shape[0]} "
                        f"rows × {UI_STATE['data']['prices'].shape[1]} "
                        f"tickers (no API calls)", 'ok')
                _refresh_tier()
                _refresh_fit()
            except Exception as e:
                _status(f"load failed: {e}", 'err')

    def on_download(_):
        with out:
            clear_output()
            try:
                tgt = _ui_parse_book(w_target.value)
            except Exception as e:
                _status(f"target: {e}", 'err')
                return
            _status('downloading from Bloomberg — this can take minutes…',
                    'info')
            try:
                UI_STATE['data'] = download_hedge_data(
                    tgt, years=int(w_years.value),
                    min_mktcap_usd_mm=w_mktcap.value,
                    exclude=_ui_list(w_excl.value),
                    force_include=_ui_list(w_force.value),
                    fetch_mktcap_hist=w_pit1.value,
                    fetch_pit_members=w_pit2.value)
                UI_STATE['loaded_from'] = None
                _status(f"downloaded and cached · "
                        f"{UI_STATE['data']['prices'].shape[0]} rows × "
                        f"{UI_STATE['data']['prices'].shape[1]} tickers",
                        'ok')
                _refresh_cache()
                _refresh_tier()
                _refresh_fit()
            except Exception as e:
                _status(f"download failed: {e}", 'err')

    def _pre_run():
        if UI_STATE['data'] is None:
            _status('no data yet — use the Data tab to download or load a '
                    'cached set', 'warn')
            return None
        try:
            tgt = _ui_parse_book(w_target.value)
        except Exception as e:
            _status(f"target: {e}", 'err')
            return None
        fit = check_cache_fit(tgt, data=UI_STATE['data'])     # [Fix 90]
        if fit['status'] == 'mismatch':
            _status(f"cannot run — {fit['message']}", 'err')
            return None
        if fit['status'] == 'stale':
            print(f"  △ {fit['message']}\n")
        return tgt

    def on_run(_):
        with out:
            clear_output()
            tgt = _pre_run()
            if tgt is None:
                return
            kw = _kwargs()
            win = _window_arg()
            UI_STATE['code'] = _ui_code(
                tgt, win, {k: v for k, v in kw.items()
                           if k != 'show_plots'}, UI_STATE['loaded_from'])
            _status('running…', 'info')
            try:
                UI_STATE['result'] = _run_engine(
                    find_best_hedge, tgt, data=UI_STATE['data'],
                    sample_window=win, **kw)
                _status(f"done in {UI_STATE.get('last_seconds', 0):.0f}s · "
                        f"UI_STATE['result'] holds everything — Export "
                        f"writes the Word report", 'ok')
            except Exception as e:
                _status(f"run failed: {e}", 'err')
                print(f"\n{type(e).__name__}: {e}")

    def on_quick(_):
        with out:
            clear_output()
            if UI_STATE['result'] is None:
                _status('run a full hedge first — RE-HEDGE pins its '
                        'validated config', 'warn')
                return
            tgt = _pre_run()
            if tgt is None:
                return
            _full = (w_requick.value == 'full')
            _status('re-hedging on the pinned config…', 'info')
            try:
                # [Fix 115] the Charts and Rolling-OOS checkboxes apply here
                # too. They used to be forced off inside quick_rehedge, so a
                # re-hedge silently threw away every chart the full run drew.
                UI_STATE['result'] = _run_engine(
                    quick_rehedge, tgt, UI_STATE['data'],
                    UI_STATE['result'], show_plots=w_plots.value,
                    rolling_oos=w_roll.value, full_validation=_full)
                _status('re-hedged · weights refit, config unchanged · '
                        + ('full validation re-run' if _full else
                           '1 validation window (fast) — see the note at the '
                           'top of the report'), 'ok')
            except Exception as e:
                _status(f"re-hedge failed: {e}", 'err')

    def on_export(_):
        with out:
            if UI_STATE['result'] is None:
                _status('nothing to export yet', 'warn')
                return
            try:
                path = export_result(UI_STATE['result'],
                                     UI_STATE.get('report_text', ''),
                                     out_dir=(w_expdir.value.strip() or None),
                                     also_csv=w_expcsv.value)
                _status(f"exported → {path}", 'ok')
            except Exception as e:
                _status(f"export failed: {e}", 'err')

    def on_code(_):
        with out:
            clear_output()
            if not UI_STATE['code']:
                _status('run something first — then this shows the exact '
                        'code for it', 'warn')
                return
            print(UI_STATE['code'])

    # [Fix 106] every action goes through _guard: single-flight, buttons
    # disabled for the duration, elapsed time reported. b_refresh and b_code
    # are instant and read-only, so they stay unguarded.
    _ACTION_BTNS.extend([b_load, b_dl, b_run, b_quick, b_export])
    b_load.on_click(_guard(on_load, 'loading cache', 'Load selected'))
    b_refresh.on_click(_refresh_cache)
    b_dl.on_click(_guard(on_download, 'downloading from Bloomberg',
                         'Download from Bloomberg'))
    b_run.on_click(_guard(on_run, 'running the hedge', 'RUN HEDGE'))
    b_quick.on_click(_guard(on_quick, 're-hedging', 'RE-HEDGE (fast)'))
    b_export.on_click(_guard(on_export, 'writing the report', 'Export Word'))
    b_code.on_click(on_code)
    for wdg in (w_win, w_from, w_to, w_look, w_basket, w_horizon,
                w_horizon_n):
        wdg.observe(_refresh_tier, names='value')
    # [Fix 115] preset ⇄ number box, and the date pickers' enabled state
    w_horizon.observe(_on_horizon_preset, names='value')
    w_horizon_n.observe(_on_horizon_typed, names='value')
    w_win.observe(_dates_hint, names='value')
    for wdg in (w_from, w_to):
        wdg.observe(_dates_hint, names='value')
    w_target.observe(_refresh_parsed, names='value')

    _refresh_cache()
    _refresh_parsed()
    _on_horizon_preset()
    _seed_dates()
    _dates_hint()
    _refresh_tier()
    _refresh_fit()
    root = W.VBox([
        W.HTML(_UI_CSS),
        W.HTML("<div class='hui-title'>PCA Hedge Finder</div>"
               "<div class='hui-sub'>Set the trade up, press "
               "<b>RUN</b>. Results stay in <code>UI_STATE['result']</code>; "
               "<b>Export Word</b> writes the hedge report to the "
               "folder set on the Data tab.</div>"),
        tabs, w_fit,
        W.HBox([b_run, b_quick, b_export, b_code],
               layout=W.Layout(margin='12px 0 2px 0')),
        w_status, out])
    root.add_class('hedge-ui')
    display(root)
    return UI_STATE


# =============================================================================
# CELL 4 (OPTIONAL): SCRIPTED API — for power users and reproducibility
# =============================================================================
# The GUI in CELL 3 covers everything below; this cell exists so a run can be
# scripted, diffed or scheduled. Everything is commented out on purpose — the
# "Show code" button in the form prints the exact equivalent of whatever you
# just ran, ready to paste here.
#
#   preview_window(data, '3m')        # what will this window buy me?
#   list_hedge_data()                 # what downloads are cached?
#   data = load_hedge_data(0)         # reload the newest, zero API calls
#   check_cache_fit(book, data=data)  # is that download still right for me?
# =============================================================================

# =============================================================================
# CELL 3: RUN — Jupyter workflow [Fix 31]
# =============================================================================
# --- CELL A: download ONCE (the only cell that hits Bloomberg) ---------------
# Discovery + prices + USD market caps. min_mktcap_usd_mm is in USD MILLIONS
# and only shapes the DISCOVERED universe; you can tighten it again in CELL B.
# [Fix 40] pass force_include here too so the names are guaranteed downloaded.
# [Fix 44] download_hedge_data now ALSO fetches historical market-cap PATHS so
#          the size floor can be applied point-in-time (as-of each window)
#          instead of with today's cap — set fetch_mktcap_hist=False to skip it.
# [Fix 49] It ALSO pulls point-in-time INDEX MEMBERSHIP (INDX_MWEIGHT_HIST):
#          past members incl. delisted names get downloaded, and each backtest
#          window only sees names that were members at the time. If the index
#          licence-gates the field (e.g. some TW benchmarks) it just prints a
#          warning and falls back to today's members — nothing breaks.
#          Set fetch_pit_members=False to skip.
# [Fix 46] `target` may be a single name, an equal-weight list, or a weighted
#          dict (a whole book) — see the portfolio example below.
# data = download_hedge_data('GMD AU Equity', min_mktcap_usd_mm=1000)
# data = download_hedge_data('2303 TT Equity', min_mktcap_usd_mm=0,
#                            force_include=['2330 TT', '0050 TT'])
#
# [Fix 46] PORTFOLIO download — union of each leg's universe, legs auto-excluded
# data = download_hedge_data({'AAPL US Equity': 0.5,     # 50% long AAPL
#                             'MSFT US Equity': 0.3,     # 30% long MSFT
#                             'NVDA US Equity': 0.2},    # 20% long NVDA
#                            min_mktcap_usd_mm=5000)

# --- CELL B: model — rerun freely while tuning, zero API calls ---------------
# [Fix 60] TWO CONFIG TRAPS were live in the previous version of this cell:
#   (a) size_band=(0.7, 1.3) was ACTIVE while its own comment said "0.3x-3x".
#       A 0.7-1.3x band is brutally tight and silently guts the universe —
#       ETFs and force-included names are exempt, but every other candidate
#       had to sit within ±30% of the target's market cap.
#   (b) min_mktcap_usd_mm was commented out, so the function DEFAULT of 5000
#       ($5bn) applied — even though CELL A downloaded with a $1bn floor.
#       Whenever size_band is off, that silent $5bn floor filters everything.
# Never rely on the default here: state the floor explicitly so CELL A and
# CELL B agree, and if you use size_band, keep it wide.
# result = find_best_hedge(
#     'GMD AU Equity', data=data,
#     min_mktcap_usd_mm=1000,         # [Fix 35/60] USD MILLIONS — matches the
#                                     # CELL A download floor. Explicit on
#                                     # purpose; the default is 5000 ($5bn).
#     restrict_to_target_tz='auto',   # [Fix 33] 'auto' / True / False
#     tz_tolerance_hours=3.0,         # [Fix 32] closes within 3h = synchronous
#     # exclude=['002230', '3156 JP Equity'],   # [Fix 36] manual filter-out:
#     #                                 # short names or full tickers, any mix
#     # force_include=['2330 TT'],      # [Fix 40] must-have names — held in
#     #                                 # EVERY basket and marked ★ in all
#     #                                 # outputs; short names or full tickers
#     gross_penalty=0.02,             # [Fix 38] composite penalty per 1.0x gross
#     max_gross=None,                 # [Fix 38] e.g. 3.0 caps gross at $300/$100
#     winsorize_pct=0.01,             # [Fix 42] clip at 1%/99% before PCA
#     robust_cov='ledoit',            # [Fix 42] shrinkage-covariance PCA
#     pit_size=True,                  # [Fix 44] point-in-time size floor
#     pit_members=True,               # [Fix 49] point-in-time index membership
#     mask_stale_days=7,              # [Fix 56] suspension guard — mask prices
#                                     # flat for >7 sessions; None disables
#     notional_mm=None,               # [Fix 72] e.g. 10 → the trade ticket adds
#                                     # a $mm column sized for $10mm gross
#     # sample_window='3m',           # [Fix 76] restrict the analysis window:
#     #                               # '3m'/'90d'/'26w'/'1y' = last N up to the
#     #                               # end of the data, or absolute dates:
#     # sample_window=('2026-07-13', '2026-07-28'),
#     #                               # [Fix 77] validation AUTO-ADAPTS to the
#     #                               # window: FULL → REDUCED → HOLDOUT →
#     #                               # IN-SAMPLE (short windows can only give
#     #                               # descriptive numbers — the ticket says
#     #                               # which tier ran and what it proves)
#     # report='full',                # [Fix 79] long-form report (adds per-
#     #                               # config weight tables + explanations)
#     # rolling_oos=True,             # [Fix 71] ON by default: walk-forward OOS
#     #                               # curve + [Fix 78] rebalance-cadence and
#     #                               # basket-turnover diagnostics
#     # boot_n=2000,                  # [Fix 70] bootstrap CI draws; 0 disables
#     # size_band=(0.3, 3.0),         # [Fix 45/60] OFF by default. If you use
#     #                               # it, keep it WIDE (0.3x-3x the target's
#     #                               # cap). It overrides the flat floor.
#     # size_override={'XYZ': 'large'},  # [Fix 45] forward-looking cap views
# )

# --- Portfolio / multi-name books [Fix 46/64] --------------------------------
# The target can be a book instead of a single name. Three accepted formats:
#
#   result = find_best_hedge('AAPL US Equity', data=data)          # 1 name
#
#   result = find_best_hedge(                       # equal-weight LONG book
#       ['AAPL US Equity', 'NVDA US Equity', 'MSFT US Equity'],
#       data=data)
#
#   result = find_best_hedge(                       # weighted book: values =
#       {'AAPL US Equity': +0.5,                    # relative notionals in
#        'NVDA US Equity': +0.3,                    # any unit ($mm, %, ...);
#        'MSFT US Equity': -0.2},                   # sign = direction.
#       data=data)                                  # Normalized by GROSS.
#
# Weights need NOT sum to 1 — {'AAPL US': 50, 'NVDA US': 30, 'MSFT US': -20}
# (e.g. $mm notionals) is normalized to (+0.5, +0.3, -0.2) internally.
# [Fix 67] every ticker MUST carry an exchange code ('AAPL US' at minimum) —
# a bare 'AAPL' now raises immediately instead of dying mid-download. The book is
# hedged as ONE synthetic return series (the NET book), the recipe's
# "per $100" means per $100 of GROSS book value, and section [2b] shows the
# per-leg residual attribution. NOTE: the legs must be in the CELL A
# download — pass the same book (or its tickers) to download_hedge_data.
#
# --- Re-hedging an updated book without a full re-run [Fix 63/66] ------------
# For an index-rebalance book that churns daily, pin the validated config and
# refit weights only (seconds, no API calls). Re-run the full find_best_hedge
# weekly or whenever the book changes materially.
# [Fix 66] the new book's legs MAY differ from the downloaded book — the only
# requirement is that every leg exists in the downloaded data (new names can
# be guaranteed via force_include= in CELL A). Ex-legs become candidates.
# new_book = {'AAA AU Equity': 0.5, 'BBB AU Equity': -0.5}
# result2 = quick_rehedge(new_book, data, result)

# =============================================================================
# [Fix 107] AUTO-LAUNCH — the form comes up with CELL 1, no second cell
# =============================================================================
# Until now the workflow was "paste CELL 1, run it, then type hedge_ui() in a
# second cell". That second cell is pure ceremony: there is nothing you can
# usefully do with this module that does not start at the form. It also made
# the tool look broken on first use — CELL 1 runs, prints a few import lines,
# and nothing appears.
#
# The form now renders at the end of CELL 1. Guards, in order:
#   · only inside a notebook (a script/CLI import must stay silent);
#   · only when ipywidgets is importable — otherwise hedge_ui() would print
#     its install banner on every import, which is noise for scripted users;
#   · HEDGE_UI_AUTOSTART = False before the cell (or in the same cell, above
#     this line) suppresses it;
#   · exceptions are swallowed to a one-line hint. Auto-launch is a
#     convenience — it must never stop the module from loading, or a widget
#     problem would take the whole scripted API down with it.
# Re-running the cell re-renders the form, which is the expected notebook
# behaviour; UI_STATE (data, result, report) persists across renders, so a
# re-render costs nothing and loses nothing.
HEDGE_UI_AUTOSTART = globals().get('HEDGE_UI_AUTOSTART', True)

if HEDGE_UI_AUTOSTART and _IN_JUPYTER:
    try:
        import ipywidgets as _W_probe          # noqa: F401  (availability probe)
        hedge_ui()
    except ImportError:
        print("[Fix 107] hedge_ui() not auto-started: ipywidgets is missing.\n"
              "          pip install ipywidgets  (then restart the kernel), "
              "or use the scripted path in CELL 4.")
    except Exception as _e_ui:                 # never break the import
        print(f"[Fix 107] hedge_ui() auto-start failed ({type(_e_ui).__name__}"
              f": {_e_ui}). The module is loaded and usable — call hedge_ui() "
              f"yourself, or use the scripted path in CELL 4.")
